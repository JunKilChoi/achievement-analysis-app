# -*- coding: utf-8 -*-
"""
성취수준별 평가결과 분석 웹앱 v1.100

버전 기록
- v1.1: 학생답 정오표 여러 파일 업로드/추가 업로드/중복 제외, 문항정보표 C6에서 선택형·서답형 만점 자동 추출
- v1.2: 화면·AI 프롬프트·엑셀 출력의 정답률/비율/변별도/수준간격차를 % 기준으로 표시
- v1.3: 표 화면/엑셀 출력에서 비율값은 52.3%처럼 데이터에 %를 포함하고, 헤더에 (명)/(점)/(%) 단위 표시
- v1.4: 업로드 후 문항정보표를 웹앱에서 직접 수정하고, 수정값을 분석/엑셀/AI에 반영
- v1.5: 문항정보 수정표의 데이터 타입 충돌로 인한 st.data_editor 오류 수정
- v1.6: 성취도 분석 탭에 학급별 최고·최저·평균 그래프와 선택 학급 성취수준 분포 그래프 추가
- v1.7: matplotlib 의존성 제거, Streamlit 기본 차트로 성취도 분석 그래프 표시
- v1.8: 성취도 분석 탭의 전체 분석 그래프를 최고점·최저점·평균 캔들형으로 변경하고, 전체/개별 반 그래프를 좌우 배치
- v1.9: 전체 분석 캔들형 그래프의 최고·최저 범위선과 평균점을 강조하고, “한 반 분석” 용어를 “개별 반 분석”으로 변경
- v1.10: 전체 분석 그래프에 점수 구간별 도수 폭을 반영한 항아리형 분포를 추가하여 최고·최저·평균과 분포를 함께 표시
- v1.11: 전체 분석 그래프의 Vega-Lite 데이터 연결 방식을 수정하여 항아리형 분포와 평균점이 실제로 표시되도록 보정
- v1.12: 학생답 정오표 다중 업로드를 파일별 독립 처리로 안정화하고, 성공/중복/실패 파일을 개별 표시
- v1.13: 정답을 .으로 표시하는 정오표와 정답 번호로 표시하는 정오표를 모두 지원하고, A~Z 복수답안코드 규약 반영
- v1.14: 성취수준 산출 기준을 환산점수에서 시험지 원점수 기준으로 변경하고, 정오표/정답 표시의 1.0 같은 정수형 소수 표시를 1로 정규화
- v1.15: 전체 분석 항아리형 그래프의 최고점·최저점 범위선 끝에 최고/최저 라벨과 점수 표시 추가
- v1.16: 전체 분석 항아리형 그래프에 평균점수 라벨을 추가하고, 점수 구간 툴팁을 한 줄 표시로 정리
- v1.17: 개별 반 분석 그래프의 학생수 축을 0 이상 고정하고, 확대/축소형 상호작용을 제거하며 성취수준 비율 표시 오류 수정
- v1.18: 문항별 분석 탭에 예상 난이도와 실제 정답률 기준 난이도의 괴리 분석, 기준 조절, 괴리 문항 강조 기능 추가
- v1.19: 난이도 괴리 분석 그래프의 Vega-Lite 데이터 연결 오류 수정
- v1.20: 난이도 괴리 분석을 예상 난이도 기준 구간과 실제 정답률의 차이(%p) 중심으로 단순화
- v1.21: 난이도 괴리 분석 그래프를 제거하고 일치/예상보다 어려웠음/예상보다 쉬웠음만 표시
- v1.22: 난이도 괴리 분석의 예상 난이도를 편집값이 아니라 업로드한 문항정보표 원본 난이도 기준으로 고정
- v1.23: 난이도 괴리 분석에서 체크박스를 제거하고 문항번호 순으로 전체 표시, 일치/불일치/예상보다 어려움/쉬움 개수 요약 추가
- v1.24: 난이도 괴리 분석 기준 입력을 어려움/보통 난이도 구분 정답률, 보통/쉬움 난이도 구분 정답률로 분리
- v1.25: 난이도 일치 여부 영역을 카드형으로 정리하고 기준 입력/요약 지표의 가시성 개선
- v1.26: 데이터 확인·성취도·학급별·평가영역별·성취수준별·학생 개별·AI 분석 탭의 주요 항목을 카드형 박스로 정리
- v1.27: 평가영역별 분석 옆에 성취기준별 분석 탭을 추가하고, 성취기준별 정답률/학생별 점수를 표시
- v1.28: 평가영역별/성취기준별 개인별 분석을 전체 나열 대신 반·학생 선택 후 해당 학생 표만 표시
- v1.29: AI 분석 탭을 기본 분석(통계 기반 해석)과 고급 분석(원안지 기반 심층 해석) 구조로 개편하고, 원안지 PDF 업로드 영역과 분석별 프롬프트 초안을 추가
- v1.30: 기본/고급 AI 분석 응답이 중간에 끊기지 않도록 max_output_tokens를 8000으로 확대
- v1.31: 기본/고급 AI 분석 max_output_tokens를 16000으로 확대
- v1.32: 기본 AI 분석 프롬프트를 통계 기반 진단 중심으로 개편하고, 고급 분석 필요 문항 추천 항목 추가
- v1.33: 기본 AI 분석 프롬프트를 이전 해석형 구조로 복원하고, 선택지 반응 및 오답 경향 분석 항목 추가
- v1.35: requirements.txt에 python-docx 의존성 누락 없이 반영
- v1.36: 기본 AI 분석 결과를 일반 대기형 출력에서 실시간 스트리밍 출력으로 변경하고, 생성 완료 후 Word 다운로드 버튼 표시
- v1.37: 기본 AI 분석 프롬프트 마지막에 원안지 기반 고급 분석이 필요한 문항을 우선순위로 추천하는 항목 추가
- v1.38: AI 분석 결과를 session_state에 저장하여 Word 다운로드 버튼 클릭 시 앱 rerun으로 결과가 초기화되지 않도록 수정
- v1.39: 고급 분석에서 업로드한 원안지 PDF를 OpenAI 파일 입력으로 연결하고, 원안지+통계 기반 심층 분석을 스트리밍 출력/Word 저장하도록 추가

- v1.40: 고급 분석에서 분석 유형별 문항번호 입력 기능을 추가하고, 원안지 기반 학생 개별 문항 분석에서 반·학생 선택 후 해당 학생의 전체 풀이 결과와 원안지를 함께 분석하도록 개선
- v1.41: AI 분석 결과를 현재 분석 조건별로 관리하여, 다운로드 시 초기화되지 않으면서 분석 유형·학생·문항 변경 시 이전 결과가 함께 저장되거나 화면에 남지 않도록 정리
- v1.42: 복수정답 표기(예: 2,3 또는 23)와 학생답 복수답안코드 A~Z를 선택지 조합으로 변환하여 정답 판정/검증/분석에 반영
- v1.43: 서답형이 포함된 시험에서 성취도·총점·평균·성취수준 분석은 정오표의 영역총점을 우선 사용하고, 선택형 문항 계산점수는 검증/문항 분석용으로 분리
- v1.44: 모든 AI 분석에 중점 분석 요청 입력 박스를 추가하고, 입력 시 기존 프롬프트 마지막에 평가 전문가 관점의 추가 분석 요청으로 반영
- v1.45: 문항정보 수정표 편집값을 세션에 저장해 수정 내용이 분석에 확실히 반영되도록 하고, 평가요소 확인용 하이라이트 표 추가
- v1.46: 문항정보 수정 단계에 평가요소를 구체적으로 수정하도록 안내하는 강조 문구 추가
- v1.47: 문항정보 수정 단계의 평가요소 하이라이트 확인용 표 제거
- v1.48: 수정된 문항정보 요약을 성취기준별 평가요소·난이도·총점 구조로 재정리
- v1.49: 수정된 문항정보 요약과 성취기준-평가요소별 상세 구성 표를 줄바꿈형 표로 표시해 긴 셀 내용을 읽기 쉽게 개선
- v1.50: 수정된 문항정보 요약을 표 대신 성취기준별 카드형 목록으로 재구성
- v1.51: 성취기준별 카드형 요약에서 불필요한 “성취기준 1/2” 번호 제목을 제거하고 성취기준 내용을 바로 표시
- v1.52: 문항정보 수정표를 적용 버튼 방식으로 변경하여 평가요소 수정값이 셀 편집 중간에 일부 반영되지 않는 문제를 안정화
- v1.53: 문항정보 수정 방식을 표형 data_editor에서 문항별 가로 입력칸 방식으로 변경하고 문항 추가/삭제 기능 추가
- v1.54: 문항번호 입력칸의 +/- 버튼을 제거하고 헤더 중앙 정렬 및 문항 삭제 후 검증 오류 방지 처리
- v1.55: 문항정보표 업로드 파일을 제거하면 문항정보 수정 세션을 초기화하여 같은 파일을 다시 올릴 때 원본 문항정보를 다시 읽도록 수정
- v1.56: 문항정보 수정 영역에서 문항 추가 버튼을 적용 버튼 옆으로 이동하고 선택 문항 삭제 버튼을 삭제 체크 열 하단으로 배치
- v1.57: 문항정보 수정 영역의 문항정보 수정값 적용, 문항 추가, 선택 문항 삭제 버튼을 한 줄에 나란히 배치
- v1.59: 평가정보 자동 인식값 수정 영역을 즉시 반영 방식에서 적용 버튼 방식으로 변경
- v1.60: 1~4단계 대분류를 시각적으로 더 뚜렷하게 구분하는 단계형 헤더와 설명 문구 추가
- v1.61: 1~4단계 사이에 은은하지만 명확한 가로 구분선을 추가하여 단계 전환을 시각적으로 분리
- v1.62: 성취도 분석 표 아래에 표준편차의 의미와 만점 대비 표준편차 비율 기준 해석 안내 문구 추가
- v1.63: 성취도 분석의 표준편차 해석 안내를 다른 부연 설명과 비슷한 간결한 안내 박스 형태로 조정
- v1.64: 성취도 분석 탭에서 전체 분석 그래프와 개별 반 분석 그래프를 각각 카드형 영역으로 분리하고, 전체/개별 반 성취수준별 비율 표를 그래프 아래에 나란히 표시
- v1.66: 학급별 문항 분석 탭에 학급 간 정답률 차이, 학급별 취약 문항 TOP 5, 문항 선택형 학급별 선택지 반응 비교 추가
- v1.67: 분석 결과 영역의 메인 탭을 한 번에 하나만 렌더링하는 선택형 메뉴로 변경하여 드롭박스 변경 후 탭 내용이 한 페이지에 펼쳐지는 현상 방지
- v1.68: 학급별 문항 분석에서 취약 문항 TOP 5, 선택지 반응 비교, 원자료 보기를 제거하고 핵심 표 중심으로 정리
- v1.69: 학급별 문항 분석 표의 기본 정렬을 정답률 낮은 순에서 문항번호 순으로 변경
- v1.70: 평가영역별 분석의 표 높이를 확대하고 환산점수 의미 안내 문구를 추가
- v1.71: 문항별 분석과 평가영역별 분석 표 높이를 크게 확대하여 한 화면에서 더 많은 행을 확인하도록 조정
- v1.72: 성취도 분석의 개별 반 그래프에 전체 성취수준 비율을 빨간색 기준점으로 함께 표시하여 선택 반과 전체를 비교하도록 개선
- v1.73: 성취도 분석의 개별 반 그래프를 선택 반 파란색 막대와 전체 빨간색 막대가 성취수준별로 나란히 보이도록 변경하고 동적 범례를 추가
- v1.74: 개별 반 성취수준 비교 그래프의 범례를 제거하고 문항별 분석의 난이도 비교 표 높이를 확대
- v1.75: 학급별 문항 분석 표에서 정답률, 변별도, 학급간 최대차 열을 서로 다른 색으로 강조
- v1.76: 학급별 문항 분석 표 강조 대상을 전체 정답률, 변별도, 학급간최대차 세 열로 한정하고 같은 색으로 통일
- v1.77: 문항별 분석 표에서도 평가영역, 전체 정답률, 변별도 세 열을 같은 색으로 강조
- v1.78: 평가영역별 분석 표에서 평균정답률을 숨기고 정답률 열을 강조 표시
- v1.79: 성취기준별 분석 표에서 성취기준과 정답률 열을 강조하고 기본 정렬을 성취기준 오름차순으로 변경
- v1.80: 성취수준별 문항 분석 탭에서 문항번호 순 정렬, 빈 성취수준 표시 보정, 안내 문구를 추가하여 표가 비어 보이는 문제를 개선
- v1.82: 화면 표시용 점수 계열 숫자를 모든 탭에서 소수점 첫째 자리까지 표시하도록 통일
- v1.84: AI 분석의 전체 분석, 학생 개별 분석, 원안지 기반 고급 분석에 분석 요청 프리셋 드롭박스 추가
- v1.85: 학생 개별 분석 프리셋 중 맞춤형 피드백과 다음 시험 학습 전략 조언에 과목 특성을 고려한 학습 방식 개선 중심 안내 반영
- v1.86: AI 분석 내부의 기본/고급 분석 전환을 탭에서 선택형 메뉴로 변경하여 기본 분석 선택 시 고급 분석 영역이 함께 펼쳐지는 현상 수정
- v1.87: 원안지 기반 고급 분석 전용 프리셋이 확실히 표시되도록 상태키를 분리하고 프리셋 안내 문구를 보강
- v1.88: 고급 분석에서 별도의 분석 범위 선택을 제거하고, 원안지 기반 고급 분석 프리셋 중심으로 실행되도록 단순화
- v1.89: Word AI 분석 보고서의 평가 정보 영역을 공문서형 요약 표와 강조 박스 중심 양식으로 개선
- v1.90: 나이스 문항정보표가 없는 학교를 위해 앱 전용 문항정보표 양식 다운로드와 업로드 인식 기능 추가
- v1.95: 파일 업로드 전 나이스 XLS data 다운로드 경로와 강의실별 정오표 저장 안내 문구 추가
- v1.96: 파일 업로드 안내에 나이스 문항정보표가 없는 경우 앱 전용 양식 작성 후 업로드 가능 안내 문구 추가
- v1.97: AI 분석 3종에 기본 프롬프트 보기와 기본 프롬프트 사용 방식 선택 기능 추가
- v1.98: 기본 프롬프트 보기에서 실제 통계 데이터 블록을 제외하고 분석 지시 구조만 표시
- v1.99: 직접 작성 프롬프트 모드에도 분석 요청 프리셋을 추가하고 프롬프트 사용 방식 도움말을 보강
- v1.100: 직접 작성 프롬프트 모드의 빈 입력 안내 박스 제거
- v1.94: 데이터 확인의 학생 정오표에서 학번이 정수형 식별값으로 표시되도록 보정
- v1.83: 성취수준별 문항 분석 표에서 평가영역을 앞쪽에 배치하고 수준간격차 열을 강조 표시
- v1.65: 문항별 분석 탭에 정답률 정렬, 열 제목 클릭 정렬, 변별도 계산식과 해석 기준 안내 문구 추가
- v1.58: 자동 인식 결과에 표시되는 교과목, 학년/학기, 문항수, 학생수, 정오표 파일 수, 만점 정보를 자동 인식값 수정에서 모두 수정할 수 있도록 확장
- v1.34: AI 분석 결과 다운로드를 TXT에서 Word(.docx) 보고서 형식으로 변경하고, 문서 상단에 평가 정보를 자동 삽입

주요 기능
- 나이스 문항정보표 + 학생답 정오표 업로드
- 자동 파싱/검증/점수 계산
- 웹앱 내 분석표 확인
- 확인용 엑셀 및 5종 분석 엑셀 ZIP 다운로드
- OpenAI API 선택 연동: 기본 분석(통계 기반 해석) / 고급 분석(원안지 기반 심층 해석) 초안 생성
"""

from __future__ import annotations

import hashlib
import html
import io
import re
import zipfile
from datetime import datetime
from dataclasses import dataclass
from typing import Any, Dict, List, Optional, Tuple

import numpy as np
import pandas as pd
import streamlit as st
from openpyxl import load_workbook

try:
    from openai import OpenAI
except Exception:  # 배포 환경에서 openai 미설치/오류 시 앱 기본 기능은 계속 사용
    OpenAI = None


APP_VERSION = "v1.100"
MULTI_CODE_MAP = {
    "A": [1, 2], "B": [1, 3], "C": [1, 4], "D": [1, 5], "E": [2, 3],
    "F": [2, 4], "G": [2, 5], "H": [3, 4], "I": [3, 5], "J": [4, 5],
    "K": [1, 2, 3], "L": [1, 2, 4], "M": [1, 2, 5], "N": [1, 3, 4],
    "O": [1, 3, 5], "P": [1, 4, 5], "Q": [2, 3, 4], "R": [2, 3, 5],
    "S": [2, 4, 5], "T": [3, 4, 5], "U": [1, 2, 3, 4], "V": [1, 2, 3, 5],
    "W": [1, 2, 4, 5], "X": [1, 3, 4, 5], "Y": [2, 3, 4, 5], "Z": [1, 2, 3, 4, 5],
}


@dataclass
class ParsedData:
    exam_info: Dict[str, Any]
    question_df: pd.DataFrame
    students_df: pd.DataFrame
    long_df: pd.DataFrame
    validation_df: pd.DataFrame
    answer_key_df: pd.DataFrame
    original_question_df: Optional[pd.DataFrame] = None


# -----------------------------------------------------------------------------
# 기본 유틸
# -----------------------------------------------------------------------------

def clean_text(v: Any) -> str:
    if v is None:
        return ""
    try:
        if pd.isna(v):
            return ""
    except Exception:
        pass
    text = str(v).replace("\u3000", " ").strip()
    text = re.sub(r"\s+", " ", text)
    return text


def to_number(v: Any) -> Optional[float]:
    if v is None or v == "":
        return None
    if isinstance(v, (int, float)) and not pd.isna(v):
        return float(v)
    text = clean_text(v).replace(",", "")
    try:
        return float(text)
    except Exception:
        return None


def to_int_if_possible(v: Any) -> Any:
    n = to_number(v)
    if n is None:
        return clean_text(v)
    if abs(n - int(n)) < 1e-9:
        return int(n)
    return n


def normalize_identifier_text(v: Any) -> str:
    """학번처럼 계산 대상이 아닌 식별값을 화면에서 소수점 없이 표시한다."""
    if v is None:
        return ""
    try:
        if pd.isna(v):
            return ""
    except Exception:
        pass
    if isinstance(v, (int, np.integer)):
        return str(int(v))
    if isinstance(v, (float, np.floating)) and abs(float(v) - int(float(v))) < 1e-9:
        return str(int(float(v)))
    text = clean_text(v)
    if re.fullmatch(r"\d+\.0+", text):
        return text.split(".")[0]
    return text


def normalize_answer_value(v: Any) -> str:
    """정답/표시값 비교용 기본 정규화. 1, 1.0, '1'을 같은 값으로 본다."""
    text = clean_text(v).upper()
    if text == "":
        return ""
    n = to_number(text)
    if n is not None and abs(n - int(n)) < 1e-9:
        return str(int(n))
    return text


def parse_choice_set(v: Any) -> Tuple[int, ...]:
    """선택지 표기를 정렬된 선택지 조합으로 변환한다.

    - 단일 선택: 2, 2.0, '2' -> (2,)
    - 복수 정답: '2,3', '2, 3', '23', '2 3' -> (2, 3)
    - 복수답안코드: A~Z -> 규약에 따른 조합
    """
    text = clean_text(v).upper()
    if text in ["", ".", "-"]:
        return tuple()
    if text in MULTI_CODE_MAP:
        return tuple(MULTI_CODE_MAP[text])

    normalized = normalize_answer_value(text)
    if normalized in MULTI_CODE_MAP:
        return tuple(MULTI_CODE_MAP[normalized])
    if re.fullmatch(r"[1-5]", normalized):
        return (int(normalized),)

    # '2,3', '2, 3', '2번 3번', '23'처럼 들어온 복수정답 표기를 모두 지원한다.
    digits = re.findall(r"[1-5]", normalized)
    if len(digits) >= 2 and len(digits) == len(set(digits)):
        compact_source = re.sub(r"[^1-5]", "", normalized)
        if compact_source == "".join(digits):
            return tuple(sorted(int(d) for d in digits))
    return tuple()


def normalize_choice_answer_value(v: Any) -> str:
    """정답 표시를 분석용 표준 표기로 변환한다. 복수정답은 '2,3'처럼 표시한다."""
    choices = parse_choice_set(v)
    if choices:
        return ",".join(map(str, choices))
    return normalize_answer_value(v)


def normalize_mark_value(v: Any) -> str:
    """정오표 원본 표시용 정규화. 학생답의 A~Z 코드는 원본 코드로 보존한다."""
    text = clean_text(v).upper()
    if text in ["", ".", "-"]:
        return text
    if text in MULTI_CODE_MAP:
        return text
    choices = parse_choice_set(text)
    if choices:
        return ",".join(map(str, choices))
    return normalize_answer_value(text)


def get_selected_display(raw: Any) -> Tuple[str, str]:
    """정오표 셀 값을 선택지 표시와 원본 유형으로 변환한다."""
    text = clean_text(raw).upper()
    if text in ["", "-"]:
        return "무표기", "무표기"
    if text in MULTI_CODE_MAP:
        return ",".join(map(str, MULTI_CODE_MAP[text])), "복수답안"
    choices = parse_choice_set(text)
    if len(choices) >= 2:
        return ",".join(map(str, choices)), "복수답안"
    if len(choices) == 1:
        return str(choices[0]), "단일답안"
    return normalize_answer_value(text), "단일답안"


def is_correct_mark(raw: Any, correct_answer: Any) -> bool:
    """두 종류의 정오표와 복수답안코드를 모두 지원한다.

    1) 정답을 '.'으로 표시하는 정오표
    2) 정답이어도 정답 번호를 그대로 표시하는 정오표
    3) 복수답안코드 A~Z로 표시되는 학생답 정오표
    """
    text = clean_text(raw).upper()
    if text == ".":
        return True
    if text in ["", "-"]:
        return False
    selected_choices = parse_choice_set(text)
    correct_choices = parse_choice_set(correct_answer)
    if selected_choices and correct_choices:
        return selected_choices == correct_choices
    return normalize_answer_value(text) == normalize_answer_value(correct_answer)


def answers_match(a: Any, b: Any) -> bool:
    """문항정보표 정답과 정오표 정답 검증용 비교.

    사용자가 문항을 삭제하거나 정답 칸을 비워 둔 경우처럼 한쪽 값이 없는 행도
    검증표에서 확인 필요로 표시하되 앱 오류로 중단되지 않도록 안전하게 비교한다.
    """
    try:
        a_text = clean_text(a)
        b_text = clean_text(b)
        if a_text == "" or b_text == "":
            return False
        a_choices = parse_choice_set(a_text)
        b_choices = parse_choice_set(b_text)
        if a_choices or b_choices:
            return a_choices == b_choices
        return normalize_answer_value(a_text) == normalize_answer_value(b_text)
    except Exception:
        return False


def is_question_no(v: Any) -> bool:
    n = to_number(v)
    return n is not None and abs(n - int(n)) < 1e-9 and 1 <= int(n) <= 300


def is_student_id(v: Any) -> bool:
    return bool(re.match(r"^\d+\s*/\s*\d+$", clean_text(v)))


def parse_exam_info_from_text(text: str) -> Dict[str, Any]:
    info: Dict[str, Any] = {}
    m = re.search(r"(\d{4})\s*학년도", text) or re.search(r"(\d{4})\s*년도", text)
    if m:
        info["학년도"] = m.group(1)
    m = re.search(r"([1-2])\s*학기", text)
    if m:
        info["학기"] = f"{m.group(1)}학기"
    m = re.search(r"(\d+)\s*학년", text)
    if m:
        info["학년"] = f"{m.group(1)}학년"
    m = re.search(r"(\d+)\s*반", text)
    if m:
        info["학급"] = f"{m.group(1)}반"
    for name in ["중간고사", "기말고사", "1차", "2차", "수행평가"]:
        if name in text:
            info["평가구분"] = name
            break
    return info


def extract_subject_from_question_sheet(rows: List[List[Any]]) -> str:
    for row in rows[:8]:
        for cell in row:
            text = clean_text(cell)
            m = re.search(r"\(\s*([^()]+?)\s*\)\s*과목", text)
            if m:
                return m.group(1).strip()
    return ""


def safe_sheet_name(name: str) -> str:
    name = re.sub(r"[\\/*?:\[\]]", "_", str(name))[:31]
    return name or "Sheet"




def make_question_info_template_xlsx() -> bytes:
    """나이스 문항정보표가 없는 학교에서 직접 입력해 업로드할 수 있는 간편 양식 생성."""
    output = io.BytesIO()
    input_cols = ["문항번호", "평가요소", "성취기준", "난이도", "배점", "정답"]
    input_df = pd.DataFrame([{c: "" for c in input_cols} for _ in range(40)])
    guide_df = pd.DataFrame(
        [
            ["문항번호", "숫자로 입력합니다. 예: 1, 2, 3"],
            ["평가요소", "문항이 실제로 평가하는 개념·자료 해석·사고 과정이 드러나도록 구체적으로 입력합니다."],
            ["성취기준", "문항정보표의 성취기준 또는 교사가 사용하는 성취기준을 입력합니다."],
            ["난이도", "쉬움, 보통, 어려움 중 하나로 입력합니다."],
            ["배점", "문항 배점을 숫자로 입력합니다. 예: 2, 2.5"],
            ["정답", "단일 정답은 1~5, 복수정답은 2,3처럼 쉼표로 입력합니다."],
            ["업로드", "이 파일을 작성한 뒤 앱의 문항정보표 업로드에 그대로 올리면 나이스 문항정보표처럼 인식됩니다."],
        ],
        columns=["항목", "작성 방법"],
    )
    example_df = pd.DataFrame(
        [
            [1, "자료를 바탕으로 물질의 특성 판단하기", "[9과08-01] 물질의 특성의 의미를 알고, 실험을 통해 밀도, 용해도, 녹는점, 끓는점 등을 설명할 수 있다.", "쉬움", 2.0, "4"],
            [2, "그래프 자료에서 용해도 변화 해석하기", "[9과08-01] 물질의 특성의 의미를 알고, 실험을 통해 밀도, 용해도, 녹는점, 끓는점 등을 설명할 수 있다.", "보통", 2.0, "2,3"],
        ],
        columns=input_cols,
    )
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        input_df.to_excel(writer, index=False, sheet_name="문항정보입력", startrow=4)
        guide_df.to_excel(writer, index=False, sheet_name="작성안내")
        example_df.to_excel(writer, index=False, sheet_name="작성예시")

        wb = writer.book
        title_fmt = wb.add_format({"bold": True, "font_size": 15, "font_color": "#1F2937"})
        desc_fmt = wb.add_format({"text_wrap": True, "font_color": "#4B5563"})
        header_fmt = wb.add_format({"bold": True, "bg_color": "#E8F0FE", "border": 1, "align": "center", "valign": "vcenter"})
        input_fmt = wb.add_format({"border": 1, "valign": "top", "text_wrap": True})
        note_fmt = wb.add_format({"bg_color": "#F8FAFC", "border": 1, "text_wrap": True, "valign": "top"})

        ws = writer.sheets["문항정보입력"]
        ws.write(0, 0, "성취수준별 평가결과 분석 웹앱 문항정보표 입력 양식", title_fmt)
        ws.write(1, 0, "나이스 문항정보표를 사용하지 않는 경우 이 양식에 문항정보를 입력한 뒤, 앱의 문항정보표 업로드에 그대로 올려 주세요.", desc_fmt)
        ws.write(2, 0, "필수 입력 열: 문항번호, 평가요소, 성취기준, 난이도, 배점, 정답 / 복수정답은 2,3처럼 입력합니다.", desc_fmt)
        for col, value in enumerate(input_cols):
            ws.write(4, col, value, header_fmt)
        widths = [10, 36, 64, 12, 10, 14]
        for col, width in enumerate(widths):
            ws.set_column(col, col, width, input_fmt)
        ws.freeze_panes(5, 0)
        ws.data_validation(5, 3, 44, 3, {"validate": "list", "source": ["쉬움", "보통", "어려움"]})

        for sheet_name in ["작성안내", "작성예시"]:
            gws = writer.sheets[sheet_name]
            df = guide_df if sheet_name == "작성안내" else example_df
            for col, value in enumerate(df.columns):
                gws.write(0, col, value, header_fmt)
            for col in range(len(df.columns)):
                gws.set_column(col, col, 24 if col == 0 else 80, note_fmt)
            gws.freeze_panes(1, 0)
    return output.getvalue()


def try_parse_app_question_template(rows: List[List[Any]]) -> Optional[Tuple[Dict[str, Any], pd.DataFrame]]:
    """앱 전용 문항정보표 양식 업로드를 감지해 문항정보 DataFrame으로 변환한다."""
    required = {"문항번호", "배점", "정답"}
    header_idx = None
    header_map: Dict[str, int] = {}
    for i, row in enumerate(rows[:30]):
        cells = [clean_text(c).replace(" ", "") for c in row]
        names = set(cells)
        has_domain = any(x in names for x in ["평가요소", "평가영역"])
        if required.issubset(names) and has_domain:
            header_idx = i
            for j, c in enumerate(cells):
                if c:
                    header_map[c] = j
            break
    if header_idx is None:
        return None

    def cell(row: List[Any], *names: str) -> Any:
        for name in names:
            key = name.replace(" ", "")
            if key in header_map:
                idx = header_map[key]
                return row[idx] if len(row) > idx else None
        return None

    items: List[Dict[str, Any]] = []
    for row in rows[header_idx + 1:]:
        q_raw = cell(row, "문항번호")
        if not is_question_no(q_raw):
            if not any(clean_text(c) for c in row):
                continue
            continue
        qno = int(to_number(q_raw) or 0)
        domain = clean_text(cell(row, "평가요소", "평가영역"))
        standard = clean_text(cell(row, "성취기준"))
        difficulty = clean_text(cell(row, "난이도"))
        if difficulty not in ["쉬움", "보통", "어려움"]:
            difficulty = difficulty or "보통"
        items.append(
            {
                "문항번호": qno,
                "평가영역": domain,
                "성취기준": standard,
                "난이도": difficulty,
                "배점": to_number(cell(row, "배점")),
                "정답": normalize_choice_answer_value(cell(row, "정답")),
            }
        )
    if not items:
        raise ValueError("앱 전용 문항정보표 양식에서 입력된 문항을 찾지 못했습니다.")
    qdf = pd.DataFrame(items).drop_duplicates(subset=["문항번호"], keep="first")
    qdf = qdf.sort_values("문항번호").reset_index(drop=True)
    exam_info: Dict[str, Any] = {
        "문항정보표양식": "앱 전용 간편 양식",
        "선택형문항수": int(qdf["문항번호"].max()),
        "선택형만점": float(pd.to_numeric(qdf["배점"], errors="coerce").fillna(0).sum()),
        "서답형문항수": 0,
        "서답형만점": 0.0,
    }
    exam_info["과목만점"] = exam_info["선택형만점"] + exam_info["서답형만점"]
    return exam_info, qdf


# -----------------------------------------------------------------------------
# 나이스 원본 파서
# -----------------------------------------------------------------------------

def read_workbook_rows(uploaded_file: Any) -> List[List[Any]]:
    uploaded_file.seek(0)
    wb = load_workbook(uploaded_file, data_only=True)
    ws = wb[wb.sheetnames[0]]
    rows: List[List[Any]] = []
    for row in ws.iter_rows(values_only=True):
        rows.append(list(row))
    return rows


def parse_question_info(uploaded_file: Any) -> Tuple[Dict[str, Any], pd.DataFrame]:
    rows = read_workbook_rows(uploaded_file)
    app_template_result = try_parse_app_question_template(rows)
    if app_template_result is not None:
        return app_template_result
    subject = extract_subject_from_question_sheet(rows)
    all_text = " ".join(clean_text(c) for row in rows[:10] for c in row if clean_text(c))
    exam_info = parse_exam_info_from_text(all_text)
    if subject:
        exam_info["교과목"] = subject

    # 총점/선택형/서답형 점수 추출
    # 나이스 문항정보표는 보통 C6 셀에
    # "선택형 100.00 점  서답형 0.00 점"과 같은 문구가 들어간다.
    c6_text = clean_text(rows[5][2]) if len(rows) > 5 and len(rows[5]) > 2 else ""
    if c6_text:
        exam_info["문항정보표_C6"] = c6_text
    scan_texts = [c6_text] + [" ".join(clean_text(c) for c in row if clean_text(c)) for row in rows[:10]]
    for text in scan_texts:
        m = re.search(r"총점\s*([\d.]+)", text)
        if m:
            exam_info["과목만점"] = float(m.group(1))
        m = re.search(r"선택형\s*([\d.]+)\s*점", text)
        if m:
            exam_info["선택형만점"] = float(m.group(1))
        m = re.search(r"서답형\s*([\d.]+)\s*점", text)
        if m:
            exam_info["서답형만점"] = float(m.group(1))
    if "선택형만점" in exam_info or "서답형만점" in exam_info:
        exam_info["과목만점"] = float(exam_info.get("선택형만점", 0) or 0) + float(exam_info.get("서답형만점", 0) or 0)
    if float(exam_info.get("서답형만점", 0) or 0) == 0:
        exam_info["서답형문항수"] = 0

    items: List[Dict[str, Any]] = []
    last_item_idx: Optional[int] = None

    for row in rows:
        # A~H열 중심으로 처리: 문항번호, 내용영역, 성취기준, 어려움/보통/쉬움, 배점, 정답
        c0 = row[0] if len(row) > 0 else None
        if is_question_no(c0):
            qno = int(to_number(c0) or 0)
            difficulty = ""
            for label, idx in [("어려움", 3), ("보통", 4), ("쉬움", 5)]:
                if len(row) > idx and "○" in clean_text(row[idx]):
                    difficulty = label
            item = {
                "문항번호": qno,
                "평가영역": clean_text(row[1] if len(row) > 1 else ""),
                "성취기준": clean_text(row[2] if len(row) > 2 else ""),
                "난이도": difficulty,
                "배점": to_number(row[6] if len(row) > 6 else None),
                "정답": normalize_choice_answer_value(row[7] if len(row) > 7 else None),
            }
            items.append(item)
            last_item_idx = len(items) - 1
        else:
            # 페이지가 나뉘면서 성취기준만 다음 줄에 이어지는 경우 직전 문항에 붙임
            cont = clean_text(row[2] if len(row) > 2 else "")
            if last_item_idx is not None and cont and not any(x in cont for x in ["성취기준", "총합계", "비율"]):
                old = clean_text(items[last_item_idx].get("성취기준", ""))
                if cont not in old:
                    items[last_item_idx]["성취기준"] = (old + " " + cont).strip()

    qdf = pd.DataFrame(items).drop_duplicates(subset=["문항번호"], keep="first")
    if not qdf.empty:
        qdf = qdf.sort_values("문항번호").reset_index(drop=True)
        exam_info["선택형문항수"] = int(qdf["문항번호"].max())
        if "선택형만점" not in exam_info:
            exam_info["선택형만점"] = float(qdf["배점"].fillna(0).sum())
        if "서답형문항수" not in exam_info:
            exam_info["서답형문항수"] = 0 if float(exam_info.get("서답형만점", 0) or 0) == 0 else None
    return exam_info, qdf


def parse_answer_sheet(uploaded_file: Any) -> Tuple[Dict[str, Any], pd.DataFrame, pd.DataFrame]:
    rows = read_workbook_rows(uploaded_file)
    all_text = " ".join(clean_text(c) for row in rows[:8] for c in row if clean_text(c))
    exam_info = parse_exam_info_from_text(all_text)
    if "과학" in all_text and "교과목" not in exam_info:
        exam_info["교과목"] = "과학"

    header_idx = None
    for i, row in enumerate(rows):
        texts = [clean_text(c) for c in row]
        q_count = sum(1 for c in row if is_question_no(c))
        has_identity_header = any(t in ["반/번호", "번호", "성명", "이름"] for t in texts)
        next_texts = [clean_text(c) for c in rows[i + 1]] if i + 1 < len(rows) else []
        next_has_answer = any(t == "정답" for t in next_texts)
        if (has_identity_header and q_count >= 3) or (q_count >= 5 and next_has_answer):
            header_idx = i
            break
    if header_idx is None:
        raise ValueError("학생답 정오표에서 학생 목록 시작 행을 찾지 못했습니다. '반/번호', '번호', '이름'이 포함된 표 헤더를 확인해 주세요.")

    header_row = rows[header_idx]
    answer_row = rows[header_idx + 1] if header_idx + 1 < len(rows) else []
    point_row = rows[header_idx + 2] if header_idx + 2 < len(rows) else []

    # 문항번호 열: 실제 문항번호가 들어 있는 열만 수집
    q_cols: List[Tuple[int, int]] = []
    for col_idx, v in enumerate(header_row):
        if is_question_no(v):
            q_cols.append((col_idx, int(to_number(v) or 0)))

    ans_records = []
    for col_idx, qno in q_cols:
        ans_records.append({
            "문항번호": qno,
            "정오표_정답": normalize_choice_answer_value(answer_row[col_idx] if col_idx < len(answer_row) else None),
            "정오표_배점": to_number(point_row[col_idx] if col_idx < len(point_row) else None),
        })
    answer_key_df = pd.DataFrame(ans_records)

    extra_cols: Dict[str, int] = {}
    for col_idx, v in enumerate(header_row):
        text = clean_text(v)
        if text in ["선택형점수", "서답형점수", "기타점수", "영역총점", "총점"]:
            extra_cols[text] = col_idx

    students: List[Dict[str, Any]] = []
    for row in rows[header_idx + 3:]:
        if len(row) == 0:
            continue
        first = row[0] if len(row) > 0 else None
        if not is_student_id(first):
            continue
        rec: Dict[str, Any] = {
            "반/번호": clean_text(row[0]),
            "학번": normalize_identifier_text(row[1] if len(row) > 1 else ""),
            "이름": clean_text(row[2] if len(row) > 2 else ""),
        }
        m = re.match(r"^(\d+)\s*/\s*(\d+)$", rec["반/번호"])
        rec["반"] = int(m.group(1)) if m else None
        rec["번호"] = int(m.group(2)) if m else None
        for col_idx, qno in q_cols:
            rec[f"문항{qno:02d}"] = normalize_mark_value(row[col_idx] if col_idx < len(row) else "")
        for name, col_idx in extra_cols.items():
            rec[name] = to_number(row[col_idx] if col_idx < len(row) else None)
        students.append(rec)

    sdf = pd.DataFrame(students)
    return exam_info, answer_key_df, sdf


# -----------------------------------------------------------------------------
# 계산/분석
# -----------------------------------------------------------------------------

def merge_exam_info(a: Dict[str, Any], b: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(a)
    for k, v in b.items():
        if k not in merged or merged[k] in [None, ""]:
            merged[k] = v
    if "평가구분" not in merged:
        merged["평가구분"] = "정기고사"
    return merged


def make_validation(qdf: pd.DataFrame, answer_key_df: pd.DataFrame) -> pd.DataFrame:
    if qdf.empty or answer_key_df.empty:
        return pd.DataFrame()
    vdf = qdf[["문항번호", "정답", "배점"]].merge(answer_key_df, on="문항번호", how="outer")
    vdf["정답일치"] = vdf.apply(lambda r: answers_match(r.get("정답"), r.get("정오표_정답")), axis=1)
    vdf["배점일치"] = np.isclose(pd.to_numeric(vdf["배점"], errors="coerce"), pd.to_numeric(vdf["정오표_배점"], errors="coerce"), equal_nan=True)
    vdf["검증결과"] = np.where(vdf["정답일치"] & vdf["배점일치"], "정상", "확인 필요")
    return vdf


def classify_level(score: float, cut_a: float, cut_b: float, cut_c: float, cut_d: float) -> str:
    if score >= cut_a:
        return "A"
    if score >= cut_b:
        return "B"
    if score >= cut_c:
        return "C"
    if score >= cut_d:
        return "D"
    return "E"


def build_long_data(qdf: pd.DataFrame, sdf: pd.DataFrame) -> pd.DataFrame:
    qmap = qdf.set_index("문항번호").to_dict(orient="index")
    rows = []
    for _, s in sdf.iterrows():
        for qno, q in qmap.items():
            raw = clean_text(s.get(f"문항{int(qno):02d}", ""))
            correct_answer = normalize_choice_answer_value(q.get("정답"))
            point = float(q.get("배점") or 0)
            selected, mark_type = get_selected_display(raw)
            is_correct = is_correct_mark(raw, correct_answer)
            if is_correct:
                selected = normalize_choice_answer_value(correct_answer)
                status = "정답"
                score = point
            elif mark_type == "무표기":
                status = "무표기"
                score = 0.0
            elif mark_type == "복수답안":
                status = "복수답안오답"
                score = 0.0
            else:
                status = "오답"
                score = 0.0
            rows.append({
                "반/번호": s.get("반/번호"), "반": s.get("반"), "번호": s.get("번호"),
                "학번": s.get("학번"), "이름": s.get("이름"), "문항번호": int(qno),
                "평가영역": q.get("평가영역"), "성취기준": q.get("성취기준"), "난이도": q.get("난이도"),
                "배점": point, "정답": normalize_choice_answer_value(correct_answer), "원본표시": normalize_mark_value(raw), "선택지": normalize_choice_answer_value(selected),
                "정오": status, "정답여부": is_correct, "점수": score,
            })
    return pd.DataFrame(rows)


def add_student_scores(sdf: pd.DataFrame, long_df: pd.DataFrame, total_full_score: float, cuts: Dict[str, float]) -> pd.DataFrame:
    # 문항별 정오표로 다시 계산한 점수는 선택형 문항 검증/문항 분석용으로 보존한다.
    # 서답형이 있는 시험의 총점·성취도·성취수준은 나이스 정오표의 영역총점을 우선 사용한다.
    score_df = long_df.groupby("반/번호", as_index=False)["점수"].sum().rename(columns={"점수": "선택형계산점수"})
    out = sdf.merge(score_df, on="반/번호", how="left")
    out["선택형계산점수"] = pd.to_numeric(out["선택형계산점수"], errors="coerce").fillna(0)

    for col in ["선택형점수", "서답형점수", "기타점수", "영역총점", "총점"]:
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce")

    if "선택형점수" in out.columns:
        out["선택형점수차이"] = out["선택형계산점수"] - out["선택형점수"].fillna(0)

    if "영역총점" in out.columns and out["영역총점"].notna().any():
        out["원본총점"] = out["영역총점"]
        final_score = out["영역총점"]
        score_source = "영역총점"
    elif "총점" in out.columns and out["총점"].notna().any():
        out["원본총점"] = out["총점"]
        final_score = out["총점"]
        score_source = "총점"
    else:
        constructed = pd.Series(0.0, index=out.index)
        for col in ["서답형점수", "기타점수"]:
            if col in out.columns:
                constructed = constructed + out[col].fillna(0)
        final_score = out["선택형계산점수"] + constructed
        out["원본총점"] = final_score
        score_source = "선택형계산점수+서답형점수+기타점수" if constructed.abs().sum() else "선택형계산점수"

    out["최종점수"] = pd.to_numeric(final_score, errors="coerce").fillna(out["선택형계산점수"])
    # 기존 화면/엑셀/AI 코드와 호환되도록 계산점수는 최종점수와 같은 값으로 둔다.
    out["계산점수"] = out["최종점수"]
    out["성취도점수출처"] = score_source
    out["환산점수"] = out["최종점수"] / total_full_score * 100 if total_full_score else out["최종점수"]
    # 성취수준은 환산점수가 아니라 시험지 자체의 원점수 기준으로 산출한다.
    out["성취수준기준점수"] = out["최종점수"]
    out["성취수준"] = out["성취수준기준점수"].apply(lambda x: classify_level(float(x), cuts["A"], cuts["B"], cuts["C"], cuts["D"]))
    return out


def calc_cronbach_alpha(long_df: pd.DataFrame) -> Optional[float]:
    if long_df.empty:
        return None
    pivot = long_df.pivot_table(index="반/번호", columns="문항번호", values="점수", aggfunc="sum", fill_value=0)
    k = pivot.shape[1]
    n = pivot.shape[0]
    if k <= 1 or n <= 1:
        return None
    item_var_sum = pivot.var(axis=0, ddof=1).sum()
    total_var = pivot.sum(axis=1).var(ddof=1)
    if total_var == 0 or pd.isna(total_var):
        return None
    return float(k / (k - 1) * (1 - item_var_sum / total_var))


def make_class_score_summary_chart_df(class_achievement: pd.DataFrame) -> pd.DataFrame:
    df = class_achievement.copy()
    if df.empty:
        return pd.DataFrame()
    df = df.sort_values("반")
    df["학급"] = df["반"].apply(lambda x: f"{int(x)}반" if pd.notna(x) and float(x).is_integer() else f"{x}반")
    out = df[["학급", "최고점", "평균", "최저점"]].copy()
    for col in ["최고점", "평균", "최저점"]:
        out[col] = pd.to_numeric(out[col], errors="coerce")
    return out.set_index("학급")


def make_class_score_distribution_chart_data(individual_df: pd.DataFrame, class_achievement: pd.DataFrame, full_score: float = 100.0, bin_size: int = 5) -> Tuple[pd.DataFrame, pd.DataFrame, str, float]:
    """학급별 최고·최저·평균과 점수 구간별 도수를 함께 그리기 위한 Vega-Lite용 데이터 생성."""
    if individual_df.empty or class_achievement.empty or "계산점수" not in individual_df.columns:
        return pd.DataFrame(), pd.DataFrame(), "", 100.0

    class_values = sorted(individual_df["반"].dropna().unique().tolist(), key=lambda x: float(x) if str(x).replace('.', '', 1).isdigit() else str(x))
    class_pos_map = {str(c): i + 1 for i, c in enumerate(class_values)}
    class_label_map = {i + 1: (f"{int(c)}반" if pd.notna(c) and float(c).is_integer() else f"{c}반") for i, c in enumerate(class_values)}

    dist_rows = []
    for c in class_values:
        scores = pd.to_numeric(individual_df[individual_df["반"].astype(str) == str(c)]["계산점수"], errors="coerce").dropna()
        if scores.empty:
            continue
        y_max = float(full_score or max(scores.max(), 100))
        y_max = max(y_max, float(scores.max()), bin_size)
        bins = np.arange(0, y_max + bin_size, bin_size)
        counts, edges = np.histogram(scores.clip(lower=0, upper=y_max), bins=bins)
        max_count = max(int(counts.max()), 1)
        pos = class_pos_map[str(c)]
        label = class_label_map[pos]
        for count, low, high in zip(counts, edges[:-1], edges[1:]):
            if count <= 0:
                continue
            half_width = 0.06 + (float(count) / max_count) * 0.34
            dist_rows.append({
                "학급위치": pos,
                "학급": label,
                "점수구간하한": float(low),
                "점수구간상한": float(high),
                "도수(명)": int(count),
                "점수구간": f"{low:.0f}점-{high:.0f}점",
                "왼쪽폭": pos - half_width,
                "오른쪽폭": pos + half_width,
            })

    summary = class_achievement.copy()
    summary["학급위치"] = summary["반"].astype(str).map(class_pos_map)
    summary = summary.dropna(subset=["학급위치"]).copy()
    summary["학급위치"] = summary["학급위치"].astype(float)
    summary["학급"] = summary["학급위치"].astype(int).map(class_label_map)
    for col in ["최고점", "평균", "최저점"]:
        summary[col] = pd.to_numeric(summary[col], errors="coerce")
    summary["최고라벨"] = summary["최고점"].apply(lambda x: f"최고 {x:.1f}" if pd.notna(x) else "최고")
    summary["최저라벨"] = summary["최저점"].apply(lambda x: f"최저 {x:.1f}" if pd.notna(x) else "최저")
    summary["평균라벨"] = summary["평균"].apply(lambda x: f"평균 {x:.1f}" if pd.notna(x) else "평균")
    summary = summary[["학급위치", "학급", "최고점", "평균", "최저점", "최고라벨", "최저라벨", "평균라벨"]]

    label_array = "[" + ",".join([repr(class_label_map[i]) for i in sorted(class_label_map)]) + "]"
    label_expr = f"{label_array}[datum.value-1]"
    y_max = float(full_score or 100.0)
    if not summary.empty:
        y_max = max(y_max, float(pd.to_numeric(summary["최고점"], errors="coerce").max() or y_max))
    return pd.DataFrame(dist_rows), summary, label_expr, y_max




def make_total_level_distribution_chart_df(individual_df: pd.DataFrame) -> pd.DataFrame:
    df = individual_df.copy()
    if df.empty or "성취수준" not in df.columns:
        return pd.DataFrame()
    levels = list("ABCDE")
    counts = df["성취수준"].value_counts().reindex(levels, fill_value=0)
    total = int(counts.sum())
    return pd.DataFrame({
        "성취수준": levels,
        "학생수(명)": counts.astype(int).values,
        "비율": [(count / total) if total else 0 for count in counts.values],
    })

def make_class_level_distribution_chart_df(individual_df: pd.DataFrame, selected_class: Any) -> pd.DataFrame:
    df = individual_df[individual_df["반"].astype(str) == str(selected_class)].copy()
    if df.empty:
        return pd.DataFrame()
    levels = list("ABCDE")
    counts = df["성취수준"].value_counts().reindex(levels, fill_value=0)
    total = int(counts.sum())
    # 내부 비율값은 앱 전체 규칙에 맞춰 0~1로 유지한다.
    # 화면 표시 단계에서만 52.3% 형식으로 변환한다.
    out = pd.DataFrame({
        "성취수준": levels,
        "학생수(명)": counts.astype(int).values,
        "비율": [(count / total) if total else 0 for count in counts.values],
    })
    return out


def item_discrimination(long_df: pd.DataFrame, student_scores: pd.DataFrame) -> pd.DataFrame:
    scores = student_scores[["반/번호", "계산점수"]].dropna().sort_values("계산점수", ascending=False)
    n = len(scores)
    if n < 4:
        return pd.DataFrame({"문항번호": sorted(long_df["문항번호"].unique()), "변별도": np.nan})
    group_n = max(1, int(round(n * 0.27)))
    high_ids = set(scores.head(group_n)["반/번호"])
    low_ids = set(scores.tail(group_n)["반/번호"])
    rows = []
    for qno, g in long_df.groupby("문항번호"):
        high_rate = g[g["반/번호"].isin(high_ids)]["정답여부"].mean()
        low_rate = g[g["반/번호"].isin(low_ids)]["정답여부"].mean()
        rows.append({"문항번호": qno, "상위집단정답률": high_rate, "하위집단정답률": low_rate, "변별도": high_rate - low_rate})
    return pd.DataFrame(rows)


DIFFICULTY_SCORE = {"어려움": 1, "보통": 2, "쉬움": 3}


def normalize_difficulty_value(value: Any) -> str:
    text = clean_text(value)
    if text in DIFFICULTY_SCORE:
        return text
    if "어려" in text:
        return "어려움"
    if "보통" in text or "중" == text:
        return "보통"
    if "쉬" in text:
        return "쉬움"
    return ""


def actual_difficulty_from_rate(rate: Any, hard_cut_percent: float, easy_cut_percent: float) -> str:
    r = pd.to_numeric(pd.Series([rate]), errors="coerce").iloc[0]
    if pd.isna(r):
        return ""
    percent = float(r) * 100
    if percent < hard_cut_percent:
        return "어려움"
    if percent < easy_cut_percent:
        return "보통"
    return "쉬움"


def expected_rate_bounds(difficulty: str, hard_cut_percent: float, easy_cut_percent: float) -> Tuple[Optional[float], Optional[float]]:
    diff = normalize_difficulty_value(difficulty)
    if diff == "어려움":
        return 0.0, float(hard_cut_percent)
    if diff == "보통":
        return float(hard_cut_percent), float(easy_cut_percent)
    if diff == "쉬움":
        return float(easy_cut_percent), 100.0
    return None, None


def difficulty_rate_gap(rate: Any, difficulty: str, hard_cut_percent: float, easy_cut_percent: float) -> float:
    r = pd.to_numeric(pd.Series([rate]), errors="coerce").iloc[0]
    if pd.isna(r):
        return np.nan
    actual_pct = float(r) * 100
    low, high = expected_rate_bounds(difficulty, hard_cut_percent, easy_cut_percent)
    if low is None or high is None:
        return np.nan
    if actual_pct < low:
        return actual_pct - low
    if actual_pct >= high and high < 100:
        return actual_pct - high
    return 0.0


def difficulty_gap_label(gap: Any) -> str:
    g = pd.to_numeric(pd.Series([gap]), errors="coerce").iloc[0]
    if pd.isna(g):
        return "비교 불가"
    if abs(float(g)) < 0.05:
        return "일치"
    if float(g) < 0:
        return "예상보다 어려웠음"
    return "예상보다 쉬웠음"


def make_difficulty_gap_analysis(
    item_df: pd.DataFrame,
    hard_cut_percent: float = 33.0,
    easy_cut_percent: float = 66.0,
    expected_question_df: Optional[pd.DataFrame] = None,
) -> pd.DataFrame:
    out = item_df.copy()

    # 예상 난이도는 실제 정답률이나 기준 슬라이더가 아니라
    # 업로드된 문항정보표 원본의 난이도에 고정한다.
    if expected_question_df is not None and not expected_question_df.empty and {"문항번호", "난이도"}.issubset(expected_question_df.columns):
        expected_map = expected_question_df[["문항번호", "난이도"]].copy()
        expected_map["문항번호"] = pd.to_numeric(expected_map["문항번호"], errors="coerce")
        expected_map["예상난이도"] = expected_map["난이도"].apply(normalize_difficulty_value)
        expected_map = expected_map.dropna(subset=["문항번호"]).drop_duplicates("문항번호", keep="first")[["문항번호", "예상난이도"]]
        out["문항번호"] = pd.to_numeric(out["문항번호"], errors="coerce")
        out = out.drop(columns=["예상난이도"], errors="ignore").merge(expected_map, on="문항번호", how="left")
        out["예상난이도"] = out["예상난이도"].fillna("")
    else:
        out["예상난이도"] = out.get("난이도", "").apply(normalize_difficulty_value) if "난이도" in out.columns else ""
    out["정답률_pct"] = pd.to_numeric(out["정답률"], errors="coerce") * 100
    bounds = out["예상난이도"].apply(lambda x: expected_rate_bounds(x, hard_cut_percent, easy_cut_percent))
    out["기준하한"] = bounds.apply(lambda x: x[0])
    out["기준상한"] = bounds.apply(lambda x: x[1])
    out["기대정답률구간"] = out.apply(
        lambda r: "" if pd.isna(r["기준하한"]) or pd.isna(r["기준상한"]) else f"{r['기준하한']:.0f}%-{r['기준상한']:.0f}%",
        axis=1,
    )
    out["차이(%p)"] = out.apply(
        lambda r: difficulty_rate_gap(r["정답률"], r["예상난이도"], hard_cut_percent, easy_cut_percent),
        axis=1,
    )
    out["차이해석"] = out["차이(%p)"].apply(difficulty_gap_label)
    out["괴리여부"] = np.where(out["차이해석"] == "일치", "일치", np.where(out["차이해석"] == "비교 불가", "비교 불가", "불일치"))
    keep_cols = [
        "문항번호", "평가영역", "예상난이도", "기대정답률구간", "정답률", "정답률_pct",
        "차이(%p)", "차이해석", "괴리여부", "배점", "정답", "변별도"
    ]
    return out[[c for c in keep_cols if c in out.columns]].sort_values("문항번호", ascending=True)


def analyze_all(parsed: ParsedData, total_full_score: float, cuts: Dict[str, float]) -> Dict[str, pd.DataFrame | float | None]:
    qdf, sdf0, long_df0 = parsed.question_df.copy(), parsed.students_df.copy(), parsed.long_df.copy()
    students = add_student_scores(sdf0, long_df0, total_full_score, cuts)
    long_df = long_df0.merge(students[["반/번호", "계산점수", "최종점수", "선택형계산점수", "환산점수", "성취수준"]], on="반/번호", how="left")

    # 성취도 분석
    level_counts = students["성취수준"].value_counts().reindex(list("ABCDE"), fill_value=0)
    achievement = pd.DataFrame({
        "구분": ["전체"],
        "응시자수": [len(students)],
        "평균": [students["계산점수"].mean()],
        "표준편차": [students["계산점수"].std(ddof=1)],
        "최고점": [students["계산점수"].max()],
        "최저점": [students["계산점수"].min()],
        "A인원": [level_counts["A"]], "B인원": [level_counts["B"]], "C인원": [level_counts["C"]],
        "D인원": [level_counts["D"]], "E인원": [level_counts["E"]],
    })
    class_achievement = students.groupby("반", dropna=False).agg(
        응시자수=("반/번호", "count"), 평균=("계산점수", "mean"), 표준편차=("계산점수", "std"),
        최고점=("계산점수", "max"), 최저점=("계산점수", "min")
    ).reset_index()
    for level in list("ABCDE"):
        cnt = students[students["성취수준"] == level].groupby("반")["반/번호"].count()
        class_achievement[f"{level}인원"] = class_achievement["반"].map(cnt).fillna(0).astype(int)

    # 문항별 분석
    item = long_df.groupby("문항번호").agg(
        배점=("배점", "first"), 정답=("정답", "first"), 평가영역=("평가영역", "first"), 난이도=("난이도", "first"),
        응시자수=("반/번호", "count"), 정답자수=("정답여부", "sum"), 평균점수=("점수", "mean")
    ).reset_index()
    item["정답률"] = item["정답자수"] / item["응시자수"]
    item = item.merge(item_discrimination(long_df, students), on="문항번호", how="left")
    # 선택지 반응률
    for opt in [1, 2, 3, 4, 5, "무표기"]:
        col = f"선택지{opt}비율" if opt != "무표기" else "무표기비율"
        rates = long_df.assign(_sel=long_df["선택지"].astype(str)).groupby("문항번호").apply(lambda g, o=str(opt): (g["_sel"] == o).mean()).reset_index(name=col)
        item = item.merge(rates, on="문항번호", how="left")

    # 예상 난이도-실제 난이도 기본 괴리 분석
    difficulty_gap = make_difficulty_gap_analysis(item, hard_cut_percent=33.0, easy_cut_percent=66.0, expected_question_df=parsed.original_question_df)

    # 학급별 분석
    class_item = long_df.groupby(["반", "문항번호"]).agg(응시자수=("반/번호", "count"), 정답률=("정답여부", "mean"), 평균점수=("점수", "mean")).reset_index()
    class_item_pivot = class_item.pivot(index="문항번호", columns="반", values="정답률").reset_index()
    class_item_pivot.columns = ["문항번호"] + [f"{int(c)}반_정답률" if not pd.isna(c) else "미상반_정답률" for c in class_item_pivot.columns[1:]]

    # 평가영역별 분석
    # 영역 만점은 문항정보표 기준으로 문항 배점을 합산하고,
    # 영역 평균은 학생별 영역 총점의 평균으로 계산한다.
    domain_scores = long_df.groupby(["반/번호", "평가영역"]).agg(
        영역점수=("점수", "sum"),
        영역배점=("배점", "sum"),
        영역정답률=("정답여부", "mean"),
    ).reset_index()
    student_meta = students[["반/번호", "반", "번호", "이름"]].drop_duplicates("반/번호")
    domain_scores = student_meta.merge(domain_scores, on="반/번호", how="right")
    domain_scores["영역환산점수"] = np.where(domain_scores["영역배점"] > 0, domain_scores["영역점수"] / domain_scores["영역배점"] * 100, np.nan)
    domain_max = qdf.groupby("평가영역", dropna=False).agg(
        문항수=("문항번호", "nunique"),
        배점합계=("배점", "sum"),
    ).reset_index()
    domain_avg = domain_scores.groupby("평가영역", dropna=False).agg(
        평균점수=("영역점수", "mean"),
        평균정답률=("영역정답률", "mean"),
    ).reset_index()
    domain_rate = long_df.groupby("평가영역", dropna=False).agg(정답률=("정답여부", "mean")).reset_index()
    domain = domain_max.merge(domain_avg, on="평가영역", how="left").merge(domain_rate, on="평가영역", how="left")
    domain["환산평균"] = np.where(domain["배점합계"] > 0, domain["평균점수"] / domain["배점합계"] * 100, np.nan)

    # 성취기준별 분석
    # 같은 성취기준에 속한 문항을 묶어 문항 수, 배점, 평균점수, 정답률을 계산한다.
    long_df["성취기준"] = long_df["성취기준"].fillna("").astype(str).str.strip().replace("", "미입력")
    qdf["성취기준"] = qdf["성취기준"].fillna("").astype(str).str.strip().replace("", "미입력")
    standard_scores = long_df.groupby(["반/번호", "성취기준"]).agg(
        성취기준점수=("점수", "sum"),
        성취기준배점=("배점", "sum"),
        성취기준정답률=("정답여부", "mean"),
    ).reset_index()
    standard_scores = student_meta.merge(standard_scores, on="반/번호", how="right")
    standard_scores["성취기준환산점수"] = np.where(standard_scores["성취기준배점"] > 0, standard_scores["성취기준점수"] / standard_scores["성취기준배점"] * 100, np.nan)
    standard_max = qdf.groupby("성취기준", dropna=False).agg(
        문항수=("문항번호", "nunique"),
        배점합계=("배점", "sum"),
        문항번호=("문항번호", lambda s: ", ".join(map(str, sorted(pd.to_numeric(s, errors="coerce").dropna().astype(int).unique())))),
        평가영역=("평가영역", lambda s: ", ".join(map(str, sorted(set(str(x).strip() for x in s if str(x).strip()))))),
    ).reset_index()
    standard_avg = standard_scores.groupby("성취기준", dropna=False).agg(
        평균점수=("성취기준점수", "mean"),
        평균정답률=("성취기준정답률", "mean"),
    ).reset_index()
    standard_rate = long_df.groupby("성취기준", dropna=False).agg(정답률=("정답여부", "mean")).reset_index()
    standard = standard_max.merge(standard_avg, on="성취기준", how="left").merge(standard_rate, on="성취기준", how="left")
    standard["환산평균"] = np.where(standard["배점합계"] > 0, standard["평균점수"] / standard["배점합계"] * 100, np.nan)

    # 성취수준별 문항 분석
    level_item = long_df.groupby(["성취수준", "문항번호"]).agg(응시자수=("반/번호", "count"), 정답률=("정답여부", "mean")).reset_index()
    level_item_pivot = level_item.pivot(index="문항번호", columns="성취수준", values="정답률").reset_index()
    for level in list("ABCDE"):
        if level not in level_item_pivot.columns:
            level_item_pivot[level] = np.nan
    level_item_pivot = level_item_pivot[["문항번호", "A", "B", "C", "D", "E"]]
    level_item_pivot = item[["문항번호", "정답률", "평가영역"]].merge(level_item_pivot, on="문항번호", how="left")
    level_item_pivot["수준간격차"] = level_item_pivot[["A", "B", "C", "D", "E"]].max(axis=1) - level_item_pivot[["A", "B", "C", "D", "E"]].min(axis=1)
    # 화면에서 비어 보이지 않도록 문항번호 순서를 명시적으로 고정한다.
    level_item_pivot = level_item_pivot.sort_values("문항번호", ascending=True).reset_index(drop=True)

    # 학생 개별 분석용 요약
    individual_cols = [
        "반/번호", "반", "번호", "이름", "선택형계산점수", "선택형점수", "서답형점수",
        "기타점수", "영역총점", "최종점수", "계산점수", "환산점수", "성취수준기준점수", "성취수준"
    ]
    individual = students[[c for c in individual_cols if c in students.columns]].copy()
    weak_items = long_df[~long_df["정답여부"]].groupby("반/번호")["문항번호"].apply(lambda s: ", ".join(map(str, sorted(s.tolist())))).reset_index(name="오답문항")
    individual = individual.merge(weak_items, on="반/번호", how="left")
    individual["오답문항"] = individual["오답문항"].fillna("")

    alpha = calc_cronbach_alpha(long_df)
    return {
        "students": students,
        "long": long_df,
        "achievement": achievement,
        "class_achievement": class_achievement,
        "item": item,
        "difficulty_gap": difficulty_gap,
        "class_item": class_item,
        "class_item_pivot": class_item_pivot,
        "domain": domain,
        "standard": standard,
        "standard_scores": standard_scores,
        "level_item": level_item_pivot,
        "domain_scores": domain_scores,
        "individual": individual,
        "alpha": alpha,
    }


# -----------------------------------------------------------------------------
# 엑셀 출력
# -----------------------------------------------------------------------------

def autosize_worksheet(writer: pd.ExcelWriter, sheet_name: str, df: pd.DataFrame) -> None:
    ws = writer.sheets[sheet_name]
    for idx, col in enumerate(df.columns):
        max_len = max([len(str(col))] + [len(str(v)) for v in df[col].head(300).fillna("")])
        ws.set_column(idx, idx, min(max(max_len + 2, 10), 45))
    header_fmt = writer.book.add_format({"bold": True, "bg_color": "#E8F0FE", "border": 1, "align": "center"})
    for col_num, value in enumerate(df.columns.values):
        ws.write(0, col_num, value, header_fmt)
    ws.freeze_panes(1, 0)


def df_to_excel_bytes(sheets: Dict[str, pd.DataFrame]) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        for name, df in sheets.items():
            safe = safe_sheet_name(name)
            out_df = format_output_df(df, add_units=True)
            out_df.to_excel(writer, index=False, sheet_name=safe)
            autosize_worksheet(writer, safe, out_df)
            ws = writer.sheets[safe]
            number_fmt = writer.book.add_format({"num_format": "0.00"})
            integer_fmt = writer.book.add_format({"num_format": "0"})
            for i, col in enumerate(out_df.columns):
                col_name = str(col)
                if any(unit in col_name for unit in ["(점)"]):
                    ws.set_column(i, i, 12, number_fmt)
                elif any(unit in col_name for unit in ["(명)", "(개)"]):
                    ws.set_column(i, i, 10, integer_fmt)
    return output.getvalue()


def make_confirm_excel(parsed: ParsedData, analysis: Dict[str, Any]) -> bytes:
    sheets = {
        "평가정보": pd.DataFrame([parsed.exam_info]),
        "문항정보": parsed.question_df,
        "학생정오표": parsed.students_df,
        "학생점수": analysis["students"],
        "성취기준별분석": analysis.get("standard", pd.DataFrame()),
        "검증결과": parsed.validation_df,
        "문항별긴자료": analysis["long"],
    }
    return df_to_excel_bytes(sheets)


def make_analysis_zip(parsed: ParsedData, analysis: Dict[str, Any]) -> bytes:
    files = {
        "성취도분석.xlsx": {
            "전체": analysis["achievement"],
            "학급별": analysis["class_achievement"],
            "학생별": analysis["individual"],
        },
        "선다형분석_문항별.xlsx": {
            "문항별분석": analysis["item"],
            "난이도괴리분석": analysis.get("difficulty_gap", pd.DataFrame()),
            "응답긴자료": analysis["long"],
        },
        "선다형분석_학급별.xlsx": {
            "학급별문항": analysis["class_item"],
            "문항별학급비교": analysis["class_item_pivot"],
        },
        "평가영역별분석.xlsx": {
            "영역별분석": analysis["domain"],
            "학생영역별": analysis["domain_scores"],
        },
        "성취기준별분석.xlsx": {
            "성취기준별분석": analysis.get("standard", pd.DataFrame()),
            "학생성취기준별": analysis.get("standard_scores", pd.DataFrame()),
        },
        "선다형분석_성취수준별.xlsx": {
            "성취수준별문항": analysis["level_item"],
        },
    }
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for filename, sheets in files.items():
            zf.writestr(filename, df_to_excel_bytes(sheets))
    return zip_buf.getvalue()


# -----------------------------------------------------------------------------
# AI 분석
# -----------------------------------------------------------------------------

def build_basic_statistics_ai_prompt(parsed: ParsedData, analysis: Dict[str, Any]) -> str:
    """기본 분석: 원안지 없이 통계 결과만으로 확장 해석을 생성한다."""
    item = analysis["item"].copy().sort_values("정답률")
    domain = analysis["domain"].copy().sort_values("정답률")
    standard = analysis.get("standard", pd.DataFrame()).copy()
    if not standard.empty and "정답률" in standard.columns:
        standard = standard.sort_values("정답률")
    level_item = analysis["level_item"].copy().sort_values("정답률")
    difficulty_gap = analysis.get("difficulty_gap", pd.DataFrame()).copy()
    if not difficulty_gap.empty:
        difficulty_gap = difficulty_gap[difficulty_gap["괴리여부"].isin(["불일치", "차이 있음"])].copy()
        difficulty_gap = difficulty_gap.sort_values(["정답률", "문항번호"], ascending=[True, True])
    alpha = analysis.get("alpha")
    exam = parsed.exam_info
    return f"""
너는 중학교 과학 평가 결과를 해석하는 교육평가 전문가이다.

아래 제공된 통계 자료를 바탕으로 평가 결과를 최대한 풍성하게 해석하라.
원안지 문항 내용은 제공되지 않았으므로 문항의 구체적 표현, 자료, 선지 구성은 직접 볼 수 없다.
그러나 정답률, 변별도, 선택지 반응률, 평가영역, 성취기준, 성취수준별 결과를 종합하여 가능한 교육적 해석을 적극적으로 제시하라.

[평가 정보]
{exam}

[성취도 요약]
{ai_percent_df(analysis['achievement']).to_string(index=False)}
검사신뢰도 알파: {None if alpha is None else round(alpha, 3)}

[학급별 성취도]
{ai_percent_df(analysis['class_achievement']).to_string(index=False)}

[정답률 낮은 문항]
{ai_percent_df(item[['문항번호','평가영역','난이도','배점','정답','정답률','변별도']].head(10)).to_string(index=False)}

[선택지 반응 포함 문항별 분석]
{ai_percent_df(item.head(15)).to_string(index=False)}

[평가영역별 분석]
{ai_percent_df(domain).to_string(index=False)}

[성취기준별 분석]
{ai_percent_df(standard).to_string(index=False) if not standard.empty else '성취기준별 분석 데이터 없음'}

[성취수준별 문항 분석 중 정답률 낮은 문항]
{ai_percent_df(level_item[['문항번호','평가영역','정답률','A','B','C','D','E','수준간격차']].head(12)).to_string(index=False)}

[예상 난이도와 실제 정답률 판정]
{difficulty_gap[['문항번호','평가영역','예상난이도','기대정답률구간','정답률_pct','차이해석']].head(12).to_string(index=False) if not difficulty_gap.empty else '불일치 문항 없음'}

다음 구조로 작성하라.

1. 전체 평가 결과 요약
- 전체 평균, 표준편차, 최고점, 최저점, 성취수준 분포를 바탕으로 시험의 전반적 난이도와 점수 분포 특징을 해석하라.
- 수치 나열에 그치지 말고, 해당 분포가 수업 및 평가 측면에서 어떤 의미를 갖는지 설명하라.

2. 학급별 결과 해석
- 학급별 평균, 최고점, 최저점, 성취수준 분포를 비교하라.
- 전체 경향과 다르게 나타나는 학급이 있으면 가능한 원인을 폭넓게 해석하라.

3. 문항별 결과 해석
- 정답률이 낮은 문항, 높은 문항, 변별도가 낮은 문항, 변별도가 높은 문항을 중심으로 해석하라.
- 예상 난이도와 실제 정답률이 어긋난 문항을 별도로 다루라.
- 특정 선택지에 응답이 몰린 문항은 가능한 개념 혼동이나 판단 과정의 가능성을 추론하라.

4. 선택지 반응 및 오답 경향 분석
- 특정 오답 선택지에 응답이 몰린 문항을 찾아라.
- 해당 선택지가 학생들에게 매력적으로 작용했을 가능성을 추론하라.
- 가능한 개념 혼동, 자료 해석 오류, 절차적 실수 가능성을 구분하여 제시하라.

5. 평가영역 및 성취기준별 해석
- 학생들이 강점을 보인 영역과 어려움을 보인 영역을 구분하라.
- 낮은 정답률을 보인 성취기준은 보충 지도가 필요한 개념이나 사고 과정을 추론하라.

6. 성취수준별 집단 해석
- 상위권과 하위권을 가르는 문항, 상위권도 어려워한 문항, 하위권에게 특히 부담이 큰 문항을 구분하라.
- 수준별 피드백 방향을 제안하라.

7. 수업 개선 및 피드백 방향
- 다음 수업에서 우선적으로 보완해야 할 내용, 재지도 방식, 학생 피드백 방향을 제안하라.

8. 원안지 기반 고급 분석 추천 문항
- 위 통계 분석 결과를 종합하여, 이후 고급 분석에서 원안지 PDF를 확인하며 문항 내용을 구체적으로 살펴볼 필요가 있는 문항을 우선순위 순서대로 추천하라.
- 추천 기준은 정답률이 낮은 문항, 예상 난이도와 실제 결과가 다른 문항, 변별도가 낮은 문항, 특정 오답 선택지에 응답이 몰린 문항, 성취수준별 정답률 차이가 큰 문항, 평가영역·성취기준상 보충 필요성이 큰 문항 등을 종합하라.
- 추천 문항은 많으면 많은 대로 제시하되, 실제로 고급 분석이 필요하다고 볼 만한 문항이 거의 없으면 억지로 추천하지 않아도 된다.
- 각 추천 문항마다 추천 순위, 문항번호, 추천 이유, 고급 분석에서 확인해야 할 점을 함께 제시하라.

작성 원칙:
- 수치를 근거로 하되, 가능한 해석을 풍성하게 제시하라.
- 원안지 문항 내용이 제공되지 않은 상태이므로 문항 표현, 자료 구성, 선지 유인, 구체적 오개념에 대한 해석은 가능성으로 제시하라.
- 문항 내용을 확인할 수 없다는 말만 반복하지 말고, 현재 주어진 통계 자료로 가능한 해석을 최대한 끌어내라.
- 문항번호, 성취기준, 평가영역, 정답률, 선택지 반응률을 활용해 구체적으로 작성하라.
""".strip()


# 이전 버전 호환용 별칭
build_overall_ai_prompt = build_basic_statistics_ai_prompt


def append_focus_request_to_prompt(prompt: str, focus_request: str) -> str:
    """사용자가 입력한 중점 분석 요청을 기존 AI 프롬프트 마지막에 추가한다."""
    focus_request = str(focus_request or "").strip()
    if not focus_request:
        return prompt
    focus_block = f"""

[의뢰자 중점 분석 요청]
아래 내용은 의뢰자가 평가 전문가에게 특별히 중점적으로 분석해 달라고 요청한 부분이다.
기존 분석 지시사항은 그대로 수행한 뒤, 마지막 단계에서 아래 요청을 반드시 반영하라.
아래 요청이 기존 분석 항목과 겹치는 경우에는 같은 내용을 단순 반복하지 말고, 해당 항목을 더 구체적이고 깊이 있게 분석하라.
분석 결과 안에서 의뢰자의 중점 분석 요청이 어떤 부분에 어떻게 반영되었는지 명시하라.

의뢰자 요청:
{focus_request}
""".rstrip()
    return f"{prompt.rstrip()}\n{focus_block}"


PROMPT_MODE_DEFAULT = "기본 프롬프트 그대로 사용"
PROMPT_MODE_WITH_FOCUS = "기본 프롬프트 + 추가 의뢰 사용"
PROMPT_MODE_DIRECT = "직접 작성한 프롬프트만 사용"
PROMPT_MODE_OPTIONS = [PROMPT_MODE_DEFAULT, PROMPT_MODE_WITH_FOCUS, PROMPT_MODE_DIRECT]
PROMPT_MODE_HELP = (
    "웹앱에는 분석 유형별 기본 프롬프트가 내장되어 있습니다. "
    "기본 프롬프트는 현재 통계 자료나 원안지 PDF를 어떤 관점과 구조로 분석할지 정해 둔 분석 틀이며, "
    "아래의 '웹앱 기본 프롬프트 보기'에서 실제 분석 구조를 확인할 수 있습니다. "
    "'기본 프롬프트 그대로 사용'은 내장된 분석 틀만 사용합니다. "
    "'기본 프롬프트 + 추가 의뢰 사용'은 기본 분석 틀은 유지하되, 중점 분석 요청 프리셋이나 직접 입력한 요구사항을 덧붙입니다. "
    "'직접 작성한 프롬프트만 사용'은 웹앱 기본 분석 틀을 사용하지 않고, 사용자가 작성한 지시문을 중심으로 분석합니다. "
    "단, 직접 작성 모드에서도 앱이 계산한 통계 데이터와 고급 분석의 원안지 PDF는 AI에게 함께 전달됩니다."
)


def extract_prompt_data_context(base_prompt: str) -> str:
    """직접 작성 프롬프트 모드에서도 앱 분석 데이터는 함께 전달하기 위해 기본 프롬프트의 데이터 블록만 추출한다."""
    prompt_text = str(base_prompt or "")
    start_candidates = [
        prompt_text.find("[고급 분석 범위]"),
        prompt_text.find("[학생]"),
        prompt_text.find("[평가 정보]"),
    ]
    start_candidates = [idx for idx in start_candidates if idx >= 0]
    start_idx = min(start_candidates) if start_candidates else 0

    end_candidates = []
    for marker in ["\n\n다음 구조", "\n\n다음 형식", "\n\n작성 원칙:"]:
        idx = prompt_text.find(marker, start_idx)
        if idx >= 0:
            end_candidates.append(idx)
    end_idx = min(end_candidates) if end_candidates else len(prompt_text)
    return prompt_text[start_idx:end_idx].strip()


def extract_prompt_instruction_preview(base_prompt: str) -> str:
    """웹앱 기본 프롬프트 보기에는 실제 통계 데이터가 아니라 분석 지시 구조만 표시한다."""
    prompt_text = str(base_prompt or "").strip()
    if not prompt_text:
        return ""

    for marker in ["다음 구조", "다음 형식"]:
        idx = prompt_text.find(marker)
        if idx >= 0:
            return prompt_text[idx:].strip()

    # 예외적으로 구조 안내 문구를 찾지 못하면 앱 데이터 블록은 최대한 제외하고 앞부분만 보여준다.
    data_markers = ["[평가 정보]", "[학생]", "[고급 분석 범위]", "[원안지 파일]"]
    data_indices = [prompt_text.find(marker) for marker in data_markers if prompt_text.find(marker) >= 0]
    end_idx = min(data_indices) if data_indices else len(prompt_text)
    return prompt_text[:end_idx].strip()


def build_direct_custom_ai_prompt(base_prompt: str, custom_prompt: str, analysis_label: str) -> str:
    """기본 분석 목차는 빼고 사용자가 직접 작성한 지시문과 앱 데이터만 결합한다."""
    data_context = extract_prompt_data_context(base_prompt)
    custom_prompt = str(custom_prompt or "").strip()
    return f"""
너는 중학교 과학 평가 결과를 해석하는 교육평가 전문가이다.

[분석 유형]
{analysis_label}

[사용자 직접 작성 프롬프트]
{custom_prompt}

[분석에 사용할 앱 데이터]
{data_context}

작성 원칙:
- 위의 사용자 직접 작성 프롬프트를 최우선으로 따르라.
- 웹앱의 기본 분석 목차나 기본 분석 요구사항을 따르지 말고, 사용자가 요청한 범위와 형식에 맞춰 작성하라.
- 단, 분석 근거는 위에 제공된 앱 데이터와 원안지 PDF가 있는 경우 해당 PDF에 한정하라.
- 문항번호, 정답률, 변별도, 평가영역, 성취기준 등 제공된 수치를 근거로 활용하라.
- 학생 이름은 직접 쓰지 말고 필요한 경우 '해당 학생' 또는 '학생'으로 표현하라.
""".strip()


def compose_ai_prompt_by_mode(
    base_prompt: str,
    prompt_mode: str,
    focus_request: str = "",
    custom_prompt: str = "",
    analysis_label: str = "AI 분석",
) -> str:
    """선택한 프롬프트 사용 방식에 따라 실제 AI 전달 프롬프트를 구성한다."""
    if prompt_mode == PROMPT_MODE_WITH_FOCUS:
        return append_focus_request_to_prompt(base_prompt, focus_request)
    if prompt_mode == PROMPT_MODE_DIRECT:
        return build_direct_custom_ai_prompt(base_prompt, custom_prompt, analysis_label)
    return base_prompt


BASIC_OVERALL_FOCUS_PRESETS: Dict[str, str] = {
    "직접 입력": "",
    "전체 평가 결과 종합 해석": "전체 평가 결과를 평가 전문가 관점에서 종합적으로 분석해 주세요. 평균, 표준편차, 성취수준 분포, 문항별 정답률, 평가영역별 결과, 성취기준별 결과를 연결하여 이번 평가에서 드러난 학생들의 전반적인 성취 특성과 학습 격차를 중심으로 해석해 주세요.",
    "정답률이 낮은 문항 중심 분석": "정답률이 낮은 문항을 중심으로 분석해 주세요. 단순히 어려운 문항을 나열하는 것이 아니라, 해당 문항에서 학생들이 어려움을 겪은 개념, 사고 과정, 자료 해석 요소, 적용 능력을 평가 전문가 관점에서 구체적으로 해석해 주세요.",
    "오답 선택지 반응 중심 분석": "오답 선택지 반응을 중심으로 분석해 주세요. 특정 오답 선택지에 학생 응답이 몰린 문항을 찾아, 그 선택지가 어떤 오개념이나 사고 오류를 반영할 수 있는지 평가 전문가 관점에서 구체적으로 해석해 주세요.",
    "성취수준별 차이 중심 분석": "성취수준별 차이를 중심으로 분석해 주세요. A~E 수준 학생들이 어떤 문항, 평가요소, 성취기준에서 뚜렷한 차이를 보이는지 살펴보고, 각 성취수준 집단의 학습 특성과 보완이 필요한 지점을 구체적으로 제시해 주세요.",
    "학급 간 차이 중심 분석": "학급 간 차이를 중심으로 분석해 주세요. 같은 평가를 보았음에도 학급별 정답률이나 성취수준 분포가 다르게 나타난 문항과 평가요소를 찾아, 특정 학급에서 상대적으로 강하거나 취약한 학습 내용이 무엇인지 해석해 주세요.",
    "평가영역·성취기준별 취약점 분석": "평가영역과 성취기준별 취약점을 중심으로 분석해 주세요. 학생들이 어떤 평가요소와 성취기준에서 상대적으로 낮은 성취를 보였는지 확인하고, 그 결과가 수업 내용, 개념 이해, 자료 해석, 탐구 기능과 어떻게 연결되는지 구체적으로 해석해 주세요.",
    "예상 난이도와 실제 정답률 차이 분석": "문항정보표의 예상 난이도와 실제 정답률 사이의 차이를 중심으로 분석해 주세요. 예상보다 어렵게 나타난 문항과 예상보다 쉽게 나타난 문항을 구분하고, 그 차이가 문항 구성, 발문 방식, 선택지 설계, 학생 개념 이해와 어떻게 관련될 수 있는지 해석해 주세요.",
    "수업 보완 및 피드백 방향 중심 분석": "분석 결과를 바탕으로 이후 수업 보완과 학생 피드백 방향을 중심으로 제안해 주세요. 단순한 결과 요약보다, 어떤 개념을 다시 다루어야 하는지, 어떤 학생 집단에 어떤 방식의 피드백이 필요한지, 다음 수업에서 보완할 활동은 무엇인지 중심으로 구체적으로 제시해 주세요.",
    "평가 문항의 타당도·변별도 중심 분석": "평가 문항의 타당도와 변별도를 중심으로 분석해 주세요. 정답률과 변별도가 낮거나 지나치게 높은 문항을 검토하고, 해당 문항이 성취기준과 평가요소를 적절하게 측정했는지, 학생들의 성취 차이를 타당하게 드러냈는지 평가 전문가 관점에서 해석해 주세요.",
}

BASIC_INDIVIDUAL_FOCUS_PRESETS: Dict[str, str] = {
    "직접 입력": "",
    "학생의 전체 성취 특성 분석": "선택한 학생의 전체 평가 결과를 바탕으로 성취 특성을 종합적으로 분석해 주세요. 영역총점, 성취수준, 문항별 정오 결과, 평가영역별 결과, 성취기준별 결과를 연결하여 이 학생의 전반적인 학습 상태를 구체적으로 해석해 주세요.",
    "학생의 강점과 취약점 중심 분석": "선택한 학생의 강점과 취약점을 중심으로 분석해 주세요. 어떤 평가요소와 성취기준에서 상대적으로 안정적인 성취를 보였는지, 어떤 부분에서 어려움이 나타났는지 구분하여 학생의 학습 특성이 드러나도록 해석해 주세요.",
    "평가영역·성취기준별 학생 취약점 분석": "선택한 학생의 평가영역별, 성취기준별 취약점을 중심으로 분석해 주세요. 단순히 틀린 문항을 나열하지 말고, 어떤 개념 이해, 자료 해석, 탐구 기능, 적용 능력에서 보완이 필요한지 구체적으로 제시해 주세요.",
    "정답률이 낮은 문항에서의 학생 반응 분석": "전체적으로 정답률이 낮았던 문항에서 선택한 학생이 어떤 반응을 보였는지 중심으로 분석해 주세요. 학생이 어려운 문항을 해결했는지, 또는 다른 학생들도 어려워한 문항에서 함께 어려움을 보였는지 비교하여 학습 특성을 해석해 주세요.",
    "성취수준 대비 학생의 특징 분석": "선택한 학생의 성취수준을 기준으로, 같은 성취수준 학생들과 비교했을 때 두드러지는 특징을 분석해 주세요. 해당 학생이 같은 수준의 학생들보다 강한 영역과 약한 영역을 구분하고, 추가 보완이 필요한 지점을 제시해 주세요.",
    "학생 맞춤형 피드백 문장 제안": "선택한 학생의 평가 결과를 바탕으로 학생에게 제공할 수 있는 맞춤형 피드백을 제안해 주세요. 단순히 점수를 설명하기보다, 이 학생이 이번 평가에서 비교적 안정적으로 수행한 평가영역과 성취기준, 상대적으로 보완이 필요한 평가영역과 성취기준, 다시 확인하면 좋을 문항 유형, 다음 평가 준비 방향이 드러나도록 구체적으로 작성해 주세요. 과목 특성을 고려하여 모든 과목을 수학처럼 누적 개념 복습이 다음 단원 학습의 필수 전제인 것처럼 해석하지 말고, 사회·과학처럼 단원별 성격이 강한 과목에서는 이번 평가 결과로 확인된 취약 영역, 자료 해석 방식, 개념 확인 방식, 문제 접근 전략을 중심으로 피드백해 주세요. 학생에게는 무조건 많이 복습하라는 식의 일반적인 조언보다, 이번 시험에서 강했던 부분은 유지하고 약했던 부분은 다음 시험에서 어떻게 보완하면 좋을지 구체적인 학습 행동으로 제안해 주세요.",
    "보충 학습이 필요한 개념 중심 분석": "선택한 학생에게 보충 학습이 필요한 개념을 중심으로 분석해 주세요. 틀린 문항과 낮은 성취를 보인 평가요소를 바탕으로, 다시 학습해야 할 핵심 개념과 확인해야 할 사고 과정을 구체적으로 제시해 주세요.",
    "다음 시험 준비를 위한 학습 전략 조언": "선택한 학생이 다음 시험을 준비할 때 활용할 수 있는 학습 전략을 중심으로 조언해 주세요. 이번 평가에서 드러난 강점과 취약점을 바탕으로, 다음 평가에서 우선 강화해야 할 학습 영역, 자주 틀린 문항 유형, 개념을 확인하는 방식, 자료를 읽고 해석하는 방식, 오답 정리 방법, 시험 준비 순서와 시간 배분 방향을 구체적으로 제안해 주세요. 단, 모든 과목을 수학처럼 누적 개념 복습이 다음 단원 학습의 필수 전제인 것처럼 해석하지 마세요. 사회·과학 등 단원별 성격이 강한 과목에서는 이 개념을 모르면 다음 학습이 치명적으로 어렵다는 식으로 과도하게 말하지 말고, 이번 평가 결과를 통해 확인된 학생의 학습 습관, 자료 해석 방식, 개념 확인 방식, 문제 접근 방식, 취약 영역 보완 전략을 중심으로 다음 시험 준비 방향을 제안해 주세요. 학생이 다음 시험에서 더 나은 결과를 얻기 위해 무엇을 더 많이 외워야 하는지가 아니라, 어떤 영역을 우선 점검하고 어떤 유형의 문제를 다시 풀어보며 어떤 방식으로 오답을 정리해야 하는지가 드러나도록 작성해 주세요.",
    "상위 성취로 도약하기 위한 보완점 분석": "선택한 학생이 현재 성취수준에서 한 단계 더 성장하기 위해 보완해야 할 지점을 중심으로 분석해 주세요. 단순 암기보다 개념 적용, 자료 해석, 고난도 문항 접근, 오답 분석 측면에서 어떤 학습 전략이 필요한지 구체적으로 제시해 주세요.",
}

ADVANCED_FOCUS_PRESETS: Dict[str, str] = {
    "직접 입력": "",
    "원안지 기반 전체 문항 설계 분석": "원안지의 실제 문항 구성과 통계 결과를 연결하여 전체 문항 설계를 분석해 주세요. 문항의 발문, 자료 제시 방식, 선택지 구성, 정답 근거, 요구 사고 과정이 학생 응답 결과와 어떻게 연결되는지 평가 전문가 관점에서 구체적으로 해석해 주세요.",
    "주요 문항 집중 분석": "통계 결과와 원안지 내용을 바탕으로 주요 문항을 집중 분석해 주세요. 정답률이 낮은 문항, 변별도가 낮거나 지나치게 높은 문항, 오답 선택지 반응이 두드러진 문항, 예상 난이도와 실제 결과가 어긋난 문항을 중심으로 문항 내용과 학생 반응을 연결해 해석해 주세요.",
    "오답 선택지 설계 분석": "원안지의 선택지 구성을 중심으로 오답 선택지 설계를 분석해 주세요. 오답 선택지가 단순히 틀린 답으로 기능하는지, 학생의 오개념이나 사고 오류를 드러낼 수 있는 매력적인 오답으로 구성되어 있는지, 특정 오답 선택지에 응답이 몰린 이유가 무엇일 수 있는지 평가 전문가 관점에서 검토해 주세요.",
    "발문과 자료 제시 방식 분석": "원안지의 실제 발문과 자료 제시 방식을 중심으로 분석해 주세요. 학생들이 문항을 해결하기 위해 어떤 정보를 읽고 해석해야 했는지, 발문이 명확했는지, 자료가 평가하려는 사고 과정을 적절히 유도했는지 평가 전문가 관점에서 구체적으로 검토해 주세요.",
    "성취기준과 문항 적합성 분석": "각 문항이 문항정보표의 성취기준과 평가요소를 적절하게 측정하고 있는지 분석해 주세요. 문항의 실제 요구 사고 과정이 성취기준과 일치하는지, 평가요소가 문항 내용을 충분히 설명하는지, 특정 문항이 성취기준보다 과도하거나 부족한 수준을 요구하지 않는지 검토해 주세요.",
    "정답 근거와 요구 사고 과정 분석": "각 문항의 정답 근거와 요구 사고 과정을 중심으로 분석해 주세요. 학생이 정답에 도달하기 위해 어떤 개념 이해, 자료 해석, 추론, 적용 과정을 거쳐야 하는지 설명하고, 이 과정이 학생 응답 결과와 어떻게 연결되는지 평가 전문가 관점에서 해석해 주세요.",
    "예상 난이도와 실제 정답률 차이 분석": "문항정보표의 예상 난이도와 실제 정답률 사이의 차이를 원안지 내용과 연결하여 분석해 주세요. 예상보다 어렵게 나타난 문항과 예상보다 쉽게 나타난 문항을 구분하고, 그 차이가 발문, 자료 제시, 선택지 구성, 정답 근거, 학생의 개념 이해와 어떻게 관련될 수 있는지 검토해 주세요.",
    "변별도가 낮거나 높은 문항 검토": "변별도가 낮거나 지나치게 높은 문항을 중심으로 원안지 내용을 검토해 주세요. 해당 문항이 성취 수준이 다른 학생들을 타당하게 구분했는지, 너무 쉽거나 어려웠는지, 발문이나 선택지 구성 때문에 변별도가 왜곡되었을 가능성은 없는지 평가 전문가 관점에서 분석해 주세요.",
    "학생들이 헷갈렸을 가능성이 큰 문항 분석": "학생들이 헷갈렸을 가능성이 큰 문항을 중심으로 분석해 주세요. 발문 표현, 조건 제시 방식, 자료 해석 요구, 선택지 간 유사성, 용어 사용, 정답 근거의 명확성을 검토하고, 실제 오답 반응과 연결하여 학생들이 어떤 지점에서 혼란을 겪었을 수 있는지 해석해 주세요.",
    "문항 완성도·엄밀성 검토": "원안지의 문항 완성도와 엄밀성을 중심으로 분석해 주세요. 발문이 평가 상황에 적합한 문어체로 명확하게 작성되었는지, 용어 사용이 정확하고 일관적인지, 조건이 누락되거나 모호하지 않은지, 정답이 하나로 타당하게 결정되는지, 오답 선택지의 매력도와 기능이 적절한지 검토해 주세요. 문항의 표현이 지나치게 구어적이거나, 선택지가 쉽게 배제되거나, 평가하려는 사고 과정보다 표현상의 혼란이 더 크게 작용할 가능성이 있는 문항이 있다면 구체적으로 지적해 주세요.",
    "수업 보완이 필요한 문항 추천": "원안지 내용과 통계 결과를 바탕으로 수업 보완이 필요한 문항을 추천해 주세요. 학생들이 어려움을 보인 문항 중 수업에서 다시 다루어야 할 개념, 자료 해석 방식, 탐구 기능, 적용 상황이 드러나는 문항을 우선적으로 제시하고, 각 문항을 어떻게 수업 보완이나 피드백 활동으로 연결할 수 있을지 제안해 주세요.",
}


def render_focus_request_with_preset(
    preset_label: str,
    presets: Dict[str, str],
    preset_key: str,
    text_key: str,
    placeholder: str,
    help_text: str,
    height: int = 110,
    text_area_label: str = "중점적으로 분석을 의뢰할 부분",
) -> str:
    """중점 분석 요청 프리셋 선택과 수정 가능한 입력 박스를 함께 표시한다."""
    selected_preset = st.selectbox(preset_label, list(presets.keys()), key=preset_key)
    last_key = f"{preset_key}_last_value"
    if st.session_state.get(last_key) != selected_preset:
        st.session_state[text_key] = presets.get(selected_preset, "")
        st.session_state[last_key] = selected_preset
    return st.text_area(
        text_area_label,
        placeholder=placeholder,
        help=help_text,
        height=height,
        key=text_key,
    )


def infer_advanced_scope_from_preset(preset_name: str) -> str:
    """고급 분석 프리셋 이름을 내부 분석 범위값으로 매핑한다."""
    if preset_name == "오답 선택지 설계 분석":
        return "오답 선택지 분석"
    if preset_name in ["주요 문항 집중 분석", "학생들이 헷갈렸을 가능성이 큰 문항 분석"]:
        return "주요 문항 집중 분석"
    return "원안지 기반 전체 시험 분석"


def parse_question_number_input(raw: str, max_question: Optional[int] = None) -> List[int]:
    """쉼표/공백/범위(예: 3, 5-7)로 입력한 문항번호를 정수 목록으로 정리한다."""
    if not raw:
        return []
    nums: List[int] = []
    for part in re.split(r"[,\s]+", str(raw).strip()):
        if not part:
            continue
        if re.fullmatch(r"\d+\s*[-~]\s*\d+", part):
            a, b = re.split(r"[-~]", part)
            a, b = int(a.strip()), int(b.strip())
            if a <= b:
                nums.extend(range(a, b + 1))
            else:
                nums.extend(range(b, a + 1))
        elif re.fullmatch(r"\d+", part):
            nums.append(int(part))
    nums = sorted(set(nums))
    if max_question is not None:
        nums = [n for n in nums if 1 <= n <= max_question]
    return nums


def _advanced_scope_instruction(scope: str, item_numbers: List[int], student_label: Optional[str] = None) -> str:
    item_text = ", ".join(f"{n}번" for n in item_numbers) if item_numbers else "지정 문항 없음"
    if scope == "원안지 기반 전체 시험 분석":
        return f"""
[고급 분석 범위]
- 분석 유형: 원안지 기반 전체 시험 분석
- 사용자가 지정한 문항: {item_text}
- 원안지 전체와 통계 결과를 함께 보며 시험 전체 구성, 문항 유형, 전체 성취도와 문항 구성의 관계를 중심으로 분석하라.
- 문항번호가 지정되어 있으면 해당 문항은 전체 분석 안에서 반드시 별도로 언급하라.
""".strip()
    if scope == "주요 문항 집중 분석":
        return f"""
[고급 분석 범위]
- 분석 유형: 주요 문항 집중 분석
- 분석할 문항번호: {item_text}
- 위 문항번호가 지정되어 있으면 반드시 해당 문항을 원안지에서 찾아 문항 내용, 자료, 선지, 정답률, 변별도, 성취기준과 연결해 분석하라.
- 문항번호가 비어 있으면 통계 자료상 우선 확인 필요성이 큰 문항을 직접 선정해 분석하라.
""".strip()
    if scope == "오답 선택지 분석":
        return f"""
[고급 분석 범위]
- 분석 유형: 오답 선택지 분석
- 분석할 문항번호: {item_text}
- 지정 문항의 선택지별 반응률과 원안지 선지 내용을 연결하여, 어떤 오답 선택지가 학생들에게 매력적으로 작용했을 가능성이 있는지 분석하라.
- 문항번호가 비어 있으면 특정 오답 선택지 응답이 두드러진 문항을 우선 선정해 분석하라.
""".strip()
    if scope == "성취기준 적합성 분석":
        return f"""
[고급 분석 범위]
- 분석 유형: 성취기준 적합성 분석
- 사용자가 지정한 문항: {item_text}
- 원안지 문항 내용과 문항정보표의 평가영역·성취기준이 실제로 잘 연결되는지 검토하라.
- 문항번호가 지정되어 있으면 해당 문항을 우선 분석하고, 비어 있으면 성취기준별 정답률이 낮거나 문항 구성상 검토 필요성이 큰 문항을 중심으로 분석하라.
""".strip()
    if scope == "원안지 기반 학생 개별 분석":
        return f"""
[고급 분석 범위]
- 분석 유형: 원안지 기반 학생 개별 문항 분석
- 분석 대상 학생: {student_label or '선택 학생'}
- 분석할 문항번호: {item_text if item_numbers else '학생의 전체 풀이 결과를 기준으로 오답·취약 문항 전체'}
- 선택 학생의 전체 문제 풀이 결과와 원안지를 함께 보고, 맞힌 문항/틀린 문항/성취기준/문항 내용을 연결하여 학생의 강점, 취약점, 보충 지도 방향을 분석하라.
- 문항번호가 지정되어 있으면 해당 문항을 우선 분석하고, 비어 있으면 학생의 오답 문항과 취약 성취기준을 중심으로 분석하라.
""".strip()
    return f"""
[고급 분석 범위]
- 분석 유형: {scope}
- 분석할 문항번호: {item_text}
""".strip()


def build_advanced_exam_ai_prompt(
    parsed: ParsedData,
    analysis: Dict[str, Any],
    pdf_name: str = "원안지 PDF",
    scope: str = "원안지 기반 전체 시험 분석",
    item_numbers: Optional[List[int]] = None,
    student_key: Optional[str] = None,
    anonymize_student: bool = True,
) -> str:
    """고급 분석: 업로드한 원안지 PDF와 통계 결과를 연결해 심층 해석을 생성한다."""
    item_numbers = item_numbers or []
    item = analysis["item"].copy().sort_values("정답률")
    domain = analysis["domain"].copy().sort_values("정답률")
    standard = analysis.get("standard", pd.DataFrame()).copy()
    if not standard.empty and "정답률" in standard.columns:
        standard = standard.sort_values("정답률")
    difficulty_gap = analysis.get("difficulty_gap", pd.DataFrame()).copy()
    if not difficulty_gap.empty:
        difficulty_gap = difficulty_gap[difficulty_gap["괴리여부"].isin(["불일치", "차이 있음"])].copy()
        difficulty_gap = difficulty_gap.sort_values(["정답률", "문항번호"], ascending=[True, True])

    selected_item = pd.DataFrame()
    if item_numbers:
        selected_item = analysis["item"].copy()
        selected_item["문항번호"] = pd.to_numeric(selected_item["문항번호"], errors="coerce").astype("Int64")
        selected_item = selected_item[selected_item["문항번호"].isin(item_numbers)].sort_values("문항번호")

    student_block = "학생 개별 분석 대상이 선택되지 않음"
    student_label = None
    if student_key:
        students = analysis.get("students", pd.DataFrame()).copy()
        long_df = analysis.get("long", pd.DataFrame()).copy()
        one = students[students["반/번호"] == student_key]
        if not one.empty:
            srow = one.iloc[0]
            student_label = f"{srow['반/번호']} 학생" if anonymize_student else f"{srow['반/번호']} {srow['이름']} 학생"
            student_all = long_df[long_df["반/번호"] == student_key].copy().sort_values("문항번호")
            if item_numbers:
                student_all = student_all[pd.to_numeric(student_all["문항번호"], errors="coerce").isin(item_numbers)]
            view_cols = ["문항번호", "평가영역", "성취기준", "난이도", "배점", "정답", "선택지", "정답여부", "점수"]
            existing = [c for c in view_cols if c in student_all.columns]
            student_block = f"""
[학생 개별 요약]
- 대상: {student_label}
- 최종점수: {round(float(srow['최종점수']), 2)}점
- 선택형 계산점수: {round(float(srow.get('선택형계산점수', srow['계산점수'])), 2)}점
- 환산점수: {round(float(srow['환산점수']), 2)}점
- 성취수준: {srow['성취수준']}

[학생 문항별 풀이 결과]
{student_all[existing].to_string(index=False) if not student_all.empty else '선택 학생의 문항별 풀이 결과 없음'}
""".strip()

    scope_instruction = _advanced_scope_instruction(scope, item_numbers, student_label)
    exam = parsed.exam_info
    selected_item_text = ai_percent_df(selected_item).to_string(index=False) if not selected_item.empty else "지정 문항 상세 통계 없음"

    return f"""
너는 중학교 과학 평가 문항을 분석하는 교육평가 전문가이다.

아래에는 원안지 PDF 파일과 평가 통계 결과가 함께 제공된다.
원안지의 문항 내용, 선택지 구성, 자료 제시 방식, 성취기준, 정답률, 선택지 반응률, 변별도, 성취수준별 정답률을 종합하여 시험 결과를 심층 분석하라.

{scope_instruction}

[원안지 파일]
- 파일명: {pdf_name}
- 이 PDF 전체를 확인하고, 아래 통계 자료와 연결하라.

[평가 정보]
{exam}

[성취도 요약]
{ai_percent_df(analysis['achievement']).to_string(index=False)}

[사용자 지정 문항 상세 통계]
{selected_item_text}

[정답률 낮은 문항]
{ai_percent_df(item[['문항번호','평가영역','난이도','배점','정답','정답률','변별도']].head(12)).to_string(index=False)}

[선택지 반응 포함 문항별 분석]
{ai_percent_df(item.head(20)).to_string(index=False)}

[평가영역별 분석]
{ai_percent_df(domain).to_string(index=False)}

[성취기준별 분석]
{ai_percent_df(standard).to_string(index=False) if not standard.empty else '성취기준별 분석 데이터 없음'}

[예상 난이도와 실제 정답률 판정]
{difficulty_gap[['문항번호','평가영역','예상난이도','기대정답률구간','정답률_pct','차이해석']].head(15).to_string(index=False) if not difficulty_gap.empty else '불일치 문항 없음'}

{student_block if scope == '원안지 기반 학생 개별 분석' else ''}

다음 구조를 기본으로 하되, 선택한 분석 유형에 맞지 않는 항목은 간략히 처리하고 필요한 항목을 더 자세히 작성하라.

1. 분석 대상 요약
- 선택한 분석 유형, 지정 문항, 선택 학생이 있으면 먼저 정리하라.

2. 원안지 기반 문항 내용 분석
- 원안지에서 해당 문항을 찾아 핵심 개념, 자료 제시 방식, 선지 구성, 요구 사고 과정을 분석하라.
- 지정 문항이 있으면 지정 문항을 우선적으로 다루라.

3. 통계 결과와 문항 내용의 연결
- 정답률, 변별도, 선택지 반응률, 예상 난이도와 실제 결과를 문항 내용과 연결해 해석하라.
- 특정 오답 선택지에 응답이 몰린 경우, 해당 선택지가 학생들에게 매력적으로 작용했을 가능성을 설명하라.

4. 평가영역·성취기준과의 연결
- 문항 내용이 문항정보표의 평가영역 및 성취기준과 잘 연결되는지 검토하라.
- 성취기준에 비해 문항이 지나치게 단순하거나 복합적인 경우를 설명하라.

5. 학생 개별 분석인 경우
- 선택 학생의 전체 풀이 결과를 바탕으로 맞힌 문항과 틀린 문항의 특성을 비교하라.
- 학생이 어려워한 문항의 공통점, 취약 성취기준, 보충 지도 방향을 원안지 내용과 연결해 제시하라.
- 학생 이름은 직접 쓰지 말고 '해당 학생' 또는 '학생'으로 표현하라.

6. 수업 개선 및 평가 개선 제안
- 다음 수업에서 보완할 개념, 활동, 피드백 방식을 제안하라.
- 문항 개선이 필요한 경우 문항 표현, 자료 제시, 선지 구성 측면에서 개선 방향을 제시하라.

작성 원칙:
- 원안지 내용과 통계 결과를 반드시 연결하라.
- 문항번호를 명시하며 구체적으로 분석하라.
- 단순 요약이 아니라 교사가 수업 개선과 평가 개선에 활용할 수 있는 분석을 작성하라.
""".strip()


def build_individual_ai_prompt(parsed: ParsedData, analysis: Dict[str, Any], student_key: str, anonymize: bool = True) -> str:
    students = analysis["students"]
    long_df = analysis["long"]
    domain_scores = analysis["domain_scores"]
    one = students[students["반/번호"] == student_key]
    if one.empty:
        raise ValueError("학생을 찾지 못했습니다.")
    s = one.iloc[0]
    student_label = f"{s['반/번호']} 학생" if anonymize else f"{s['반/번호']} {s['이름']} 학생"
    wrong = long_df[(long_df["반/번호"] == student_key) & (~long_df["정답여부"])].copy()
    wrong_view = wrong[["문항번호", "평가영역", "난이도", "배점", "정답", "선택지", "성취기준"]].head(20)
    domain_view = domain_scores[domain_scores["반/번호"] == student_key][["평가영역", "영역점수", "영역배점", "영역정답률"]]
    return f"""
너는 중학교 과학 교사의 학생별 평가 피드백 작성을 돕는 전문가다.
아래 데이터만 근거로 개별 학생의 학습 특성을 해석하라. 단정적 진단, 인성 평가, 과도한 추측은 금지한다.

[학생]
{student_label}

[점수 요약]
최종점수: {round(float(s['최종점수']), 2)}점
선택형 계산점수: {round(float(s.get('선택형계산점수', s['계산점수'])), 2)}점
환산점수: {round(float(s['환산점수']), 2)}점
성취수준: {s['성취수준']}

[평가영역별 결과]
{ai_percent_df(domain_view).to_string(index=False)}

[오답 문항]
{wrong_view.to_string(index=False)}

다음 형식으로 작성하라.
1. 강점으로 볼 수 있는 부분
2. 보완이 필요한 평가영역
3. 오답 문항에서 드러나는 학습 점검 지점
4. 학생에게 제공할 수 있는 피드백 문장 3개
학생 이름을 직접 쓰지 말고, '해당 학생' 또는 '학생'으로 표현하라.
""".strip()


def call_openai(api_key: str, model: str, prompt: str) -> str:
    if OpenAI is None:
        raise RuntimeError("openai 라이브러리를 불러오지 못했습니다. requirements.txt에 openai가 포함되어 있는지 확인하세요.")
    client = OpenAI(api_key=api_key)
    # Responses API 사용: max_tokens 대신 max_output_tokens 사용
    resp = client.responses.create(
        model=model,
        input=prompt,
        max_output_tokens=16000,
    )
    return getattr(resp, "output_text", "").strip()


def call_openai_stream(api_key: str, model: str, prompt: str, placeholder) -> str:
    """OpenAI Responses API를 스트리밍으로 호출하고, 생성 중인 본문을 화면에 즉시 표시한다."""
    if OpenAI is None:
        raise RuntimeError("openai 라이브러리를 불러오지 못했습니다. requirements.txt에 openai가 포함되어 있는지 확인하세요.")

    client = OpenAI(api_key=api_key)
    full_text = ""

    with client.responses.stream(
        model=model,
        input=prompt,
        max_output_tokens=16000,
    ) as stream:
        for event in stream:
            event_type = getattr(event, "type", "")
            if event_type == "response.output_text.delta":
                delta = getattr(event, "delta", "") or ""
                full_text += delta
                placeholder.markdown(full_text + "▌")

        final_response = stream.get_final_response()

    final_text = (getattr(final_response, "output_text", "") or full_text).strip()
    placeholder.markdown(final_text)
    return final_text


def call_openai_stream_with_pdf(api_key: str, model: str, prompt: str, pdf_bytes: bytes, pdf_name: str, placeholder) -> str:
    """업로드한 원안지 PDF를 OpenAI 파일 입력으로 연결하고 스트리밍 분석 결과를 표시한다."""
    if OpenAI is None:
        raise RuntimeError("openai 라이브러리를 불러오지 못했습니다. requirements.txt에 openai가 포함되어 있는지 확인하세요.")

    client = OpenAI(api_key=api_key)
    file_obj = io.BytesIO(pdf_bytes)
    file_obj.name = pdf_name

    uploaded_file = client.files.create(
        file=file_obj,
        purpose="user_data",
    )
    file_id = getattr(uploaded_file, "id", None)
    if not file_id:
        raise RuntimeError("원안지 PDF 파일 업로드에 실패했습니다.")

    full_text = ""
    response_input = [
        {
            "role": "user",
            "content": [
                {"type": "input_file", "file_id": file_id},
                {"type": "input_text", "text": prompt},
            ],
        }
    ]

    with client.responses.stream(
        model=model,
        input=response_input,
        max_output_tokens=16000,
    ) as stream:
        for event in stream:
            event_type = getattr(event, "type", "")
            if event_type == "response.output_text.delta":
                delta = getattr(event, "delta", "") or ""
                full_text += delta
                placeholder.markdown(full_text + "▌")

        final_response = stream.get_final_response()

    final_text = (getattr(final_response, "output_text", "") or full_text).strip()
    placeholder.markdown(final_text)
    return final_text




def make_ai_report_docx(
    parsed: ParsedData,
    analysis: Dict[str, Any],
    report_title: str,
    report_body: str,
    report_type: str = "AI 분석",
    focus_request: str = "",
) -> bytes:
    """AI 분석 결과를 평가정보가 포함된 Word 보고서 형태로 만든다."""
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, Inches, RGBColor
    except Exception as e:
        raise RuntimeError("Word(.docx) 저장을 위해 python-docx가 필요합니다. requirements.txt에 python-docx를 추가하세요.") from e

    def is_blank_value(v: Any) -> bool:
        return v is None or str(v).strip() in ["", "nan", "None"]

    def number_or_none(v: Any) -> Optional[float]:
        try:
            if v is None or str(v).strip() == "":
                return None
            n = float(v)
            if math.isnan(n):
                return None
            return n
        except Exception:
            return None

    def fmt_score(v: Any) -> str:
        n = number_or_none(v)
        if n is None:
            return "" if is_blank_value(v) else str(v)
        return f"{n:.1f}점"

    def fmt_count(v: Any, unit: str) -> str:
        n = number_or_none(v)
        if n is None:
            return "" if is_blank_value(v) else str(v)
        return f"{int(round(n))}{unit}"

    def fmt_raw(v: Any) -> str:
        if is_blank_value(v):
            return ""
        return str(v)

    def fmt_alpha(v: Any) -> str:
        n = number_or_none(v)
        if n is None:
            return "" if is_blank_value(v) else str(v)
        return f"{n:.3f}"

    def set_cell_shading(cell: Any, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tc_pr.append(shd)
        shd.set(qn('w:fill'), fill)

    def set_cell_border(cell: Any, color: str = "D0D7DE", size: str = "6") -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            tag = 'w:{}'.format(edge)
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), size)
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), color)

    def set_cell_text(cell: Any, text_value: str, bold: bool = False, font_size: float = 9.5, color: str = "111827") -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(str(text_value))
        r.bold = bold
        r.font.size = Pt(font_size)
        r.font.name = "맑은 고딕"
        r.font.color.rgb = RGBColor.from_string(color)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def style_table_cells(table: Any) -> None:
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "맑은 고딕"

    def add_label_value_table(title_text: str, rows: List[Tuple[str, str]], cols_per_row: int = 2) -> None:
        if not rows:
            return
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(4)
        hr = heading.add_run(title_text)
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.name = "맑은 고딕"
        hr.font.color.rgb = RGBColor.from_string("1F2937")

        pair_count = cols_per_row
        table = doc.add_table(rows=0, cols=pair_count * 2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        col_widths = [Inches(1.15), Inches(2.0)] * pair_count
        for idx, row_start in enumerate(range(0, len(rows), pair_count)):
            row = table.add_row()
            pair = rows[row_start:row_start + pair_count]
            for col in range(pair_count):
                label_cell = row.cells[col * 2]
                value_cell = row.cells[col * 2 + 1]
                if col < len(pair):
                    label, value = pair[col]
                    set_cell_text(label_cell, label, bold=True, font_size=9.0, color="374151")
                    set_cell_text(value_cell, value, bold=False, font_size=9.0, color="111827")
                    set_cell_shading(label_cell, "F3F4F6")
                    set_cell_shading(value_cell, "FFFFFF")
                else:
                    set_cell_text(label_cell, "", font_size=9.0)
                    set_cell_text(value_cell, "", font_size=9.0)
                    set_cell_shading(label_cell, "FFFFFF")
                    set_cell_shading(value_cell, "FFFFFF")
                label_cell.width = col_widths[col * 2]
                value_cell.width = col_widths[col * 2 + 1]
        style_table_cells(table)

    def add_box(title_text: str, body_text: str, fill: str = "F8FAFC", border: str = "CBD5E1") -> None:
        if not str(body_text or "").strip():
            return
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, fill)
        set_cell_border(cell, color=border, size="8")
        cell.text = ""
        p_title = cell.paragraphs[0]
        p_title.paragraph_format.space_after = Pt(4)
        r_title = p_title.add_run(title_text)
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.name = "맑은 고딕"
        r_title.font.color.rgb = RGBColor.from_string("1F2937")
        p_body = cell.add_paragraph()
        p_body.paragraph_format.line_spacing = 1.15
        p_body.paragraph_format.space_after = Pt(0)
        r_body = p_body.add_run(str(body_text).strip())
        r_body.font.size = Pt(9.5)
        r_body.font.name = "맑은 고딕"
        r_body.font.color.rgb = RGBColor.from_string("111827")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "맑은 고딕"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "맑은 고딕"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(report_title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "맑은 고딕"
    run.font.color.rgb = RGBColor.from_string("111827")

    exam = parsed.exam_info or {}
    achievement = analysis.get("achievement", pd.DataFrame())
    overall = achievement.iloc[0].to_dict() if isinstance(achievement, pd.DataFrame) and not achievement.empty else {}
    alpha = analysis.get("alpha")

    subtitle_parts = [
        fmt_raw(exam.get("학년도", "")),
        fmt_raw(exam.get("학년", "")),
        fmt_raw(exam.get("학기", "")),
        fmt_raw(exam.get("평가구분", "")),
        fmt_raw(exam.get("교과목", "")),
    ]
    subtitle_text = " ".join([p for p in subtitle_parts if p]).strip()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    sub_run = subtitle.add_run(subtitle_text if subtitle_text else report_type)
    sub_run.font.size = Pt(10)
    sub_run.font.name = "맑은 고딕"
    sub_run.font.color.rgb = RGBColor.from_string("374151")

    generated = doc.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.paragraph_format.space_after = Pt(10)
    gen_run = generated.add_run(f"{report_type} · 생성일시 {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    gen_run.font.size = Pt(8.5)
    gen_run.font.name = "맑은 고딕"
    gen_run.font.color.rgb = RGBColor.from_string("6B7280")

    doc.add_paragraph()

    base_info_rows = [
        ("학년도", fmt_raw(exam.get("학년도", ""))),
        ("학년", fmt_raw(exam.get("학년", ""))),
        ("학기", fmt_raw(exam.get("학기", ""))),
        ("평가구분", fmt_raw(exam.get("평가구분", ""))),
        ("교과목", fmt_raw(exam.get("교과목", ""))),
        ("응시자 수", fmt_count(overall.get("응시자수(명)", overall.get("응시자수", "")), "명")),
        ("과목 만점", fmt_score(exam.get("과목만점", ""))),
        ("선택형 만점", fmt_score(exam.get("선택형만점", ""))),
        ("서답형 만점", fmt_score(exam.get("서답형만점", ""))),
        ("선택형 문항 수", fmt_count(exam.get("선택형문항수", ""), "문항")),
        ("서답형 문항 수", fmt_count(exam.get("서답형문항수", ""), "문항")),
    ]
    base_info_rows = [(k, v) for k, v in base_info_rows if v]
    add_label_value_table("평가 기본 정보", base_info_rows, cols_per_row=2)

    summary_rows = [
        ("전체 평균", fmt_score(overall.get("평균(점)", overall.get("평균", ""))),),
        ("표준편차", fmt_score(overall.get("표준편차", ""))),
        ("최고점", fmt_score(overall.get("최고점(점)", overall.get("최고점", ""))),),
        ("최저점", fmt_score(overall.get("최저점(점)", overall.get("최저점", ""))),),
        ("검사신뢰도 α", fmt_alpha(alpha)),
    ]
    summary_rows = [(k, v) for k, v in summary_rows if v]
    add_label_value_table("전체 성취 요약", summary_rows, cols_per_row=2)

    focus_request = str(focus_request or "").strip()
    if focus_request:
        doc.add_paragraph()
        add_box("중점 분석 요청", focus_request, fill="F8FAFC", border="CBD5E1")

    doc.add_paragraph()
    result_heading = doc.add_paragraph()
    result_heading.paragraph_format.space_before = Pt(8)
    result_heading.paragraph_format.space_after = Pt(6)
    rr = result_heading.add_run("AI 분석 결과")
    rr.bold = True
    rr.font.size = Pt(13)
    rr.font.name = "맑은 고딕"
    rr.font.color.rgb = RGBColor.from_string("111827")

    def add_formatted_line(line: str) -> None:
        raw = line.rstrip()
        if raw == "":
            doc.add_paragraph()
            return
        text = raw.strip()
        if re.match(r"^#{1,3}\s+", text):
            level = min(len(text) - len(text.lstrip('#')), 3)
            doc.add_heading(re.sub(r"^#{1,3}\s+", "", text), level=level)
            return
        if re.match(r"^\d+\.\s+", text) and len(text) <= 90:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = "맑은 고딕"
            return
        if text.startswith(('-', '•', '*')):
            p = doc.add_paragraph(text.lstrip('-•* ').strip(), style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
            return
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.name = "맑은 고딕"

    for line in str(report_body).splitlines():
        add_formatted_line(line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

# -----------------------------------------------------------------------------
# Streamlit UI
# -----------------------------------------------------------------------------

def uploaded_file_signature(uploaded_file: Any) -> str:
    data = uploaded_file.getvalue()
    return hashlib.sha256(data).hexdigest()


def bytes_to_uploaded_like(data: bytes) -> io.BytesIO:
    bio = io.BytesIO(data)
    bio.seek(0)
    return bio


def preview_answer_file(data: bytes) -> Dict[str, Any]:
    """정오표 파일을 저장하기 전에 단독으로 읽어 기본 정보를 확인한다.
    여러 파일을 한꺼번에 올릴 때 한 파일의 오류가 전체 업로드를 막지 않게 하기 위한 사전 검사다.
    """
    a_exam, answer_key_df, sdf = parse_answer_sheet(bytes_to_uploaded_like(data))
    classes = []
    if not sdf.empty and "반" in sdf.columns:
        cls = pd.to_numeric(sdf["반"], errors="coerce").dropna().astype(int).tolist()
        classes = sorted(set(cls))
    return {
        "exam_info": a_exam,
        "question_count": int(len(answer_key_df)),
        "student_count": int(len(sdf)),
        "classes": classes,
    }


def prepare_parsed(question_file: Any, answer_files: List[Dict[str, Any]]) -> ParsedData:
    q_exam, qdf = parse_question_info(question_file)

    answer_keys: List[pd.DataFrame] = []
    student_frames: List[pd.DataFrame] = []
    answer_exam_infos: List[Dict[str, Any]] = []
    for item in answer_files:
        a_exam, answer_key_df, sdf = parse_answer_sheet(bytes_to_uploaded_like(item["data"]))
        answer_exam_infos.append(a_exam)
        answer_keys.append(answer_key_df.assign(정오표파일=item.get("name", "")))
        student_frames.append(sdf.assign(원본파일=item.get("name", "")))

    if not student_frames:
        raise ValueError("학생답 정오표 파일이 없습니다.")

    sdf_all = pd.concat(student_frames, ignore_index=True)
    # 동일 학생이 중복 업로드된 경우 마지막 업로드본만 남긴다.
    if "반/번호" in sdf_all.columns:
        sdf_all = sdf_all.drop_duplicates(subset=["반/번호"], keep="last").sort_values(["반", "번호"]).reset_index(drop=True)

    # 정답/배점 검증은 첫 번째 정오표를 대표 기준으로 삼고,
    # 여러 파일의 정오표 정답/배점이 서로 다른 경우 별도 검증 행을 추가한다.
    base_answer_key = answer_keys[0].drop(columns=["정오표파일"], errors="ignore") if answer_keys else pd.DataFrame()
    validation = make_validation(qdf, base_answer_key)
    if len(answer_keys) > 1:
        all_keys = pd.concat(answer_keys, ignore_index=True)
        key_check = all_keys.groupby("문항번호").agg(
            정답종류=("정오표_정답", lambda x: len(set(map(str, x.dropna())))),
            배점종류=("정오표_배점", lambda x: len(set(map(str, x.dropna())))),
        ).reset_index()
        diff = key_check[(key_check["정답종류"] > 1) | (key_check["배점종류"] > 1)]
        if not diff.empty:
            extra = diff.assign(정답=np.nan, 배점=np.nan, 정오표_정답="파일별 상이", 정오표_배점="파일별 상이", 정답일치=False, 배점일치=False, 검증결과="정오표 파일 간 확인 필요")
            validation = pd.concat([validation, extra[validation.columns]], ignore_index=True)

    exam_info = q_exam
    for a_exam in answer_exam_infos:
        exam_info = merge_exam_info(exam_info, a_exam)
    exam_info["정오표파일수"] = len(answer_files)
    exam_info["학생수"] = len(sdf_all)

    long_df = build_long_data(qdf, sdf_all)
    return ParsedData(
        exam_info=exam_info,
        question_df=qdf,
        students_df=sdf_all,
        long_df=long_df,
        validation_df=validation,
        answer_key_df=base_answer_key,
        original_question_df=qdf.copy(),
    )



def make_question_editor_signature(question_file: Any, answer_files: List[Dict[str, Any]]) -> str:
    """문항정보 수정표를 어떤 업로드 묶음에 연결할지 판별하는 서명."""
    h = hashlib.sha256()
    h.update(question_file.getvalue())
    for item in answer_files:
        h.update(str(item.get("name", "")).encode("utf-8"))
        h.update(item.get("data", b""))
    return h.hexdigest()


def normalize_question_editor_df(df: pd.DataFrame) -> pd.DataFrame:
    """st.data_editor에 안정적으로 넣기 위해 문항정보표 dtype과 표시 순서를 정리."""
    out = df.copy()
    preferred_cols = ["문항번호", "평가영역", "성취기준", "난이도", "배점", "정답"]
    for col in preferred_cols:
        if col not in out.columns:
            out[col] = "" if col not in ["문항번호", "배점"] else 0
    if "문항번호" in out.columns:
        out["문항번호"] = pd.to_numeric(out["문항번호"], errors="coerce").fillna(0).astype(int)
    if "배점" in out.columns:
        out["배점"] = pd.to_numeric(out["배점"], errors="coerce").fillna(0.0).astype(float)
    for text_col in ["평가영역", "성취기준", "난이도", "정답"]:
        if text_col in out.columns:
            out[text_col] = out[text_col].fillna("").astype(str)
    ordered = [c for c in preferred_cols if c in out.columns] + [c for c in out.columns if c not in preferred_cols]
    return out[ordered]


def make_eval_area_highlight_df(df: pd.DataFrame) -> pd.io.formats.style.Styler:
    """평가요소/평가영역 열을 확인하기 쉽도록 별도 미리보기 표에서 강조."""
    show_cols = [c for c in ["문항번호", "평가영역", "성취기준", "난이도", "배점", "정답"] if c in df.columns]
    view = df[show_cols].copy()

    def highlight_eval_area(row: pd.Series) -> List[str]:
        return [
            "background-color: #fff3cd; font-weight: 700;" if col == "평가영역" else ""
            for col in row.index
        ]

    return view.style.apply(highlight_eval_area, axis=1)


def is_percent_column(col: Any, series: Optional[pd.Series] = None) -> bool:
    """0~1 비율값으로 계산된 컬럼을 화면/엑셀에서 %로 표시하기 위한 판별 함수."""
    name = str(col)
    if any(k in name for k in ["률", "비율", "변별도", "수준간격차"]):
        return True
    # 성취수준별 문항 분석의 A~E 컬럼은 이름만으로는 비율인지 알기 어려우므로 값 범위로 보조 판별
    if name in list("ABCDE") and series is not None:
        nums = pd.to_numeric(series, errors="coerce").dropna()
        if not nums.empty and nums.between(-1, 1).all():
            return True
    return False


def percent_columns(df: pd.DataFrame) -> List[Any]:
    return [col for col in df.columns if is_percent_column(col, df[col])]


def unit_for_column(col: Any, series: Optional[pd.Series] = None) -> str:
    """표 출력용 헤더 단위. 내부 계산용 컬럼명은 유지하고 출력 직전에만 붙인다."""
    name = str(col)
    if is_percent_column(col, series):
        return "%"
    if name in ["응시자수", "정답자수"] or name.endswith("인원"):
        return "명"
    if name in ["문항수"]:
        return "개"
    # 표준편차는 요청에 따라 단위를 붙이지 않는다.
    if "표준편차" in name:
        return ""
    score_keywords = [
        "평균", "최고점", "최저점", "점수", "배점", "만점", "총점", "환산평균",
        "영역점수", "영역배점", "배점합계", "계산점수", "환산점수", "평균점수"
    ]
    if any(k in name for k in score_keywords):
        return "점"
    return ""


def add_unit_headers(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    rename_map = {}
    for col in out.columns:
        name = str(col)
        if re.search(r"\([%명점개]\)$", name):
            continue
        unit = unit_for_column(col, out[col])
        if unit:
            rename_map[col] = f"{name}({unit})"
    return out.rename(columns=rename_map)


def is_score_display_column(col: Any, series: Optional[pd.Series] = None) -> bool:
    """점수 계열 화면 표시용 컬럼 판별. 문항번호·정답·선택지 같은 식별/답안 컬럼은 제외한다."""
    name = str(col).strip()
    excluded_exact = {
        "문항번호", "문항", "번호", "반", "정답", "선택지", "원본표시", "정오",
        "학년도", "학기", "선택형문항수", "서답형문항수", "정오표파일수",
        "학생수", "응시자수", "정답자수", "문항수", "오답문항"
    }
    if name in excluded_exact:
        return False
    if name.endswith("문항번호") or name.endswith("번호"):
        return False
    # 비율 컬럼은 별도로 % 형식 처리한다.
    if is_percent_column(col, series):
        return False
    score_keywords = [
        "점수", "배점", "만점", "총점", "평균", "최고점", "최저점", "환산",
        "영역점수", "영역배점", "배점합계", "계산점수", "표준편차"
    ]
    if not any(k in name for k in score_keywords):
        return False
    nums = pd.to_numeric(series, errors="coerce") if series is not None else None
    return nums is None or nums.notna().any()


def score_display_columns(df: pd.DataFrame) -> List[Any]:
    return [col for col in df.columns if is_score_display_column(col, df[col])]


def format_output_df(df: pd.DataFrame, digits: int = 1, add_units: bool = True) -> pd.DataFrame:
    """화면/엑셀/AI 출력용: 비율은 %, 점수 계열은 소수점 첫째 자리로 표시하고 헤더에는 단위를 붙인다."""
    out = df.copy()
    percent_cols = percent_columns(out)
    for col in percent_cols:
        nums = pd.to_numeric(out[col], errors="coerce")
        out[col] = nums.map(lambda x: "" if pd.isna(x) else f"{x * 100:.{digits}f}%")
    for col in score_display_columns(out):
        nums = pd.to_numeric(out[col], errors="coerce")
        out[col] = nums.map(lambda x: "" if pd.isna(x) else f"{x:.1f}")
    if add_units:
        out = add_unit_headers(out)
    return out




def render_wrapped_table(df: pd.DataFrame, key: str = "wrapped-table") -> None:
    """긴 텍스트 셀이 줄바꿈되도록 읽기 전용 HTML 표로 표시한다."""
    if df is None or df.empty:
        st.info("표시할 데이터가 없습니다.")
        return

    display_df = df.copy().fillna("")

    def cell_text(value: Any) -> str:
        text = str(value)
        return html.escape(text).replace(" / ", "<br>").replace(" · ", "<br>")

    header_html = "".join(f"<th>{html.escape(str(col))}</th>" for col in display_df.columns)
    body_rows = []
    for _, row in display_df.iterrows():
        body_rows.append("<tr>" + "".join(f"<td>{cell_text(row[col])}</td>" for col in display_df.columns) + "</tr>")

    table_html = f"""
    <style>
    .{key} {{
        width: 100%;
        border-collapse: collapse;
        table-layout: fixed;
        font-size: 0.92rem;
    }}
    .{key} th {{
        background: #f7f7f8;
        border: 1px solid #e5e7eb;
        padding: 8px 10px;
        text-align: left;
        vertical-align: top;
        white-space: normal;
        word-break: keep-all;
    }}
    .{key} td {{
        border: 1px solid #e5e7eb;
        padding: 8px 10px;
        vertical-align: top;
        white-space: normal;
        overflow-wrap: anywhere;
        word-break: keep-all;
        line-height: 1.45;
    }}
    .{key} th:nth-child(1), .{key} td:nth-child(1) {{ width: 22%; }}
    .{key} th:nth-child(2), .{key} td:nth-child(2) {{ width: 10%; }}
    .{key} th:nth-child(3), .{key} td:nth-child(3) {{ width: 8%; }}
    </style>
    <table class="{key}">
        <thead><tr>{header_html}</tr></thead>
        <tbody>{''.join(body_rows)}</tbody>
    </table>
    """
    st.markdown(table_html, unsafe_allow_html=True)


def style_class_item_analysis_df(df: pd.DataFrame) -> pd.io.formats.style.Styler:
    """학급별 문항 분석에서 전체 정답률, 변별도, 학급간최대차 세 열만 같은 색으로 강조한다."""
    styles = pd.DataFrame("", index=df.index, columns=df.columns)
    highlight_cols = {"정답률", "정답률(%)", "변별도", "변별도(%)", "학급간최대차", "학급간최대차(%)", "학급간 최대차", "학급간 최대차(%)"}
    for col in df.columns:
        name = str(col).strip()
        if name in highlight_cols:
            styles[col] = "background-color: #fff3cd; font-weight: 700;"
    return df.style.apply(lambda _: styles, axis=None)




def style_domain_analysis_df(df: pd.DataFrame) -> pd.io.formats.style.Styler:
    """평가영역별 분석에서 정답률 열만 같은 색으로 강조한다."""
    styles = pd.DataFrame("", index=df.index, columns=df.columns)
    highlight_cols = {"정답률", "정답률(%)", "영역정답률", "영역정답률(%)"}
    for col in df.columns:
        name = str(col).strip()
        if name in highlight_cols:
            styles[col] = "background-color: #fff3cd; font-weight: 700;"
    return df.style.apply(lambda _: styles, axis=None)


def style_standard_analysis_df(df: pd.DataFrame) -> pd.io.formats.style.Styler:
    """성취기준별 분석에서 성취기준과 정답률 열을 같은 색으로 강조한다."""
    styles = pd.DataFrame("", index=df.index, columns=df.columns)
    highlight_cols = {"성취기준", "정답률", "정답률(%)", "성취기준정답률", "성취기준정답률(%)"}
    for col in df.columns:
        name = str(col).strip()
        if name in highlight_cols:
            styles[col] = "background-color: #fff3cd; font-weight: 700;"
    return df.style.apply(lambda _: styles, axis=None)


def style_level_item_analysis_df(df: pd.DataFrame) -> pd.io.formats.style.Styler:
    """성취수준별 문항 분석에서 수준간격차 열을 강조한다."""
    styles = pd.DataFrame("", index=df.index, columns=df.columns)
    highlight_cols = {"수준간격차", "수준간격차(%)", "수준 간 격차", "수준 간 격차(%)"}
    for col in df.columns:
        name = str(col).strip()
        if name in highlight_cols:
            styles[col] = "background-color: #fff3cd; font-weight: 700;"
    return df.style.apply(lambda _: styles, axis=None)


def style_individual_analysis_df(df: pd.DataFrame) -> pd.io.formats.style.Styler:
    """학생 개별 요약표에서 영역총점과 성취수준 열을 같은 색으로 강조한다."""
    styles = pd.DataFrame("", index=df.index, columns=df.columns)
    highlight_cols = {"영역총점", "영역총점(점)", "성취수준"}
    for col in df.columns:
        name = str(col).strip()
        if name in highlight_cols:
            styles[col] = "background-color: #fff3cd; font-weight: 700;"
    return df.style.apply(lambda _: styles, axis=None)

def style_item_analysis_df(df: pd.DataFrame) -> pd.io.formats.style.Styler:
    """문항별 분석에서 평가영역, 전체 정답률, 변별도 세 열만 같은 색으로 강조한다."""
    styles = pd.DataFrame("", index=df.index, columns=df.columns)
    highlight_cols = {"평가영역", "★ 평가요소", "정답률", "정답률(%)", "전체 정답률", "전체 정답률(%)", "변별도", "변별도(%)"}
    for col in df.columns:
        name = str(col).strip()
        if name in highlight_cols:
            styles[col] = "background-color: #fff3cd; font-weight: 700;"
    return df.style.apply(lambda _: styles, axis=None)


def fmt_percent_df(df: pd.DataFrame, digits: int = 1) -> pd.DataFrame:
    """Streamlit 화면 표시용: 데이터에 %를 포함하고 헤더에 단위를 붙인다."""
    return format_output_df(df, digits=digits, add_units=True)


def ai_percent_df(df: pd.DataFrame, digits: int = 1) -> pd.DataFrame:
    """AI 프롬프트용: 비율값을 % 문자열로 변환하고 단위 헤더를 붙인다."""
    return format_output_df(df, digits=digits, add_units=True)



def render_student_selector(individual_df: pd.DataFrame, key_prefix: str) -> Optional[str]:
    """반과 학생을 선택하고 선택 학생의 반/번호를 반환한다."""
    if individual_df.empty:
        st.info("학생 데이터가 없습니다.")
        return None

    selector_df = individual_df.copy()
    selector_df["반"] = pd.to_numeric(selector_df["반"], errors="coerce")
    selector_df["번호"] = pd.to_numeric(selector_df["번호"], errors="coerce")
    class_values = sorted([int(x) for x in selector_df["반"].dropna().unique()])
    if not class_values:
        st.info("선택할 수 있는 반 정보가 없습니다.")
        return None

    left, right = st.columns([1, 2])
    with left:
        selected_class = st.selectbox("반 선택", class_values, format_func=lambda x: f"{x}반", key=f"{key_prefix}_class")
    class_students = selector_df[selector_df["반"] == selected_class].copy().sort_values(["번호", "이름"], na_position="last")
    if class_students.empty:
        st.info("선택한 반의 학생 데이터가 없습니다.")
        return None

    class_students["선택표시"] = class_students.apply(
        lambda r: f"{int(r['번호'])}번 {r['이름']}" if pd.notna(r.get("번호")) else f"{r.get('반/번호', '')} {r.get('이름', '')}",
        axis=1,
    )
    options = class_students["반/번호"].tolist()
    label_map = dict(zip(class_students["반/번호"], class_students["선택표시"]))
    with right:
        return st.selectbox("학생 선택", options, format_func=lambda x: label_map.get(x, x), key=f"{key_prefix}_student")

def main() -> None:
    st.set_page_config(page_title="성취수준별 평가결과 분석 웹앱", layout="wide")
    st.title("성취수준별 평가결과 분석 웹앱")
    st.caption(f"{APP_VERSION} · 나이스 문항정보표/학생답 정오표 자동 분석")

    with st.expander("사용 흐름", expanded=False):
        st.markdown(
            "1. 문항정보표와 학생답 정오표를 업로드합니다.\n"
            "2. 앱이 문항정보, 정답, 배점, 학생 정오표를 자동 인식합니다.\n"
            "3. 성취수준 분할점수를 확인하고 분석 결과를 웹에서 먼저 봅니다.\n"
            "4. 확인용 엑셀과 5종 분석 엑셀 ZIP을 다운로드합니다.\n"
            "5. 필요한 경우 OpenAI API 키를 입력해 전체/개별 학생 AI 분석 초안을 생성합니다."
        )

    st.markdown(
        """
        <style>
        .big-step-header {
            display: flex;
            align-items: center;
            gap: 14px;
            margin: 34px 0 16px 0;
            padding: 18px 20px;
            border: 1px solid #d1d5db;
            border-left: 8px solid #374151;
            border-radius: 14px;
            background: linear-gradient(90deg, #f3f4f6 0%, #ffffff 100%);
            box-shadow: 0 1px 2px rgba(0,0,0,0.04);
        }
        .big-step-badge {
            min-width: 44px;
            height: 44px;
            border-radius: 999px;
            background: #374151;
            color: white;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.25rem;
            font-weight: 800;
        }
        .big-step-title {
            font-size: 1.45rem;
            font-weight: 800;
            color: #111827;
            line-height: 1.25;
        }
        .big-step-desc {
            margin-top: 4px;
            font-size: 0.95rem;
            color: #4b5563;
            line-height: 1.45;
        }
        .big-section-gap {
            height: 22px;
            margin: 36px 0 20px 0;
            border: 0;
            border-top: 2px solid rgba(75, 85, 99, 0.36);
            box-shadow: 0 -1px 0 rgba(255, 255, 255, 0.9) inset;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    def render_step_header(step_no: str, title: str, desc: str = "") -> None:
        desc_html = f"<div class='big-step-desc'>{html.escape(desc)}</div>" if desc else ""
        st.markdown(
            f"""
            <div class="big-step-header">
                <div class="big-step-badge">{html.escape(step_no)}</div>
                <div>
                    <div class="big-step-title">{html.escape(title)}</div>
                    {desc_html}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    if "answer_file_store" not in st.session_state:
        st.session_state.answer_file_store = {}

    left, right = st.columns(2)
    with left:
        question_file = st.file_uploader("문항정보표 업로드", type=["xlsx"], key="question_file")
        st.download_button(
            "문항정보표 양식 다운하기",
            make_question_info_template_xlsx(),
            file_name="문항정보표_입력양식.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            help="나이스 문항정보표를 사용하지 않는 학교는 이 양식을 내려받아 문항번호, 평가요소, 성취기준, 난이도, 배점, 정답을 입력한 뒤 문항정보표 업로드에 다시 올리면 됩니다.",
        )
        st.caption("나이스 문항정보표가 없는 경우 위 양식을 내려받아 작성한 뒤, 같은 문항정보표 업로드 칸에 올리면 나이스 문항정보표처럼 인식됩니다.")
    with right:
        uploaded_answer_files = st.file_uploader(
            "학생답 정오표 업로드",
            type=["xlsx"],
            accept_multiple_files=True,
            key="answer_files",
            help="1반 파일을 먼저 올린 뒤 2반 파일을 추가로 올려도 되고, 여러 파일을 한꺼번에 선택해도 됩니다. 같은 파일은 자동으로 중복 제외됩니다.",
        )

    added_files: List[str] = []
    duplicate_files: List[str] = []
    failed_files: List[Dict[str, str]] = []
    for f in uploaded_answer_files or []:
        data = f.getvalue()
        sig = hashlib.sha256(data).hexdigest()
        if sig in st.session_state.answer_file_store:
            duplicate_files.append(f.name)
            continue
        try:
            preview = preview_answer_file(data)
            st.session_state.answer_file_store[sig] = {
                "name": f.name,
                "size": f.size,
                "data": data,
                "student_count": preview.get("student_count", 0),
                "question_count": preview.get("question_count", 0),
                "classes": preview.get("classes", []),
            }
            added_files.append(f.name)
        except Exception as e:
            failed_files.append({"파일명": f.name, "실패 사유": str(e)})
            continue

    if st.session_state.answer_file_store or failed_files or duplicate_files:
        with st.expander("업로드된 학생답 정오표 파일", expanded=True):
            if st.session_state.answer_file_store:
                file_view = pd.DataFrame([
                    {
                        "파일명": v["name"],
                        "크기(bytes)": v["size"],
                        "인식 학생수(명)": v.get("student_count", ""),
                        "인식 문항수(개)": v.get("question_count", ""),
                        "인식 학급": ", ".join(f"{c}반" for c in v.get("classes", [])),
                    }
                    for v in st.session_state.answer_file_store.values()
                ])
                st.dataframe(file_view, use_container_width=True, hide_index=True)
            if added_files:
                st.success(f"정오표 {len(added_files)}개를 추가했습니다: " + ", ".join(added_files))
            if duplicate_files:
                st.info(f"이미 등록된 정오표 {len(duplicate_files)}개는 중복 제외했습니다: " + ", ".join(duplicate_files))
            if failed_files:
                st.warning(f"정오표 {len(failed_files)}개는 읽지 못해 제외했습니다.")
                st.dataframe(pd.DataFrame(failed_files), use_container_width=True, hide_index=True)
            if st.button("정오표 목록 초기화"):
                st.session_state.answer_file_store = {}
                st.rerun()

    answer_files = list(st.session_state.answer_file_store.values())

    # 문항정보표를 파일 업로더에서 제거한 경우에는 편집 세션을 함께 초기화합니다.
    # 같은 문항정보표 파일을 다시 올렸을 때 이전에 삭제/수정한 문항이 남지 않고,
    # 엑셀 원본을 새로 읽어 오도록 하기 위한 처리입니다.
    if not question_file:
        for key in [
            "question_info_editor_signature",
            "question_info_editor_df",
            "question_info_editor_applied_at",
            "auto_recognition_editor_signature",
            "auto_recognition_editor_values",
        ]:
            st.session_state.pop(key, None)

    if not question_file and not answer_files:
        st.markdown("""
        <div style="border:1px solid #e5e7eb; border-radius:14px; padding:18px 20px; background:#f8fafc; margin:16px 0 18px 0;">
            <div style="font-size:1.08rem; font-weight:800; color:#111827; margin-bottom:10px;">파일 업로드 안내</div>
            <div style="color:#374151; line-height:1.65; font-size:0.95rem;">
                문항정보표와 교과목별학생정오표는 반드시 나이스에서 <b>XLS data</b> 형식으로 다운로드한 파일을 업로드해 주세요.
            </div>
            <div style="margin-top:14px; color:#111827; font-weight:700;">문항정보표</div>
            <div style="color:#374151; line-height:1.65; font-size:0.94rem;">
                나이스 → [교과담임] → [정기시험] → [문항정보표관리] → 학년·과목 선택 → [조회] → [출력] → [XLS data]
            </div>
            <div style="margin-top:8px; color:#475569; line-height:1.65; font-size:0.93rem;">
                나이스 문항정보표를 사용하지 않는 경우에는 위의 <b>문항정보표 양식 다운하기</b> 버튼으로 양식을 내려받아 작성한 뒤, 문항정보표 업로드 칸에 올릴 수 있습니다.
            </div>
            <div style="margin-top:14px; color:#111827; font-weight:700;">교과목별학생정오표</div>
            <div style="color:#374151; line-height:1.65; font-size:0.94rem;">
                나이스 → [교과담임] → [정기시험조회/통계] → [교과목별학생정오표] → 강의실별 [조회] → [XLS data]
            </div>
            <div style="margin-top:8px; color:#475569; line-height:1.65; font-size:0.93rem;">
                여러 강의실을 담당하는 경우, 각 강의실을 선택하여 조회한 뒤 강의실별 정오표 파일을 각각 저장해 주세요. 저장한 파일들은 이 앱에 한꺼번에 업로드할 수 있습니다.
            </div>
            <div style="margin-top:14px; padding:10px 12px; border-radius:10px; background:#fff7ed; border:1px solid #fed7aa; color:#7c2d12; line-height:1.55; font-size:0.93rem;">
                <b>주의</b> · 두 파일 모두 XLS data 버전이어야 합니다. PDF, 화면 출력용 파일, 임의로 편집한 엑셀 파일은 정상 인식되지 않을 수 있습니다.
            </div>
        </div>
        """, unsafe_allow_html=True)

    if not question_file or not answer_files:
        st.info("문항정보표와 학생답 정오표를 모두 업로드하면 자동 분석을 시작합니다. 정오표는 여러 개 업로드할 수 있습니다.")
        return

    try:
        parsed = prepare_parsed(question_file, answer_files)
    except Exception as e:
        st.error(f"파일을 읽는 중 오류가 발생했습니다: {e}")
        st.stop()

    auto_info_signature = make_question_editor_signature(question_file, answer_files)
    auto_info_sig_key = "auto_recognition_editor_signature"
    auto_info_state_key = "auto_recognition_editor_values"

    def safe_int(value: Any, default: int = 0) -> int:
        num = pd.to_numeric(value, errors="coerce")
        if pd.isna(num):
            return int(default)
        return int(num)

    def safe_float(value: Any, default: float = 0.0) -> float:
        num = pd.to_numeric(value, errors="coerce")
        if pd.isna(num):
            return float(default)
        return float(num)

    def default_auto_info_values() -> Dict[str, Any]:
        selected_full = safe_float(parsed.exam_info.get("선택형만점"), safe_float(parsed.question_df["배점"].fillna(0).sum() if "배점" in parsed.question_df.columns else 0.0, 0.0))
        written_full = safe_float(parsed.exam_info.get("서답형만점"), 0.0)
        total_full = safe_float(parsed.exam_info.get("과목만점"), selected_full + written_full)
        return {
            "학년도": str(parsed.exam_info.get("학년도", "2026")),
            "학년": str(parsed.exam_info.get("학년", "1학년")),
            "학기": str(parsed.exam_info.get("학기", "1학기")),
            "평가구분": str(parsed.exam_info.get("평가구분", "중간고사")),
            "교과목": str(parsed.exam_info.get("교과목", "과학")),
            "선택형문항수": safe_int(parsed.exam_info.get("선택형문항수"), len(parsed.question_df)),
            "서답형문항수": safe_int(parsed.exam_info.get("서답형문항수"), 0),
            "학생수": safe_int(parsed.exam_info.get("학생수"), len(parsed.students_df)),
            "정오표파일수": safe_int(parsed.exam_info.get("정오표파일수"), len(answer_files)),
            "선택형만점": selected_full,
            "서답형만점": written_full,
            "과목만점": total_full,
        }

    if st.session_state.get(auto_info_sig_key) != auto_info_signature:
        st.session_state[auto_info_sig_key] = auto_info_signature
        st.session_state[auto_info_state_key] = default_auto_info_values()

    auto_info_values = {**default_auto_info_values(), **st.session_state.get(auto_info_state_key, {})}
    parsed.exam_info.update(auto_info_values)

    render_step_header("1", "자동 인식 결과", "업로드한 문항정보표와 학생답 정오표에서 인식한 평가 정보를 확인하고, 필요한 값은 적용 버튼으로 수정합니다.")
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("교과목", parsed.exam_info.get("교과목", "-"))
    m2.metric("학년/학기", f"{parsed.exam_info.get('학년', '-')}/{parsed.exam_info.get('학기', '-')}")
    m3.metric("선택형 문항 수", parsed.exam_info.get("선택형문항수", len(parsed.question_df)))
    m4.metric("서답형 문항 수", parsed.exam_info.get("서답형문항수", 0) if parsed.exam_info.get("서답형문항수") is not None else "확인 필요")
    m5.metric("학생 수", parsed.exam_info.get("학생수", len(parsed.students_df)))
    m6.metric("정오표 파일 수", parsed.exam_info.get("정오표파일수", len(answer_files)))

    s1, s2, s3 = st.columns(3)
    s1.metric("선택형 만점", parsed.exam_info.get("선택형만점", "-"))
    s2.metric("서답형 만점", parsed.exam_info.get("서답형만점", "-"))
    s3.metric("과목 만점", parsed.exam_info.get("과목만점", "-"))

    auto_info_key_suffix = auto_info_signature[:8]
    auto_info_widget_keys = [
        f"auto_year_{auto_info_key_suffix}",
        f"auto_grade_{auto_info_key_suffix}",
        f"auto_semester_{auto_info_key_suffix}",
        f"auto_eval_type_{auto_info_key_suffix}",
        f"auto_subject_{auto_info_key_suffix}",
        f"auto_selected_q_count_{auto_info_key_suffix}",
        f"auto_written_q_count_{auto_info_key_suffix}",
        f"auto_student_count_{auto_info_key_suffix}",
        f"auto_answer_file_count_{auto_info_key_suffix}",
        f"auto_selected_full_score_{auto_info_key_suffix}",
        f"auto_written_full_score_{auto_info_key_suffix}",
        f"auto_total_full_score_{auto_info_key_suffix}",
    ]

    with st.expander("평가정보 자동 인식값 수정", expanded=False):
        st.caption("위 자동 인식 결과에 표시된 값을 모두 수정할 수 있습니다. 수정한 값은 '자동 인식값 수정 적용' 버튼을 눌러야 아래 분석과 AI 분석에 반영됩니다. 학생 수와 정오표 파일 수는 실제 데이터 행을 바꾸는 값이 아니라 보고서와 AI 분석에 표시되는 평가 정보입니다.")
        temp_auto_info_values = auto_info_values.copy()

        cols = st.columns(5)
        temp_auto_info_values["학년도"] = cols[0].text_input("학년도", temp_auto_info_values.get("학년도", "2026"), key=f"auto_year_{auto_info_key_suffix}")
        temp_auto_info_values["학년"] = cols[1].text_input("학년", temp_auto_info_values.get("학년", "1학년"), key=f"auto_grade_{auto_info_key_suffix}")
        temp_auto_info_values["학기"] = cols[2].text_input("학기", temp_auto_info_values.get("학기", "1학기"), key=f"auto_semester_{auto_info_key_suffix}")
        temp_auto_info_values["평가구분"] = cols[3].text_input("평가구분", temp_auto_info_values.get("평가구분", "중간고사"), key=f"auto_eval_type_{auto_info_key_suffix}")
        temp_auto_info_values["교과목"] = cols[4].text_input("교과목", temp_auto_info_values.get("교과목", "과학"), key=f"auto_subject_{auto_info_key_suffix}")

        cols = st.columns(4)
        temp_auto_info_values["선택형문항수"] = int(cols[0].number_input("선택형 문항 수", min_value=0, value=safe_int(temp_auto_info_values.get("선택형문항수"), len(parsed.question_df)), step=1, key=f"auto_selected_q_count_{auto_info_key_suffix}"))
        temp_auto_info_values["서답형문항수"] = int(cols[1].number_input("서답형 문항 수", min_value=0, value=safe_int(temp_auto_info_values.get("서답형문항수"), 0), step=1, key=f"auto_written_q_count_{auto_info_key_suffix}"))
        temp_auto_info_values["학생수"] = int(cols[2].number_input("학생 수", min_value=0, value=safe_int(temp_auto_info_values.get("학생수"), len(parsed.students_df)), step=1, key=f"auto_student_count_{auto_info_key_suffix}"))
        temp_auto_info_values["정오표파일수"] = int(cols[3].number_input("정오표 파일 수", min_value=0, value=safe_int(temp_auto_info_values.get("정오표파일수"), len(answer_files)), step=1, key=f"auto_answer_file_count_{auto_info_key_suffix}"))

        cols = st.columns(3)
        temp_auto_info_values["선택형만점"] = float(cols[0].number_input("선택형 만점", min_value=0.0, value=safe_float(temp_auto_info_values.get("선택형만점"), 0.0), step=1.0, key=f"auto_selected_full_score_{auto_info_key_suffix}"))
        temp_auto_info_values["서답형만점"] = float(cols[1].number_input("서답형 만점", min_value=0.0, value=safe_float(temp_auto_info_values.get("서답형만점"), 0.0), step=1.0, key=f"auto_written_full_score_{auto_info_key_suffix}"))
        temp_auto_info_values["과목만점"] = float(cols[2].number_input("과목 만점", min_value=0.0, value=safe_float(temp_auto_info_values.get("과목만점"), 0.0), step=1.0, key=f"auto_total_full_score_{auto_info_key_suffix}"))

        btn_cols = st.columns([1, 1, 4])
        if btn_cols[0].button("자동 인식값 수정 적용", key=f"apply_auto_info_{auto_info_key_suffix}"):
            st.session_state[auto_info_state_key] = temp_auto_info_values.copy()
            st.success("자동 인식값 수정 내용이 적용되었습니다.")
            st.rerun()

        if btn_cols[1].button("원본으로 되돌리기", key=f"reset_auto_info_{auto_info_key_suffix}"):
            st.session_state[auto_info_state_key] = default_auto_info_values()
            for widget_key in auto_info_widget_keys:
                st.session_state.pop(widget_key, None)
            st.rerun()

    auto_info_values = {**default_auto_info_values(), **st.session_state.get(auto_info_state_key, {})}
    parsed.exam_info.update(auto_info_values)

    st.markdown("<div class='big-section-gap'></div>", unsafe_allow_html=True)
    render_step_header("2", "문항정보 수정", "문항별 평가요소, 성취기준, 난이도, 배점, 정답을 실제 분석 목적에 맞게 보정합니다.")
    st.caption("나이스 문항정보표에서 자동 인식한 값입니다. 평가 후 분석 자료를 더 구체화하려면 평가영역, 성취기준, 난이도 등을 여기서 수정하세요. 수정한 값은 아래 분석 결과, 확인용 엑셀, 5종 분석 엑셀, AI 분석에 모두 반영됩니다.")

    editor_signature = make_question_editor_signature(question_file, answer_files)
    editor_state_key = "question_info_editor_df"
    editor_sig_key = "question_info_editor_signature"

    # 파일 묶음이 바뀐 경우에만 원본 문항정보로 초기화하고,
    # 같은 파일에서는 사용자가 수정한 평가영역/성취기준/난이도/배점/정답을 세션에 계속 보존합니다.
    if st.session_state.get(editor_sig_key) != editor_signature:
        st.session_state[editor_sig_key] = editor_signature
        st.session_state[editor_state_key] = normalize_question_editor_df(parsed.question_df)

    editable_question_df = normalize_question_editor_df(st.session_state[editor_state_key])
    if "_row_id" not in editable_question_df.columns:
        editable_question_df["_row_id"] = [f"q_{i}_{int(num)}" for i, num in enumerate(editable_question_df["문항번호"].tolist())]
        st.session_state[editor_state_key] = editable_question_df.copy()

    st.info("가로 입력칸에서 수정한 평가요소, 성취기준, 난이도, 배점, 정답은 아래 분석과 AI 분석에 반영됩니다. 수정 후에는 반드시 아래의 '문항정보 수정값 적용' 버튼을 눌러 주세요. 파일을 다시 올리거나 정오표 목록을 초기화하기 전까지 적용한 수정값이 유지됩니다.")
    st.warning("평가요소는 이후 평가영역별 분석과 AI 분석의 핵심 기준이 되므로, 문항정보표의 내용을 그대로 사용하기보다 반드시 문항의 실제 평가 내용을 반영하도록 수정해 주세요. 특히 평가영역이 단원명이나 큰 주제처럼 넓게 입력되어 있다면, 해당 문항이 실제로 평가하는 개념, 사고 과정, 자료 해석 능력, 적용 상황 등이 드러나도록 구체적으로 보완해야 합니다. 평가요소가 자세할수록 문항별 정답률, 오답 경향, 성취수준별 차이를 더 정확하고 의미 있게 해석할 수 있습니다.")
    st.caption("문항이 잘못 인식되었거나 누락된 경우 아래의 '+ 문항 추가'를 사용하고, 삭제할 문항은 오른쪽 삭제 칸을 체크한 뒤 하단의 '선택 문항 삭제'를 누르세요.")

    editable_question_df = normalize_question_editor_df(st.session_state[editor_state_key])
    if "_row_id" not in editable_question_df.columns:
        editable_question_df["_row_id"] = [f"q_{i}_{int(num)}" for i, num in enumerate(editable_question_df["문항번호"].tolist())]

    header_cols = st.columns([0.75, 2.4, 2.4, 0.9, 0.8, 0.9, 0.55])
    headers = ["문항", "★ 평가요소", "성취기준", "난이도", "배점", "정답", "삭제"]
    for col, label in zip(header_cols, headers):
        col.markdown(f"<div style='text-align:center; font-weight:700;'>{label}</div>", unsafe_allow_html=True)

    edited_rows: List[Dict[str, Any]] = []
    with st.form(key=f"question_info_editor_form_{editor_signature[:12]}", clear_on_submit=False):
        for idx, row in editable_question_df.reset_index(drop=True).iterrows():
            row_id = str(row.get("_row_id", f"row_{idx}"))
            c1, c2, c3, c4, c5, c6, c7 = st.columns([0.75, 2.4, 2.4, 0.9, 0.8, 0.9, 0.55])
            qno_default = int(pd.to_numeric(row.get("문항번호", idx + 1), errors="coerce") or idx + 1)
            q_no_text = c1.text_input("문항번호", value=str(qno_default), key=f"qno_{editor_signature[:8]}_{row_id}", label_visibility="collapsed", placeholder="번호")
            q_no = int(pd.to_numeric(q_no_text, errors="coerce") or 0)
            eval_area = c2.text_input("평가요소", value=str(row.get("평가영역", "")), key=f"eval_{editor_signature[:8]}_{row_id}", label_visibility="collapsed", placeholder="문항의 실제 평가 내용을 구체적으로 입력")
            standard = c3.text_input("성취기준", value=str(row.get("성취기준", "")), key=f"std_{editor_signature[:8]}_{row_id}", label_visibility="collapsed")
            diff_value = str(row.get("난이도", ""))
            diff_options = ["", "어려움", "보통", "쉬움"]
            diff_index = diff_options.index(diff_value) if diff_value in diff_options else 0
            difficulty = c4.selectbox("난이도", options=diff_options, index=diff_index, key=f"diff_{editor_signature[:8]}_{row_id}", label_visibility="collapsed")
            score = c5.number_input("배점", min_value=0.0, step=0.1, value=float(pd.to_numeric(row.get("배점", 0), errors="coerce") or 0.0), key=f"score_{editor_signature[:8]}_{row_id}", label_visibility="collapsed", format="%.1f")
            answer = c6.text_input("정답", value=str(row.get("정답", "")), key=f"ans_{editor_signature[:8]}_{row_id}", label_visibility="collapsed", placeholder="예: 3 또는 2,3")
            delete_row = c7.checkbox("삭제", value=False, key=f"del_{editor_signature[:8]}_{row_id}", label_visibility="collapsed")
            edited_rows.append({"문항번호": q_no, "평가영역": eval_area, "성취기준": standard, "난이도": difficulty, "배점": score, "정답": answer, "_row_id": row_id, "삭제": delete_row})

        action_cols = st.columns([1.4, 1.0, 1.2, 3.1])
        apply_question_edits = action_cols[0].form_submit_button("문항정보 수정값 적용", type="primary", use_container_width=True)
        add_question_row = action_cols[1].form_submit_button("+ 문항 추가", use_container_width=True)
        delete_question_rows = action_cols[2].form_submit_button("선택 문항 삭제", use_container_width=True)

    if apply_question_edits or delete_question_rows or add_question_row:
        edited_question_df = pd.DataFrame(edited_rows)
        if delete_question_rows:
            edited_question_df = edited_question_df[~edited_question_df["삭제"]].copy()
        edited_question_df = edited_question_df.drop(columns=["삭제"], errors="ignore")
        edited_question_df["문항번호"] = pd.to_numeric(edited_question_df["문항번호"], errors="coerce")
        edited_question_df = edited_question_df[edited_question_df["문항번호"].notna() & (edited_question_df["문항번호"] > 0)].copy()
        if add_question_row:
            max_q = pd.to_numeric(edited_question_df["문항번호"], errors="coerce").fillna(0).max() if not edited_question_df.empty else 0
            new_id = f"new_{datetime.now().strftime('%H%M%S%f')}"
            new_row = {"문항번호": int(max_q) + 1, "평가영역": "", "성취기준": "", "난이도": "보통", "배점": 0.0, "정답": "", "_row_id": new_id}
            edited_question_df = pd.concat([edited_question_df, pd.DataFrame([new_row])], ignore_index=True)
        if edited_question_df.empty:
            st.error("적용할 문항이 없습니다. 최소 1개 이상의 문항을 남겨 주세요.")
            st.stop()
        edited_question_df = normalize_question_editor_df(edited_question_df)
        if "_row_id" not in edited_question_df.columns:
            edited_question_df["_row_id"] = [f"q_{i}_{int(num)}" for i, num in enumerate(edited_question_df["문항번호"].tolist())]
        edited_question_df = edited_question_df.sort_values("문항번호").reset_index(drop=True)
        st.session_state[editor_state_key] = edited_question_df.copy()
        st.session_state["question_info_editor_applied_at"] = datetime.now().strftime("%H:%M:%S")
        if delete_question_rows:
            st.success("선택한 문항을 삭제하고 문항정보 수정값을 적용했습니다.")
        elif add_question_row:
            st.success("새 문항을 추가했습니다. 내용을 입력한 뒤 문항정보 수정값 적용을 눌러 주세요.")
        else:
            st.success("문항정보 수정값을 적용했습니다. 아래 요약과 분석 결과에 적용된 값이 반영됩니다.")
        st.rerun()
    else:
        edited_question_df = normalize_question_editor_df(st.session_state[editor_state_key])

    applied_at = st.session_state.get("question_info_editor_applied_at")
    if applied_at:
        st.caption(f"마지막 적용 시각: {applied_at}")

    # 편집값 정규화 후 전체 분석 데이터에 재반영
    parsed.question_df = edited_question_df.drop(columns=["_row_id"], errors="ignore").copy()
    parsed.question_df["문항번호"] = pd.to_numeric(parsed.question_df["문항번호"], errors="coerce").fillna(0).astype(int)
    parsed.question_df["배점"] = pd.to_numeric(parsed.question_df["배점"], errors="coerce").fillna(0.0)
    parsed.question_df["정답"] = parsed.question_df["정답"].apply(normalize_choice_answer_value)
    parsed.long_df = build_long_data(parsed.question_df, parsed.students_df)
    parsed.validation_df = make_validation(parsed.question_df, parsed.answer_key_df)

    with st.expander("수정된 문항정보 요약", expanded=False):
        st.caption("성취기준별로 평가요소, 난이도 배치, 문항 구성, 총점을 카드형으로 확인할 수 있습니다. 평가요소를 수정한 뒤 각 성취기준 안에서 문항의 평가 내용이 충분히 구체적으로 구분되는지 점검하세요.")

        q_summary = parsed.question_df.copy()
        q_summary["성취기준"] = q_summary.get("성취기준", "").fillna("").astype(str).str.strip().replace("", "미입력")
        q_summary["평가영역"] = q_summary.get("평가영역", "").fillna("").astype(str).str.strip().replace("", "미입력")
        q_summary["난이도"] = q_summary.get("난이도", "").fillna("").astype(str).str.strip().replace("", "미입력")
        q_summary["배점"] = pd.to_numeric(q_summary.get("배점", 0), errors="coerce").fillna(0.0)
        q_summary["문항번호"] = pd.to_numeric(q_summary.get("문항번호", 0), errors="coerce").fillna(0).astype(int)

        def unique_text_list(values: pd.Series) -> List[str]:
            items: List[str] = []
            for value in values:
                text = str(value).strip()
                if text and text not in items:
                    items.append(text)
            return items

        def difficulty_parts(group: pd.DataFrame) -> List[str]:
            order = ["어려움", "보통", "쉬움", "미입력"]
            parts: List[str] = []
            for diff in order:
                sub = group[group["난이도"] == diff]
                if not sub.empty:
                    parts.append(f"{diff} {len(sub)}문항 / {sub['배점'].sum():.1f}점")
            return parts

        def item_lines(group: pd.DataFrame) -> List[str]:
            rows: List[str] = []
            for _, row in group.sort_values("문항번호").iterrows():
                rows.append(f"{int(row['문항번호'])}번 · {row['난이도']} · {row['배점']:.1f}점 · {row['평가영역']}")
            return rows

        def render_standard_card(standard: str, group: pd.DataFrame, idx: int) -> None:
            total_score = group["배점"].sum()
            item_count = len(group)
            eval_items = unique_text_list(group["평가영역"])
            diff_items = difficulty_parts(group)
            question_items = item_lines(group)

            eval_html = "".join(f"<li>{html.escape(item)}</li>" for item in eval_items) or "<li>미입력</li>"
            diff_html = "".join(f"<span class='summary-chip'>{html.escape(item)}</span>" for item in diff_items) or "<span class='summary-chip'>미입력</span>"
            question_html = "".join(f"<li>{html.escape(item)}</li>" for item in question_items) or "<li>미입력</li>"

            card_html = f"""
            <div class="standard-summary-card">
                <div class="standard-card-standard">{html.escape(str(standard))}</div>
                <div class="standard-card-metrics">
                    <div><span>총점</span><strong>{total_score:.1f}점</strong></div>
                    <div><span>문항수</span><strong>{item_count}문항</strong></div>
                </div>
                <div class="standard-card-section">
                    <div class="standard-card-label">난이도 구성</div>
                    <div class="summary-chip-wrap">{diff_html}</div>
                </div>
                <div class="standard-card-section">
                    <div class="standard-card-label">평가요소</div>
                    <ul>{eval_html}</ul>
                </div>
                <div class="standard-card-section">
                    <div class="standard-card-label">문항 구성</div>
                    <ul>{question_html}</ul>
                </div>
            </div>
            """
            st.markdown(card_html, unsafe_allow_html=True)

        st.markdown(
            """
            <style>
            .standard-summary-card {
                border: 1px solid #e5e7eb;
                border-radius: 14px;
                padding: 18px 20px;
                margin: 14px 0;
                background: #ffffff;
                box-shadow: 0 1px 3px rgba(15, 23, 42, 0.06);
            }
            .standard-card-standard {
                font-size: 1.02rem;
                font-weight: 700;
                line-height: 1.5;
                color: #111827;
                margin-bottom: 12px;
                word-break: keep-all;
                overflow-wrap: anywhere;
            }
            .standard-card-metrics {
                display: flex;
                gap: 10px;
                flex-wrap: wrap;
                margin-bottom: 12px;
            }
            .standard-card-metrics div {
                background: #f9fafb;
                border: 1px solid #eef0f3;
                border-radius: 10px;
                padding: 8px 12px;
                min-width: 110px;
            }
            .standard-card-metrics span {
                display: block;
                color: #6b7280;
                font-size: 0.82rem;
                margin-bottom: 2px;
            }
            .standard-card-metrics strong {
                color: #111827;
                font-size: 1rem;
            }
            .standard-card-section {
                margin-top: 12px;
            }
            .standard-card-label {
                font-weight: 700;
                color: #374151;
                margin-bottom: 6px;
            }
            .standard-summary-card ul {
                margin: 4px 0 0 1.1rem;
                padding: 0;
                line-height: 1.55;
            }
            .standard-summary-card li {
                margin: 3px 0;
                word-break: keep-all;
                overflow-wrap: anywhere;
            }
            .summary-chip-wrap {
                display: flex;
                gap: 6px;
                flex-wrap: wrap;
            }
            .summary-chip {
                display: inline-block;
                background: #f3f4f6;
                border: 1px solid #e5e7eb;
                border-radius: 999px;
                padding: 5px 9px;
                font-size: 0.9rem;
                color: #374151;
            }
            </style>
            """,
            unsafe_allow_html=True,
        )

        if q_summary.empty:
            st.info("표시할 문항정보가 없습니다.")
        else:
            for idx, (standard, group) in enumerate(q_summary.groupby("성취기준", dropna=False, sort=False), start=1):
                render_standard_card(str(standard), group, idx)

    st.markdown("<div class='big-section-gap'></div>", unsafe_allow_html=True)
    render_step_header("3", "분석 기준 및 결과 확인", "성취수준 분할점수를 설정한 뒤, 탭별 분석 결과와 AI 분석을 확인합니다.")
    st.caption("선택형 만점, 서답형 만점, 과목 만점은 1. 자동 인식 결과의 '평가정보 자동 인식값 수정'에서 조정할 수 있습니다.")
    parsed.exam_info["선택형만점"] = safe_float(parsed.exam_info.get("선택형만점"), safe_float(parsed.question_df["배점"].fillna(0).sum() if "배점" in parsed.question_df.columns else 0.0, 0.0))
    parsed.exam_info["서답형만점"] = safe_float(parsed.exam_info.get("서답형만점"), 0.0)
    total_full_score = safe_float(parsed.exam_info.get("과목만점"), parsed.exam_info["선택형만점"] + parsed.exam_info["서답형만점"] or 100.0)
    c1, c2, c3, c4 = st.columns(4)
    cut_a = c1.number_input("A/B 분할점수", value=round(total_full_score * 0.9, 1), step=1.0, key="cut_a_raw")
    cut_b = c2.number_input("B/C 분할점수", value=round(total_full_score * 0.8, 1), step=1.0, key="cut_b_raw")
    cut_c = c3.number_input("C/D 분할점수", value=round(total_full_score * 0.7, 1), step=1.0, key="cut_c_raw")
    cut_d = c4.number_input("D/E 분할점수", value=round(total_full_score * 0.6, 1), step=1.0, key="cut_d_raw")
    cuts = {"A": cut_a, "B": cut_b, "C": cut_c, "D": cut_d}

    analysis = analyze_all(parsed, total_full_score, cuts)

    warn_df = parsed.validation_df[parsed.validation_df["검증결과"] == "확인 필요"] if not parsed.validation_df.empty else pd.DataFrame()
    if warn_df.empty:
        st.success("문항정보표와 학생답 정오표의 정답/배점 검증 결과가 정상입니다.")
    else:
        st.warning("문항정보표와 학생답 정오표의 정답/배점이 다른 문항이 있습니다. 검증결과 탭에서 확인하세요.")

    analysis_tab_labels = [
        "데이터 확인", "성취도 분석", "문항별 분석", "학급별 분석", "평가영역별 분석", "성취기준별 분석", "성취수준별 분석", "학생 개별", "AI 분석"
    ]
    selected_analysis_tab = st.radio(
        "분석 영역 선택",
        analysis_tab_labels,
        horizontal=True,
        label_visibility="collapsed",
        key="main_analysis_section_selector",
    )

    if selected_analysis_tab == "데이터 확인":
        with st.container(border=True):
            st.markdown("#### 문항정보")
            st.dataframe(parsed.question_df, use_container_width=True, height=260)
        with st.container(border=True):
            st.markdown("#### 학생 정오표")
            st.dataframe(parsed.students_df, use_container_width=True, height=260)
        with st.container(border=True):
            st.markdown("#### 검증결과")
            st.dataframe(parsed.validation_df, use_container_width=True, height=220)

    elif selected_analysis_tab == "성취도 분석":
        alpha = analysis.get("alpha")
        with st.container(border=True):
            a1, a2, a3, a4 = st.columns(4)
            a1.metric("평균(점)", f"{analysis['achievement'].loc[0, '평균']:.2f}")
            a2.metric("표준편차", f"{analysis['achievement'].loc[0, '표준편차']:.2f}")
            a3.metric("최고/최저(점)", f"{analysis['achievement'].loc[0, '최고점']:.1f} / {analysis['achievement'].loc[0, '최저점']:.1f}")
            a4.metric("검사신뢰도 α", "-" if alpha is None else f"{alpha:.3f}")
            st.dataframe(fmt_percent_df(analysis["achievement"]), use_container_width=True)

            achievement_std = safe_float(analysis["achievement"].loc[0, "표준편차"], 0.0)
            std_ratio = (achievement_std / total_full_score * 100) if total_full_score else 0.0
            st.info(
                "표준편차는 학생들의 점수가 평균을 중심으로 얼마나 흩어져 있는지를 보여주는 값입니다. "
                "시험 만점과 점수 범위의 영향을 받으므로 절대 점수만으로 크고 작음을 판단하기보다 "
                f"만점 대비 비율을 함께 보는 것이 좋습니다. 현재 표준편차는 만점 대비 {std_ratio:.1f}%입니다. "
                "대략 10% 이하는 점수 분포가 좁은 편, 10~20%는 어느 정도 차이가 있는 편, "
                "20% 이상은 학생 간 점수 차이가 큰 편으로 참고할 수 있습니다. "
                "최종 해석은 평균 점수, 성취수준 분포, 문항별 정답률과 함께 판단하세요."
            )

        graph_col1, graph_col2 = st.columns(2)

        with graph_col1:
            with st.container(border=True):
                st.markdown("#### 전체 분석 그래프")
                dist_chart_df, summary_chart_df, class_label_expr, chart_y_max = make_class_score_distribution_chart_data(
                    analysis["individual"], analysis["class_achievement"], full_score=total_full_score, bin_size=5
                )
                if not summary_chart_df.empty:
                    dist_for_chart = dist_chart_df.copy()
                    summary_for_chart = summary_chart_df.copy()
                    dist_for_chart["그래프요소"] = "분포"
                    summary_for_chart["그래프요소"] = "요약"
                    chart_source_df = pd.concat([dist_for_chart, summary_for_chart], ignore_index=True, sort=False)
    
                    st.vega_lite_chart(
                        chart_source_df,
                        {
                            "height": 390,
                            "width": "container",
                            "resolve": {"scale": {"x": "shared", "y": "shared"}},
                            "config": {
                                "axis": {"labelFontSize": 12, "titleFontSize": 13, "grid": True},
                                "view": {"stroke": "transparent"},
                            },
                            "layer": [
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '분포'"}],
                                    "mark": {
                                        "type": "rect",
                                        "cornerRadius": 2,
                                        "color": "#94A3B8",
                                        "opacity": 0.8,
                                    },
                                    "encoding": {
                                        "x": {
                                            "field": "왼쪽폭",
                                            "type": "quantitative",
                                            "axis": {
                                                "title": "학급",
                                                "labelAngle": 0,
                                                "values": summary_chart_df["학급위치"].astype(int).tolist(),
                                                "labelExpr": class_label_expr,
                                            },
                                        },
                                        "x2": {"field": "오른쪽폭"},
                                        "y": {
                                            "field": "점수구간하한",
                                            "type": "quantitative",
                                            "axis": {"title": "점수"},
                                            "scale": {"domain": [0, chart_y_max]},
                                        },
                                        "y2": {"field": "점수구간상한"},
                                        "tooltip": [
                                            {"field": "학급", "type": "nominal", "title": "학급"},
                                            {"field": "점수구간", "type": "nominal", "title": "점수 구간"},
                                            {"field": "도수(명)", "type": "quantitative", "title": "도수(명)"},
                                        ],
                                    },
                                },
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '요약'"}],
                                    "mark": {
                                        "type": "rule",
                                        "strokeWidth": 5,
                                        "color": "#111827",
                                        "opacity": 0.95,
                                    },
                                    "encoding": {
                                        "x": {"field": "학급위치", "type": "quantitative"},
                                        "y": {"field": "최저점", "type": "quantitative"},
                                        "y2": {"field": "최고점"},
                                        "tooltip": [
                                            {"field": "학급", "type": "nominal", "title": "학급"},
                                            {"field": "최고점", "type": "quantitative", "format": ".1f", "title": "최고점"},
                                            {"field": "평균", "type": "quantitative", "format": ".1f", "title": "평균"},
                                            {"field": "최저점", "type": "quantitative", "format": ".1f", "title": "최저점"},
                                        ],
                                    },
                                },
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '요약'"}],
                                    "mark": {"type": "tick", "thickness": 4, "size": 36, "color": "#111827"},
                                    "encoding": {
                                        "x": {"field": "학급위치", "type": "quantitative"},
                                        "y": {"field": "최고점", "type": "quantitative"},
                                    },
                                },
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '요약'"}],
                                    "mark": {"type": "tick", "thickness": 4, "size": 36, "color": "#111827"},
                                    "encoding": {
                                        "x": {"field": "학급위치", "type": "quantitative"},
                                        "y": {"field": "최저점", "type": "quantitative"},
                                    },
                                },
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '요약'"}],
                                    "mark": {
                                        "type": "text",
                                        "align": "left",
                                        "baseline": "middle",
                                        "dx": 10,
                                        "fontSize": 11,
                                        "fontWeight": "bold",
                                        "color": "#111827",
                                    },
                                    "encoding": {
                                        "x": {"field": "학급위치", "type": "quantitative"},
                                        "y": {"field": "최고점", "type": "quantitative"},
                                        "text": {"field": "최고라벨", "type": "nominal"},
                                    },
                                },
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '요약'"}],
                                    "mark": {
                                        "type": "text",
                                        "align": "left",
                                        "baseline": "middle",
                                        "dx": 10,
                                        "fontSize": 11,
                                        "fontWeight": "bold",
                                        "color": "#111827",
                                    },
                                    "encoding": {
                                        "x": {"field": "학급위치", "type": "quantitative"},
                                        "y": {"field": "최저점", "type": "quantitative"},
                                        "text": {"field": "최저라벨", "type": "nominal"},
                                    },
                                },
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '요약'"}],
                                    "mark": {
                                        "type": "point",
                                        "filled": True,
                                        "shape": "diamond",
                                        "size": 260,
                                        "color": "#E11D48",
                                        "stroke": "white",
                                        "strokeWidth": 2.5,
                                    },
                                    "encoding": {
                                        "x": {"field": "학급위치", "type": "quantitative"},
                                        "y": {"field": "평균", "type": "quantitative"},
                                        "tooltip": [
                                            {"field": "학급", "type": "nominal", "title": "학급"},
                                            {"field": "평균", "type": "quantitative", "format": ".1f", "title": "평균"},
                                        ],
                                    },
                                },
                                {
                                    "transform": [{"filter": "datum.그래프요소 == '요약'"}],
                                    "mark": {
                                        "type": "text",
                                        "align": "left",
                                        "baseline": "middle",
                                        "dx": 12,
                                        "dy": -1,
                                        "fontSize": 12,
                                        "fontWeight": "bold",
                                        "color": "#E11D48",
                                    },
                                    "encoding": {
                                        "x": {"field": "학급위치", "type": "quantitative"},
                                        "y": {"field": "평균", "type": "quantitative"},
                                        "text": {"field": "평균라벨", "type": "nominal"},
                                    },
                                },
                            ],
                        },
                        use_container_width=True,
                    )
                else:
                    st.info("표시할 학급별 점수 데이터가 없습니다.")
    
                st.markdown("##### 전체 성취수준별 비율")
                total_level_chart_df = make_total_level_distribution_chart_df(analysis["individual"])
                if not total_level_chart_df.empty:
                    st.dataframe(fmt_percent_df(total_level_chart_df), use_container_width=True, hide_index=True)
                else:
                    st.info("표시할 전체 성취수준 데이터가 없습니다.")

        with graph_col2:
            with st.container(border=True):
                st.markdown("#### 개별 반 분석 그래프")
                class_values = sorted(analysis["individual"]["반"].dropna().unique().tolist())
                if class_values:
                    selected_class_for_graph = st.selectbox(
                        "분석할 반 선택",
                        class_values,
                        format_func=lambda x: f"{int(x)}반" if pd.notna(x) and float(x).is_integer() else f"{x}반",
                        key="achievement_class_graph_select",
                    )
                    class_level_chart_df = make_class_level_distribution_chart_df(analysis["individual"], selected_class_for_graph)
                    if not class_level_chart_df.empty:
                        class_chart_data = class_level_chart_df.copy()
                        total_chart_data = make_total_level_distribution_chart_df(analysis["individual"])
                        chart_data = class_chart_data.merge(
                            total_chart_data[["성취수준", "학생수(명)", "비율"]].rename(
                                columns={"학생수(명)": "전체학생수(명)", "비율": "전체비율"}
                            ),
                            on="성취수준",
                            how="left",
                        )
                        chart_data["선택반비율(%)"] = pd.to_numeric(chart_data["비율"], errors="coerce").fillna(0) * 100
                        chart_data["전체비율(%)"] = pd.to_numeric(chart_data["전체비율"], errors="coerce").fillna(0) * 100
                        y_max_ratio = max(
                            5,
                            float(chart_data[["선택반비율(%)", "전체비율(%)"]].max().max() or 0) + 5,
                        )
                        y_max_ratio = min(100, max(10, y_max_ratio))
                        selected_class_label = (
                            f"{int(selected_class_for_graph)}반"
                            if pd.notna(selected_class_for_graph) and float(selected_class_for_graph).is_integer()
                            else f"{selected_class_for_graph}반"
                        )
                        grouped_chart_data = []
                        for _, row in chart_data.iterrows():
                            grouped_chart_data.append({
                                "성취수준": row.get("성취수준"),
                                "구분": "전체",
                                "비율(%)": row.get("전체비율(%)", 0),
                                "학생수(명)": row.get("전체학생수(명)", 0),
                            })
                            grouped_chart_data.append({
                                "성취수준": row.get("성취수준"),
                                "구분": selected_class_label,
                                "비율(%)": row.get("선택반비율(%)", 0),
                                "학생수(명)": row.get("학생수(명)", 0),
                            })
                        grouped_chart_df = pd.DataFrame(grouped_chart_data)
                        st.caption(f"빨간색 막대는 전체 성취수준 비율, 파란색 막대는 {selected_class_label} 성취수준 비율입니다.")
                        st.vega_lite_chart(
                            grouped_chart_df,
                            {
                                "height": 340,
                                "width": "container",
                                "config": {
                                    "axis": {"labelFontSize": 12, "titleFontSize": 13, "grid": True},
                                    "view": {"stroke": "transparent"},
                                },
                                "mark": {
                                    "type": "bar",
                                    "cornerRadiusTopLeft": 5,
                                    "cornerRadiusTopRight": 5,
                                    "opacity": 0.88,
                                },
                                "encoding": {
                                    "x": {
                                        "field": "성취수준",
                                        "type": "nominal",
                                        "axis": {"title": "성취수준"},
                                        "sort": ["A", "B", "C", "D", "E"],
                                    },
                                    "xOffset": {
                                        "field": "구분",
                                        "type": "nominal",
                                        "sort": ["전체", selected_class_label],
                                    },
                                    "y": {
                                        "field": "비율(%)",
                                        "type": "quantitative",
                                        "axis": {"title": "비율(%)"},
                                        "scale": {"domain": [0, y_max_ratio], "nice": False, "zero": True},
                                    },
                                    "color": {
                                        "field": "구분",
                                        "type": "nominal",
                                        "scale": {
                                            "domain": ["전체", selected_class_label],
                                            "range": ["#E11D48", "#2563EB"],
                                        },
                                        "legend": None,
                                    },
                                    "tooltip": [
                                        {"field": "구분", "type": "nominal", "title": "구분"},
                                        {"field": "성취수준", "type": "nominal", "title": "성취수준"},
                                        {"field": "학생수(명)", "type": "quantitative", "format": "d", "title": "학생수"},
                                        {"field": "비율(%)", "type": "quantitative", "format": ".1f", "title": "비율(%)"},
                                    ],
                                },
                            },
                            use_container_width=True,
                        )
                        comparison_table = chart_data[["성취수준", "학생수(명)", "비율", "전체학생수(명)", "전체비율"]].rename(
                            columns={"학생수(명)": "선택반학생수(명)", "비율": "선택반비율"}
                        )
                        st.dataframe(fmt_percent_df(comparison_table), use_container_width=True, hide_index=True)
                    else:
                        st.info("선택한 반의 성취수준 데이터가 없습니다.")
                else:
                    st.info("표시할 학급 데이터가 없습니다.")
    
        with st.container(border=True):
            st.markdown("#### 학급별 성취도")
            st.dataframe(fmt_percent_df(analysis["class_achievement"]), use_container_width=True)

    elif selected_analysis_tab == "문항별 분석":
        with st.container(border=True):
            st.markdown("정답률과 변별도를 기준으로 취약 문항을 먼저 확인할 수 있습니다.")
            st.info(
                "이 표는 정답률이 낮은 문항부터 배열되어 있어 학생들이 상대적으로 어려워한 문항을 먼저 확인할 수 있습니다. "
                "각 열 제목을 클릭하면 정답률, 변별도, 평가영역, 난이도, 선택지별 응답 비율 등을 기준으로 "
                "오름차순·내림차순 정렬을 바꿀 수 있습니다.\n\n"
                "변별도는 상위 집단과 하위 집단의 정답률 차이를 나타내는 지표입니다. "
                "계산식은 `변별도 = 상위 집단 정답률 - 하위 집단 정답률`입니다. "
                "값이 클수록 성취 수준이 높은 학생과 낮은 학생을 잘 구분한 문항으로 볼 수 있습니다. "
                "일반적으로 40% 이상이면 변별도가 높은 문항, 20~40%는 어느 정도 변별력이 있는 문항, "
                "20% 미만은 변별력이 낮은 문항으로 참고할 수 있습니다.\n\n"
                "다만 변별도가 지나치게 높은 경우에도 단순히 좋은 문항이라고만 해석하기보다는 검토가 필요합니다. "
                "70% 이상처럼 상위 집단과 하위 집단의 정답률 차이가 매우 크게 나타나는 문항은 "
                "성취 수준을 강하게 구분한 문항일 수 있지만, 하위 성취 학생에게 지나치게 어려웠는지, "
                "특정 개념 결손이 크게 작용했는지, 발문이나 선택지가 일부 학생에게 과도한 혼란을 주었는지 "
                "함께 살펴볼 필요가 있습니다. 따라서 변별도는 정답률, 난이도, 선택지 반응, 평가 내용과 함께 "
                "종합적으로 해석하는 것이 좋습니다."
            )
            item_display = fmt_percent_df(analysis["item"].sort_values("정답률"))
            st.dataframe(style_item_analysis_df(item_display), use_container_width=True, height=760, hide_index=True)

        st.markdown("#### 예상 난이도-실제 정답률 일치 여부")

        with st.container(border=True):
            st.markdown("##### 실제 난이도 판정 기준")
            기준_col1, 기준_col2, 기준_col3 = st.columns([1.2, 1.2, 1.6])
            with 기준_col1:
                hard_cut = st.number_input(
                    "어려움/보통 난이도 구분 정답률(%)",
                    min_value=0,
                    max_value=100,
                    value=33,
                    step=1,
                    key="difficulty_hard_cut_input",
                )
            with 기준_col2:
                easy_cut = st.number_input(
                    "보통/쉬움 난이도 구분 정답률(%)",
                    min_value=0,
                    max_value=100,
                    value=66,
                    step=1,
                    key="difficulty_easy_cut_input",
                )
            with 기준_col3:
                st.markdown(
                    f"""
                    <div style="height:100%; min-height:74px; display:flex; align-items:center; padding:0.6rem 0.2rem; color:#475569; font-size:0.92rem; line-height:1.55;">
                        <div>
                            <b>어려움</b> &lt; {int(hard_cut)}% &nbsp;·&nbsp;
                            <b>보통</b> {int(hard_cut)}% 이상 {int(easy_cut)}% 미만 &nbsp;·&nbsp;
                            <b>쉬움</b> {int(easy_cut)}% 이상
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

        if int(hard_cut) >= int(easy_cut):
            st.warning("보통/쉬움 난이도 구분 정답률은 어려움/보통 난이도 구분 정답률보다 커야 합니다. 현재는 기본값 33%, 66%로 계산합니다.")
            calc_hard_cut, calc_easy_cut = 33.0, 66.0
        else:
            calc_hard_cut, calc_easy_cut = float(hard_cut), float(easy_cut)

        difficulty_gap_df = make_difficulty_gap_analysis(analysis["item"], hard_cut_percent=calc_hard_cut, easy_cut_percent=calc_easy_cut, expected_question_df=parsed.original_question_df)
        analysis["difficulty_gap"] = difficulty_gap_df

        match_count = int((difficulty_gap_df["차이해석"] == "일치").sum()) if not difficulty_gap_df.empty else 0
        harder_count = int((difficulty_gap_df["차이해석"] == "예상보다 어려웠음").sum()) if not difficulty_gap_df.empty else 0
        easier_count = int((difficulty_gap_df["차이해석"] == "예상보다 쉬웠음").sum()) if not difficulty_gap_df.empty else 0
        mismatch_count = harder_count + easier_count

        st.markdown(
            f"""
            <style>
            .difficulty-summary-grid {{
                display: grid;
                grid-template-columns: repeat(4, minmax(120px, 1fr));
                gap: 0.7rem;
                margin: 0.85rem 0 1.05rem 0;
            }}
            .difficulty-summary-card {{
                border: 1px solid #E5E7EB;
                border-radius: 14px;
                padding: 0.85rem 1rem;
                background: linear-gradient(180deg, #FFFFFF 0%, #F8FAFC 100%);
                box-shadow: 0 1px 2px rgba(15, 23, 42, 0.05);
            }}
            .difficulty-summary-label {{
                color: #64748B;
                font-size: 0.86rem;
                font-weight: 600;
                margin-bottom: 0.25rem;
                white-space: nowrap;
            }}
            .difficulty-summary-value {{
                color: #0F172A;
                font-size: 1.85rem;
                font-weight: 800;
                line-height: 1.05;
            }}
            .difficulty-summary-subtle {{ border-left: 5px solid #CBD5E1; }}
            .difficulty-summary-hard {{ border-left: 5px solid #EF4444; }}
            .difficulty-summary-easy {{ border-left: 5px solid #3B82F6; }}
            .difficulty-summary-match {{ border-left: 5px solid #22C55E; }}
            @media (max-width: 900px) {{
                .difficulty-summary-grid {{ grid-template-columns: repeat(2, minmax(120px, 1fr)); }}
            }}
            </style>
            <div class="difficulty-summary-grid">
                <div class="difficulty-summary-card difficulty-summary-match">
                    <div class="difficulty-summary-label">일치 문항수(개)</div>
                    <div class="difficulty-summary-value">{match_count}</div>
                </div>
                <div class="difficulty-summary-card difficulty-summary-subtle">
                    <div class="difficulty-summary-label">불일치 문항수(개)</div>
                    <div class="difficulty-summary-value">{mismatch_count}</div>
                </div>
                <div class="difficulty-summary-card difficulty-summary-hard">
                    <div class="difficulty-summary-label">예상보다 어려운 문항수(개)</div>
                    <div class="difficulty-summary-value">{harder_count}</div>
                </div>
                <div class="difficulty-summary-card difficulty-summary-easy">
                    <div class="difficulty-summary-label">예상보다 쉬운 문항수(개)</div>
                    <div class="difficulty-summary-value">{easier_count}</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        gap_view_df = difficulty_gap_df.sort_values("문항번호", ascending=True).copy()
        display_gap = gap_view_df[[c for c in ["문항번호", "평가영역", "예상난이도", "기대정답률구간", "정답률_pct", "차이해석"] if c in gap_view_df.columns]].copy()
        if "정답률_pct" in display_gap.columns:
            display_gap["실제정답률(%)"] = display_gap.pop("정답률_pct").map(lambda x: "" if pd.isna(x) else f"{x:.1f}%")
        if "차이해석" in display_gap.columns:
            display_gap = display_gap.rename(columns={"차이해석": "판정"})
        with st.container(border=True):
            st.dataframe(display_gap, use_container_width=True, height=760, hide_index=True)

    elif selected_analysis_tab == "학급별 분석":
        with st.container(border=True):
            st.markdown("#### 학급별 문항 분석")
            st.markdown(
                "문항별 전체 정답률과 학급별 정답률을 함께 확인합니다. "
                "특히 학급 간 정답률 차이를 보면 특정 반에서 유독 어려워한 문항을 찾을 수 있습니다."
            )

            class_item_pivot = analysis["class_item_pivot"].copy()
            item_meta_cols = [c for c in ["문항번호", "평가영역", "성취기준", "난이도", "배점", "정답", "정답률", "변별도"] if c in analysis["item"].columns]
            class_item_view = analysis["item"][item_meta_cols].merge(class_item_pivot, on="문항번호", how="left")
            class_rate_cols = [c for c in class_item_view.columns if str(c).endswith("반_정답률") or str(c) == "미상반_정답률"]
            if class_rate_cols:
                class_item_view["학급간최대차"] = class_item_view[class_rate_cols].max(axis=1) - class_item_view[class_rate_cols].min(axis=1)
            class_item_display = fmt_percent_df(class_item_view.sort_values("문항번호", ascending=True))
            st.dataframe(style_class_item_analysis_df(class_item_display), use_container_width=True, height=520, hide_index=True)

        with st.container(border=True):
            st.markdown("#### 학급 간 정답률 차이가 큰 문항")
            st.markdown("학급별 정답률의 최고값과 최저값 차이가 큰 문항을 먼저 보여줍니다. 학급별 수업 흐름, 활동 경험, 오개념 차이를 확인할 때 유용합니다.")
            gap_rows = []
            if class_rate_cols:
                for _, row in class_item_view.iterrows():
                    rates = row[class_rate_cols].dropna()
                    if rates.empty:
                        continue
                    max_col = rates.idxmax()
                    min_col = rates.idxmin()
                    gap_rows.append({
                        "문항번호": row.get("문항번호", ""),
                        "평가영역": row.get("평가영역", ""),
                        "성취기준": row.get("성취기준", ""),
                        "난이도": row.get("난이도", ""),
                        "전체정답률": row.get("정답률", np.nan),
                        "최고학급": str(max_col).replace("_정답률", ""),
                        "최고정답률": rates[max_col],
                        "최저학급": str(min_col).replace("_정답률", ""),
                        "최저정답률": rates[min_col],
                        "학급간차이": rates[max_col] - rates[min_col],
                    })
            gap_df = pd.DataFrame(gap_rows)
            if not gap_df.empty:
                st.dataframe(fmt_percent_df(gap_df.sort_values("학급간차이", ascending=False)), use_container_width=True, height=360, hide_index=True)
            else:
                st.info("학급 간 정답률 차이를 계산할 데이터가 없습니다.")

        st.markdown("---")

    elif selected_analysis_tab == "평가영역별 분석":
        with st.container(border=True):
            st.markdown("#### 평가영역별 분석")
            st.info(
                "환산점수는 각 평가영역에서 얻은 점수를 100점 만점 기준으로 바꾸어 표시한 값입니다. "
                "서로 배점이 다른 평가영역을 같은 기준에서 비교하기 위한 값이며, 이미 해당 시험의 만점이 100점인 경우에는 원점수와 같은 값으로 표시됩니다."
            )
            domain_display_cols = [c for c in ["평가영역", "문항수", "배점합계", "평균점수", "환산평균", "정답률"] if c in analysis["domain"].columns]
            domain_display = fmt_percent_df(analysis["domain"][domain_display_cols].sort_values("정답률"))
            st.dataframe(style_domain_analysis_df(domain_display), use_container_width=True, height=760, hide_index=True)
        with st.container(border=True):
            st.markdown("#### 개인별 평가영역 분석")
            selected_student_domain = render_student_selector(analysis["individual"], "domain_individual")
            if selected_student_domain:
                domain_one = analysis["domain_scores"][analysis["domain_scores"]["반/번호"] == selected_student_domain].copy()
                domain_cols = [c for c in ["평가영역", "영역점수", "영역배점", "영역환산점수", "영역정답률"] if c in domain_one.columns]
                domain_one_display = fmt_percent_df(domain_one[domain_cols].sort_values("평가영역"))
                st.dataframe(style_domain_analysis_df(domain_one_display), use_container_width=True, height=760, hide_index=True)

    elif selected_analysis_tab == "성취기준별 분석":
        with st.container(border=True):
            st.markdown("#### 성취기준별 분석")
            standard_cols = [c for c in ["성취기준", "평가영역", "문항번호", "문항수", "배점합계", "평균점수", "환산평균", "정답률"] if c in analysis["standard"].columns]
            standard_display = fmt_percent_df(analysis["standard"][standard_cols].sort_values("성취기준"))
            st.dataframe(style_standard_analysis_df(standard_display), use_container_width=True, height=520, hide_index=True)
        with st.container(border=True):
            st.markdown("#### 개인별 성취기준 분석")
            selected_student_standard = render_student_selector(analysis["individual"], "standard_individual")
            if selected_student_standard:
                standard_one = analysis["standard_scores"][analysis["standard_scores"]["반/번호"] == selected_student_standard].copy()
                standard_one = standard_one.merge(
                    analysis["standard"][["성취기준", "평가영역", "문항번호"]],
                    on="성취기준",
                    how="left",
                )
                standard_one = standard_one.rename(columns={"평가영역_y": "평가영역", "평가영역_x": "학생평가영역"})
                standard_cols_one = [c for c in ["성취기준", "평가영역", "문항번호", "성취기준점수", "성취기준배점", "성취기준환산점수", "성취기준정답률"] if c in standard_one.columns]
                standard_one_display = fmt_percent_df(standard_one[standard_cols_one].sort_values("성취기준"))
                st.dataframe(style_standard_analysis_df(standard_one_display), use_container_width=True, height=360, hide_index=True)

    elif selected_analysis_tab == "성취수준별 분석":
        with st.container(border=True):
            st.markdown("#### 성취수준별 문항 분석")
            st.info(
                "A~E 열은 각 성취수준 학생들의 문항별 정답률입니다. "
                "기본 정렬은 문항번호 순이며, 해당 성취수준 학생이 없는 경우에는 '-'로 표시됩니다."
            )
            level_df = analysis.get("level_item", pd.DataFrame()).copy()
            if level_df.empty:
                st.warning("성취수준별 문항 분석에 표시할 데이터가 없습니다. 성취수준 분할점수와 학생 총점이 정상적으로 산출되었는지 확인해 주세요.")
            else:
                if "문항번호" in level_df.columns:
                    level_df["문항번호"] = pd.to_numeric(level_df["문항번호"], errors="coerce")
                    level_df = level_df.sort_values("문항번호", ascending=True)
                preferred_level_cols = ["평가영역", "문항번호", "정답률", "A", "B", "C", "D", "E", "수준간격차"]
                level_df = level_df[[c for c in preferred_level_cols if c in level_df.columns] + [c for c in level_df.columns if c not in preferred_level_cols]]
                level_display = fmt_percent_df(level_df)
                for level_col in list("ABCDE"):
                    if level_col in level_display.columns:
                        level_display[level_col] = level_display[level_col].replace("", "-").fillna("-")
                st.dataframe(style_level_item_analysis_df(level_display), use_container_width=True, height=760, hide_index=True)

    elif selected_analysis_tab == "학생 개별":
        with st.container(border=True):
            st.markdown("학생 이름은 웹앱 내부 확인용입니다. AI 분석에 보낼 때는 익명화 옵션을 권장합니다.")
            individual_view = analysis["individual"].sort_values(["반", "번호"]).copy()
            hide_cols = ["선택형계산점수", "최종점수", "계산점수", "성취수준기준점수"]
            individual_view = individual_view.drop(columns=[c for c in hide_cols if c in individual_view.columns])
            if "환산점수" in individual_view.columns:
                individual_view = individual_view.rename(columns={"환산점수": "100점 만점 환산점수"})
            preferred_individual_cols = [
                "반/번호", "반", "번호", "이름", "선택형점수", "서답형점수", "기타점수",
                "영역총점", "100점 만점 환산점수", "성취수준", "오답문항"
            ]
            individual_view = individual_view[[c for c in preferred_individual_cols if c in individual_view.columns] + [c for c in individual_view.columns if c not in preferred_individual_cols]]
            st.dataframe(style_individual_analysis_df(fmt_percent_df(individual_view)), use_container_width=True, height=420, hide_index=True)
        with st.container(border=True):
            student_options = analysis["individual"].sort_values(["반", "번호"])["반/번호"].tolist()
            selected_student = st.selectbox("학생 선택", student_options)
            one_long = analysis["long"][analysis["long"]["반/번호"] == selected_student]
            st.dataframe(fmt_percent_df(one_long[["문항번호", "평가영역", "난이도", "배점", "정답", "원본표시", "선택지", "정오", "점수", "성취기준"]]), use_container_width=True, height=360)

    elif selected_analysis_tab == "AI 분석":
        with st.container(border=True):
            st.markdown("#### AI 분석")
            st.markdown("기본 분석은 원안지 없이 통계 자료를 바탕으로 해석하고, 고급 분석은 원안지 PDF를 함께 활용하는 구조입니다.")
            api_key = st.text_input("OpenAI API Key", type="password")
            model = st.text_input("모델", value="gpt-4o-mini")

        ai_section_mode = st.radio(
            "AI 분석 종류",
            ["기본 분석: 통계 기반 해석", "고급 분석: 원안지 기반 심층 해석"],
            horizontal=True,
            key="ai_section_mode",
        )

        if ai_section_mode == "기본 분석: 통계 기반 해석":
            with st.container(border=True):
                st.markdown("#### 기본 분석: 통계 기반 해석")
                st.markdown("원안지 PDF 없이 현재 분석 데이터만으로 전체 경향, 취약 영역, 문항별 이상 신호, 학생 개별 피드백을 생성합니다. 결과는 평가 정보가 포함된 Word 문서로 저장할 수 있습니다.")
                basic_mode = st.radio(
                    "기본 분석 유형",
                    ["전체 통계 분석", "학생 개별 분석"],
                    horizontal=True,
                    key="basic_ai_mode",
                )
                anonymize = st.checkbox("학생 개별 분석에서 이름을 API로 보내지 않기", value=True, key="basic_ai_anonymize")
                if basic_mode == "학생 개별 분석":
                    student_options = analysis["individual"].sort_values(["반", "번호"])["반/번호"].tolist()
                    ai_student = st.selectbox("AI 분석 대상 학생", student_options, key="basic_ai_student")
                    base_basic_prompt = build_individual_ai_prompt(parsed, analysis, ai_student, anonymize=anonymize)
                    basic_download_name = "AI_기본분석_학생개별.docx"
                    basic_prompt_mode_key = "basic_individual_prompt_mode"
                    basic_base_preview_key = "basic_individual_base_prompt_preview"
                    basic_custom_key = "basic_individual_direct_prompt"
                    basic_analysis_label = "기본 분석: 학생 개별 분석"
                else:
                    base_basic_prompt = build_basic_statistics_ai_prompt(parsed, analysis)
                    basic_download_name = "AI_기본분석_통계기반해석.docx"
                    basic_prompt_mode_key = "basic_overall_prompt_mode"
                    basic_base_preview_key = "basic_overall_base_prompt_preview"
                    basic_custom_key = "basic_overall_direct_prompt"
                    basic_analysis_label = "기본 분석: 전체 통계 분석"

                basic_prompt_mode = st.radio(
                    "프롬프트 사용 방식",
                    PROMPT_MODE_OPTIONS,
                    horizontal=True,
                    key=basic_prompt_mode_key,
                    help=PROMPT_MODE_HELP,
                )
                with st.expander("웹앱 기본 프롬프트 보기", expanded=False):
                    st.text_area("웹앱 기본 프롬프트", extract_prompt_instruction_preview(base_basic_prompt), height=420)

                basic_focus_request = ""
                basic_custom_prompt = ""
                if basic_prompt_mode == PROMPT_MODE_WITH_FOCUS:
                    if basic_mode == "학생 개별 분석":
                        basic_focus_request = render_focus_request_with_preset(
                            "중점 분석 요청 프리셋",
                            BASIC_INDIVIDUAL_FOCUS_PRESETS,
                            "basic_individual_focus_preset",
                            "basic_individual_focus_request",
                            "예: 학생 맞춤형 피드백, 다음 시험 준비 전략, 보충 학습이 필요한 개념 등을 중점적으로 분석해줘.",
                            "프리셋을 고르면 아래 요청 문구가 자동 입력됩니다. 필요하면 직접 수정할 수 있습니다. 비워두면 기본 프롬프트만 사용합니다.",
                        )
                    else:
                        basic_focus_request = render_focus_request_with_preset(
                            "중점 분석 요청 프리셋",
                            BASIC_OVERALL_FOCUS_PRESETS,
                            "basic_overall_focus_preset",
                            "basic_overall_focus_request",
                            "예: 정답률이 낮은 문항, 오답 선택지 반응, 학급 간 차이, 평가 타당도 관점 등을 중점적으로 분석해줘.",
                            "프리셋을 고르면 아래 요청 문구가 자동 입력됩니다. 필요하면 직접 수정할 수 있습니다. 비워두면 기본 프롬프트만 사용합니다.",
                        )
                elif basic_prompt_mode == PROMPT_MODE_DIRECT:
                    if basic_mode == "학생 개별 분석":
                        basic_custom_prompt = render_focus_request_with_preset(
                            "중점 분석 요청 프리셋",
                            BASIC_INDIVIDUAL_FOCUS_PRESETS,
                            "basic_individual_direct_preset",
                            basic_custom_key,
                            "예: 학생 맞춤형 피드백, 다음 시험 준비 전략, 보충 학습이 필요한 개념 등을 중심으로 분석해줘.",
                            "직접 작성 모드에서도 프리셋을 초안으로 불러올 수 있습니다. 프리셋을 고르면 아래 분석 지시문에 자동 입력되며, 원하는 대로 수정할 수 있습니다. 이 모드에서는 웹앱의 기본 분석 목차를 사용하지 않고, 여기에 작성한 지시문과 앱 분석 데이터만 AI에 전달합니다.",
                            height=170,
                            text_area_label="AI에게 직접 전달할 분석 지시문",
                        )
                    else:
                        basic_custom_prompt = render_focus_request_with_preset(
                            "중점 분석 요청 프리셋",
                            BASIC_OVERALL_FOCUS_PRESETS,
                            "basic_overall_direct_preset",
                            basic_custom_key,
                            "예: 전체 분석은 하지 말고, 정답률이 50% 미만인 문항만 표로 정리하고 학생들이 어려워했을 가능성을 간단히 분석해줘.",
                            "직접 작성 모드에서도 프리셋을 초안으로 불러올 수 있습니다. 프리셋을 고르면 아래 분석 지시문에 자동 입력되며, 원하는 대로 수정할 수 있습니다. 이 모드에서는 웹앱의 기본 분석 목차를 사용하지 않고, 여기에 작성한 지시문과 앱 분석 데이터만 AI에 전달합니다.",
                            height=170,
                            text_area_label="AI에게 직접 전달할 분석 지시문",
                        )
                else:
                    st.caption("웹앱 기본 프롬프트만 사용합니다. 추가 의뢰나 직접 작성 프롬프트는 붙지 않습니다.")

                basic_prompt = compose_ai_prompt_by_mode(
                    base_basic_prompt,
                    basic_prompt_mode,
                    focus_request=basic_focus_request,
                    custom_prompt=basic_custom_prompt,
                    analysis_label=basic_analysis_label,
                )
                basic_report_focus_request = basic_focus_request if basic_prompt_mode == PROMPT_MODE_WITH_FOCUS else (basic_custom_prompt if basic_prompt_mode == PROMPT_MODE_DIRECT else "")

            with st.container(border=True):
                with st.expander("최종 AI 전달 프롬프트 확인", expanded=False):
                    st.text_area("최종 AI 전달 프롬프트", basic_prompt, height=420)
                result_state_key = "basic_ai_result_text"
                docx_state_key = "basic_ai_docx_bytes"
                filename_state_key = "basic_ai_docx_filename"
                meta_state_key = "basic_ai_result_meta"
                context_state_key = "basic_ai_result_context"
                basic_context = hashlib.sha1(f"{basic_mode}|{basic_prompt}".encode("utf-8")).hexdigest()

                # 분석 유형·학생·프롬프트가 바뀌면 이전 결과를 현재 화면에서 제거한다.
                # 다운로드 버튼 클릭으로 인한 재실행에서는 context가 같으므로 결과가 유지된다.
                if st.session_state.get(context_state_key) and st.session_state.get(context_state_key) != basic_context:
                    for _key in [result_state_key, docx_state_key, filename_state_key, meta_state_key, context_state_key]:
                        st.session_state.pop(_key, None)

                basic_just_ran = False
                if st.button("기본 분석 실행", type="primary", key="run_basic_ai"):
                    if not api_key:
                        st.error("OpenAI API Key를 입력하세요.")
                    elif basic_prompt_mode == PROMPT_MODE_DIRECT and not basic_custom_prompt.strip():
                        st.error("직접 작성한 프롬프트만 사용하려면 분석 지시문을 입력하세요.")
                    else:
                        try:
                            # 새 분석을 시작할 때만 기존 결과를 지우고, 다운로드 버튼 클릭으로 인한 rerun에서는 유지한다.
                            st.session_state.pop(result_state_key, None)
                            st.session_state.pop(docx_state_key, None)
                            st.session_state.pop(filename_state_key, None)
                            st.session_state.pop(meta_state_key, None)
                            st.session_state.pop(context_state_key, None)
                            basic_just_ran = True

                            st.markdown("#### 기본 분석 결과")
                            st.caption("분석 결과가 생성되는 대로 아래에 실시간으로 표시됩니다.")
                            result_placeholder = st.empty()
                            result = call_openai_stream(api_key, model, basic_prompt, result_placeholder)

                            if len(result.strip()) < 300:
                                st.warning("AI 응답이 예상보다 짧습니다. 모델 출력 제한, API 오류, 프롬프트 입력량을 확인해 주세요.")

                            report_title = "성취수준별 평가결과 AI 분석 보고서"
                            report_type = f"기본 분석: {basic_mode}"
                            docx_bytes = make_ai_report_docx(parsed, analysis, report_title, result, report_type=report_type, focus_request=basic_report_focus_request)

                            st.session_state[result_state_key] = result
                            st.session_state[docx_state_key] = docx_bytes
                            st.session_state[filename_state_key] = basic_download_name
                            st.session_state[meta_state_key] = report_type
                            st.session_state[context_state_key] = basic_context
                            st.download_button(
                                "기본 분석 결과 Word 다운로드",
                                docx_bytes,
                                basic_download_name,
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="download_basic_ai_docx_fresh",
                            )
                        except Exception as e:
                            st.error(f"AI 분석 중 오류가 발생했습니다: {e}")

                if st.session_state.get(result_state_key) and not basic_just_ran:
                    st.markdown("#### 저장된 기본 분석 결과")
                    st.markdown(st.session_state[result_state_key])
                    st.download_button(
                        "기본 분석 결과 Word 다운로드",
                        st.session_state[docx_state_key],
                        st.session_state.get(filename_state_key, basic_download_name),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="download_basic_ai_docx_persisted",
                    )

        elif ai_section_mode == "고급 분석: 원안지 기반 심층 해석":
            with st.container(border=True):
                st.markdown("#### 고급 분석: 원안지 기반 심층 해석")
                st.markdown("원안지 PDF를 업로드한 뒤, 프롬프트 사용 방식과 필요 시 중점 분석 요청 프리셋을 고르고 문항 내용 기반 심층 분석을 생성합니다.")
                source_pdf = st.file_uploader("원안지 PDF 업로드", type=["pdf"], key="source_exam_pdf")

                advanced_prompt_mode = st.radio(
                    "프롬프트 사용 방식",
                    PROMPT_MODE_OPTIONS,
                    horizontal=True,
                    key="advanced_prompt_mode_v197",
                    help=PROMPT_MODE_HELP,
                )

                advanced_focus_request = ""
                advanced_custom_prompt = ""
                selected_advanced_preset = "직접 입력"
                if advanced_prompt_mode == PROMPT_MODE_WITH_FOCUS:
                    advanced_focus_request = render_focus_request_with_preset(
                        "중점 분석 요청 프리셋",
                        ADVANCED_FOCUS_PRESETS,
                        "advanced_focus_preset_v188",
                        "advanced_focus_request_v188",
                        "예: 문항 완성도·엄밀성, 오답 선택지 매력도, 발문 명확성, 성취기준 적합성 등을 중점적으로 분석해줘.",
                        "원안지 기반 고급 분석 전용 프리셋입니다. 문항 설계, 발문, 선택지, 정답 근거, 문항 완성도·엄밀성 등 원안지 PDF를 함께 보며 분석할 관점을 고를 수 있습니다. 프리셋을 고르면 아래 요청 문구가 자동 입력되고, 필요하면 직접 수정할 수 있습니다.",
                    )
                    selected_advanced_preset = st.session_state.get("advanced_focus_preset_v188", "직접 입력")
                    if selected_advanced_preset != "직접 입력":
                        st.caption(f"선택한 분석 관점: {selected_advanced_preset}")
                elif advanced_prompt_mode == PROMPT_MODE_DIRECT:
                    selected_advanced_preset = st.selectbox(
                        "중점 분석 요청 프리셋",
                        list(ADVANCED_FOCUS_PRESETS.keys()),
                        key="advanced_direct_preset_v199",
                        help="직접 작성 모드에서도 프리셋을 초안으로 불러올 수 있습니다. 프리셋을 고르면 아래 분석 지시문에 자동 입력되며, 원하는 대로 수정할 수 있습니다.",
                    )
                    direct_preset_last_key = "advanced_direct_preset_v199_last_value"
                    if st.session_state.get(direct_preset_last_key) != selected_advanced_preset:
                        st.session_state["advanced_direct_prompt_v197"] = ADVANCED_FOCUS_PRESETS.get(selected_advanced_preset, "")
                        st.session_state[direct_preset_last_key] = selected_advanced_preset
                    if selected_advanced_preset != "직접 입력":
                        st.caption(f"선택한 분석 관점: {selected_advanced_preset}")
                elif advanced_prompt_mode == PROMPT_MODE_DEFAULT:
                    st.caption("웹앱 기본 프롬프트만 사용합니다. 추가 의뢰나 직접 작성 프롬프트는 붙지 않습니다.")

                advanced_scope = infer_advanced_scope_from_preset(selected_advanced_preset)

                max_q = int(pd.to_numeric(parsed.question_df["문항번호"], errors="coerce").max()) if not parsed.question_df.empty else None
                need_item_input = advanced_scope in ["주요 문항 집중 분석", "오답 선택지 분석"]
                item_number_raw = st.text_input(
                    "분석할 문항번호",
                    value="",
                    placeholder="예: 3, 7, 12 또는 3-7",
                    help="비워두면 통계 결과를 기준으로 AI가 우선 분석 문항을 선정합니다.",
                    key="advanced_item_numbers",
                )
                selected_item_numbers = parse_question_number_input(item_number_raw, max_question=max_q)
                if item_number_raw.strip() and not selected_item_numbers:
                    st.warning("입력한 문항번호를 인식하지 못했습니다. 예: 3, 7, 12 또는 3-7 형식으로 입력하세요.")
                elif selected_item_numbers:
                    st.caption("선택 문항: " + ", ".join(f"{n}번" for n in selected_item_numbers))
                elif need_item_input:
                    st.caption("문항번호를 비워두면 앱의 통계 결과를 바탕으로 우선 분석 문항을 AI가 선정합니다.")

                advanced_student_key = None

                pdf_name = source_pdf.name if source_pdf is not None else "원안지 PDF 미업로드"
                base_advanced_prompt = build_advanced_exam_ai_prompt(
                    parsed,
                    analysis,
                    pdf_name=pdf_name,
                    scope=advanced_scope,
                    item_numbers=selected_item_numbers,
                    student_key=advanced_student_key,
                    anonymize_student=True,
                )
                with st.expander("웹앱 기본 프롬프트 보기", expanded=False):
                    st.text_area("웹앱 기본 프롬프트", extract_prompt_instruction_preview(base_advanced_prompt), height=460)

                if advanced_prompt_mode == PROMPT_MODE_DIRECT:
                    advanced_custom_prompt = st.text_area(
                        "AI에게 직접 전달할 분석 지시문",
                        placeholder="예: 전체 문항 설계 분석은 생략하고, 원안지에서 정답률이 낮은 문항의 발문과 선택지만 간단히 검토해줘.",
                        height=170,
                        key="advanced_direct_prompt_v197",
                        help="직접 작성 모드에서는 웹앱의 기본 분석 목차를 사용하지 않고, 여기에 작성한 지시문과 앱 분석 데이터, 원안지 PDF만 AI에 전달합니다. 위 프리셋을 초안으로 불러온 뒤 원하는 분석 방향으로 수정할 수 있습니다.",
                    )

                advanced_prompt = compose_ai_prompt_by_mode(
                    base_advanced_prompt,
                    advanced_prompt_mode,
                    focus_request=advanced_focus_request,
                    custom_prompt=advanced_custom_prompt,
                    analysis_label=f"고급 분석: {advanced_scope}",
                )
                advanced_report_focus_request = advanced_focus_request if advanced_prompt_mode == PROMPT_MODE_WITH_FOCUS else (advanced_custom_prompt if advanced_prompt_mode == PROMPT_MODE_DIRECT else "")

            with st.container(border=True):
                with st.expander("최종 AI 전달 프롬프트 확인", expanded=False):
                    st.text_area("최종 AI 전달 프롬프트", advanced_prompt, height=460)

                adv_result_state_key = "advanced_ai_result_text"
                adv_docx_state_key = "advanced_ai_docx_bytes"
                adv_filename_state_key = "advanced_ai_docx_filename"
                adv_meta_state_key = "advanced_ai_result_meta"
                adv_context_state_key = "advanced_ai_result_context"
                pdf_context_name = source_pdf.name if source_pdf is not None else ""
                advanced_context = hashlib.sha1(f"{advanced_scope}|{item_number_raw}|{advanced_student_key}|{pdf_context_name}|{advanced_prompt}".encode("utf-8")).hexdigest()

                # 프리셋·문항번호·PDF가 바뀌면 이전 고급 분석 결과를 현재 화면에서 제거한다.
                if st.session_state.get(adv_context_state_key) and st.session_state.get(adv_context_state_key) != advanced_context:
                    for _key in [adv_result_state_key, adv_docx_state_key, adv_filename_state_key, adv_meta_state_key, adv_context_state_key]:
                        st.session_state.pop(_key, None)

                advanced_just_ran = False
                if source_pdf is None:
                    st.info("원안지 PDF를 업로드하면 고급 분석을 실행할 수 있습니다.")

                if st.button("고급 분석 실행", type="primary", key="run_advanced_ai", disabled=(source_pdf is None)):
                    if not api_key:
                        st.error("OpenAI API Key를 입력하세요.")
                    elif advanced_prompt_mode == PROMPT_MODE_DIRECT and not advanced_custom_prompt.strip():
                        st.error("직접 작성한 프롬프트만 사용하려면 분석 지시문을 입력하세요.")
                    else:
                        try:
                            st.session_state.pop(adv_result_state_key, None)
                            st.session_state.pop(adv_docx_state_key, None)
                            st.session_state.pop(adv_filename_state_key, None)
                            st.session_state.pop(adv_meta_state_key, None)
                            st.session_state.pop(adv_context_state_key, None)
                            advanced_just_ran = True

                            pdf_bytes = source_pdf.getvalue()
                            st.markdown("#### 고급 분석 결과")
                            st.caption("원안지 PDF와 통계 자료를 함께 분석하며, 결과가 생성되는 대로 실시간으로 표시됩니다.")
                            result_placeholder = st.empty()
                            result = call_openai_stream_with_pdf(api_key, model, advanced_prompt, pdf_bytes, source_pdf.name, result_placeholder)

                            if len(result.strip()) < 300:
                                st.warning("AI 응답이 예상보다 짧습니다. PDF 인식, 모델 출력 제한, API 오류, 프롬프트 입력량을 확인해 주세요.")

                            report_title = "성취수준별 평가결과 원안지 기반 고급 분석 보고서"
                            report_type = f"고급 분석: {advanced_scope}"
                            docx_bytes = make_ai_report_docx(parsed, analysis, report_title, result, report_type=report_type, focus_request=advanced_report_focus_request)

                            st.session_state[adv_result_state_key] = result
                            st.session_state[adv_docx_state_key] = docx_bytes
                            st.session_state[adv_filename_state_key] = "AI_고급분석_원안지기반심층해석.docx"
                            st.session_state[adv_meta_state_key] = report_type
                            st.session_state[adv_context_state_key] = advanced_context
                            st.download_button(
                                "고급 분석 결과 Word 다운로드",
                                docx_bytes,
                                "AI_고급분석_원안지기반심층해석.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="download_advanced_ai_docx_fresh",
                            )
                        except Exception as e:
                            st.error(f"고급 분석 중 오류가 발생했습니다: {e}")

                if st.session_state.get(adv_result_state_key) and not advanced_just_ran:
                    st.markdown("#### 저장된 고급 분석 결과")
                    st.markdown(st.session_state[adv_result_state_key])
                    st.download_button(
                        "고급 분석 결과 Word 다운로드",
                        st.session_state[adv_docx_state_key],
                        st.session_state.get(adv_filename_state_key, "AI_고급분석_원안지기반심층해석.docx"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="download_advanced_ai_docx_persisted",
                    )

    st.markdown("<div class='big-section-gap'></div>", unsafe_allow_html=True)
    render_step_header("4", "다운로드", "확인용 입력자료와 5종 분석 결과를 엑셀 파일로 내려받습니다.")
    d1, d2 = st.columns(2)
    confirm_bytes = make_confirm_excel(parsed, analysis)
    zip_bytes = make_analysis_zip(parsed, analysis)
    d1.download_button(
        "확인용 엑셀 다운로드",
        confirm_bytes,
        file_name="확인용_분석입력자료.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )
    d2.download_button(
        "5종 분석 엑셀 ZIP 다운로드",
        zip_bytes,
        file_name="성취수준별_평가결과_분석_5종.zip",
        mime="application/zip",
        use_container_width=True,
    )


if __name__ == "__main__":
    main()
